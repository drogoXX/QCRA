# report_formatter.py
"""
Professional document formatting module for KVI Risk Assessment Reports.

Contains functions for styling tables, adding cover pages, table of contents,
headers/footers, KPI summary boxes, figure captions, and appendices.
"""

import os
from datetime import datetime
from typing import Optional, Dict, List, Any, Callable, Tuple

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from report_config import (
    KVI_COLORS, FONTS, FONT_SIZES, DOCUMENT_SETTINGS,
    GLOSSARY_TERMS, DEFAULT_ASSUMPTIONS, DEFAULT_LIMITATIONS,
    DEFAULT_DISCLAIMER, DEFAULT_APPROVAL_ROLES, CHART_CAPTIONS,
    VALID_PHASES
)
from utils import format_currency, format_percentage, hex_to_rgb


class FigureCounter:
    """Track figure numbers for captioning throughout document."""

    def __init__(self):
        self.count = 0

    def next(self) -> int:
        """Get next figure number."""
        self.count += 1
        return self.count

    def reset(self):
        """Reset counter for new document."""
        self.count = 0


class SectionCounter:
    """Track section numbers for document headings."""

    def __init__(self):
        self.level1 = 0
        self.level2 = 0
        self.level3 = 0

    def next_level1(self) -> str:
        """Get next level 1 section number."""
        self.level1 += 1
        self.level2 = 0
        self.level3 = 0
        return str(self.level1)

    def next_level2(self) -> str:
        """Get next level 2 section number."""
        self.level2 += 1
        self.level3 = 0
        return f"{self.level1}.{self.level2}"

    def next_level3(self) -> str:
        """Get next level 3 section number."""
        self.level3 += 1
        return f"{self.level1}.{self.level2}.{self.level3}"

    def reset(self):
        """Reset counters for new document."""
        self.level1 = 0
        self.level2 = 0
        self.level3 = 0


# Global counters
figure_counter = FigureCounter()
section_counter = SectionCounter()


# =============================================================================
# CELL STYLING FUNCTIONS
# =============================================================================

def set_cell_shading(cell, color_hex: str):
    """
    Set cell background color.

    Args:
        cell: python-docx table cell object
        color_hex: Hex color without '#' (e.g., '1F4E79')
    """
    color_hex = color_hex.lstrip('#')
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top: int = 50, bottom: int = 50, left: int = 100, right: int = 100):
    """
    Set cell margins in twips (1/20 of a point).

    Args:
        cell: python-docx table cell object
        top, bottom, left, right: Margin values in twips
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, color: str = 'D9D9D9', size: str = '4'):
    """
    Add thin borders to cell.

    Args:
        cell: python-docx table cell object
        color: Border color hex without '#'
        size: Border width
    """
    color = color.lstrip('#')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm: float):
    """Set row height in centimeters."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # Convert cm to twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


# =============================================================================
# TABLE FORMATTING
# =============================================================================

def apply_professional_table_style(
    table,
    has_total_row: bool = False,
    highlight_conditions: Optional[Dict[int, Tuple[Callable, str]]] = None
):
    """
    Apply professional KVI table styling.

    Args:
        table: python-docx Table object
        has_total_row: If True, style the last row as a total/summary row
        highlight_conditions: Dict with column index and (condition_func, bg_color) tuples
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_color = KVI_COLORS['table_header'].lstrip('#')
    alt_color = KVI_COLORS['table_alt_row'].lstrip('#')
    total_color = KVI_COLORS['total_row_bg'].lstrip('#')
    border_color = KVI_COLORS['table_border'].lstrip('#')

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Set cell margins
            set_cell_margins(cell, top=50, bottom=50, left=100, right=100)

            # Header row (first row)
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(FONT_SIZES['table_header'])

            # Total row (last row if has_total_row)
            elif has_total_row and i == len(table.rows) - 1:
                set_cell_shading(cell, total_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Alternating rows
            elif i % 2 == 0:
                set_cell_shading(cell, alt_color)
            else:
                set_cell_shading(cell, 'FFFFFF')

            # Apply conditional formatting if specified
            if highlight_conditions and j in highlight_conditions:
                condition_func, bg_color = highlight_conditions[j]
                cell_text = cell.text.strip()
                try:
                    if condition_func(cell_text):
                        set_cell_shading(cell, bg_color.lstrip('#'))
                except (ValueError, TypeError):
                    pass

            # Add thin borders
            set_cell_border(cell, color=border_color)


# =============================================================================
# HEADING STYLES
# =============================================================================

def configure_heading_styles(document: Document):
    """
    Configure heading styles with KVI branding.

    Args:
        document: python-docx Document object
    """
    styles = document.styles

    # Get colors
    dark_blue = hex_to_rgb(KVI_COLORS['dark'])
    primary_blue = hex_to_rgb(KVI_COLORS['primary'])
    neutral = hex_to_rgb(KVI_COLORS['neutral'])

    # Heading 1: Main sections
    h1 = styles['Heading 1']
    h1.font.name = FONTS['heading']
    h1.font.size = Pt(FONT_SIZES['heading1'])
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['table_header']))
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2: Subsections
    h2 = styles['Heading 2']
    h2.font.name = FONTS['heading']
    h2.font.size = Pt(FONT_SIZES['heading2'])
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(*primary_blue)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3: Sub-subsections
    h3 = styles['Heading 3']
    h3.font.name = FONTS['heading']
    h3.font.size = Pt(FONT_SIZES['heading3'])
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(*neutral)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Normal body text style
    normal = styles['Normal']
    normal.font.name = FONTS['body']
    normal.font.size = Pt(FONT_SIZES['body'])
    normal.font.color.rgb = RGBColor(*neutral)
    normal.paragraph_format.line_spacing = 1.15


# =============================================================================
# COVER PAGE
# =============================================================================

def add_cover_page(
    document: Document,
    config: Dict[str, Any],
    logo_path: Optional[str] = None
):
    """
    Add a professional cover page to the document.

    Args:
        document: python-docx Document object
        config: Configuration dict with keys:
            - project_name: Project name
            - contract_ref: Contract reference number
            - doc_number: Document number
            - revision: Revision letter/number
            - confidence_level: P50/P80/P95
            - classification: CONFIDENTIAL/INTERNAL
            - author: Author name
            - report_date: Report date string
        logo_path: Path to logo image file (optional)
    """
    # Add logo if provided - larger and properly positioned
    if logo_path and os.path.exists(logo_path):
        logo_para = document.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_para.add_run()
        try:
            logo_run.add_picture(logo_path, width=Cm(6))  # Larger logo
        except Exception:
            pass  # Skip if logo fails to load

    # Add spacing
    for _ in range(2):
        document.add_paragraph()

    # Main title
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('RISK ASSESSMENT REPORT')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['dark']))

    # Subtitle
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('Monte Carlo Simulation & Probabilistic\nRisk Analysis')
    sub_run.font.size = Pt(FONT_SIZES['subtitle'])
    sub_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    # Horizontal line
    document.add_paragraph()
    line_para = document.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run('─' * 60)
    line_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['primary']))

    document.add_paragraph()

    # Project details table
    details_table = document.add_table(rows=4, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    details = [
        ('Project:', config.get('project_name', 'PROTOS CC Carbon Capture')),
        ('Contract Ref:', config.get('contract_ref', '-')),
        ('Document No:', f"{config.get('doc_number', '-')} Rev {config.get('revision', 'A')}"),
        ('Confidence:', f"P{config.get('confidence_level', 80)}"),
    ]

    for i, (label, value) in enumerate(details):
        label_cell = details_table.cell(i, 0)
        value_cell = details_table.cell(i, 1)

        label_cell.text = label
        value_cell.text = value

        # Style label
        for para in label_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

        # Style value
        for para in value_cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['dark']))

    # Set column widths
    details_table.columns[0].width = Inches(1.8)
    details_table.columns[1].width = Inches(3.5)

    # Another horizontal line
    document.add_paragraph()
    line_para2 = document.add_paragraph()
    line_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run2 = line_para2.add_run('─' * 60)
    line_run2.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['primary']))

    document.add_paragraph()

    # Classification and date
    info_para = document.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    classification = config.get('classification', DOCUMENT_SETTINGS['default_classification'])
    report_date = config.get('report_date', datetime.now().strftime('%d %B %Y'))

    info_run = info_para.add_run(f"Classification: {classification}\nGenerated: {report_date}")
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    # Page break after cover (no document control here - moved to Appendix E)
    document.add_page_break()


# =============================================================================
# TABLE OF CONTENTS
# =============================================================================

def add_table_of_contents(document: Document):
    """
    Add a Table of Contents that updates when document is opened in Word.

    Args:
        document: python-docx Document object
    """
    # Add TOC title
    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = toc_title.add_run('TABLE OF CONTENTS')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['table_header']))

    document.add_paragraph()

    # Create TOC field
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Add placeholder text between separate and end
    placeholder_run = paragraph.add_run()
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = 'Press Ctrl+A then F9 to update Table of Contents, or right-click and select "Update Field"'
    placeholder_run._r.append(placeholder_text)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    paragraph._p.append(fldChar3)

    # Add styled instruction box
    document.add_paragraph()
    inst_box = document.add_table(rows=1, cols=1)
    inst_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    inst_cell = inst_box.cell(0, 0)
    set_cell_shading(inst_cell, KVI_COLORS['light'].lstrip('#'))
    set_cell_border(inst_cell, color=KVI_COLORS['primary'].lstrip('#'), size='4')
    set_cell_margins(inst_cell, top=100, bottom=100, left=150, right=150)

    inst_para = inst_cell.paragraphs[0]
    inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_run = inst_para.add_run("📋 To generate the Table of Contents:\n")
    inst_run.font.bold = True
    inst_run.font.size = Pt(10)
    inst_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['dark']))

    step_run = inst_para.add_run("1. Press Ctrl+A to select all\n2. Press F9 to update fields\n3. Or right-click above and select 'Update Field'")
    step_run.font.size = Pt(9)
    step_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    document.add_page_break()


# =============================================================================
# HEADERS AND FOOTERS
# =============================================================================

def add_page_number(paragraph):
    """Add 'Page X of Y' field to paragraph."""
    run = paragraph.add_run('Page ')
    run.font.size = Pt(FONT_SIZES['footer'])
    run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    # PAGE field
    run_page = paragraph.add_run()
    run_page.font.size = Pt(FONT_SIZES['footer'])

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    instrText1.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fldChar1)
    run_page._r.append(instrText1)
    run_page._r.append(fldChar2)

    run_of = paragraph.add_run(' of ')
    run_of.font.size = Pt(FONT_SIZES['footer'])
    run_of.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    # NUMPAGES field
    run_total = paragraph.add_run()
    run_total.font.size = Pt(FONT_SIZES['footer'])

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run_total._r.append(fldChar3)
    run_total._r.append(instrText2)
    run_total._r.append(fldChar4)


def add_header_footer(
    document: Document,
    project_name: str,
    doc_ref: str,
    report_date: str
):
    """
    Add running headers and footers to all sections.

    Args:
        document: python-docx Document object
        project_name: Project name for header
        doc_ref: Document reference for footer
        report_date: Report date for footer
    """
    neutral_color = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    for section in document.sections:
        # Different first page (no header/footer on cover)
        section.different_first_page_header_footer = True

        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()

        # Left-aligned project name
        run_left = header_para.add_run(project_name)
        run_left.font.size = Pt(FONT_SIZES['header'])
        run_left.font.color.rgb = neutral_color

        # Tab to right
        header_para.add_run('\t\t')

        # Right-aligned document title
        run_right = header_para.add_run('Risk Assessment Report')
        run_right.font.size = Pt(FONT_SIZES['header'])
        run_right.font.color.rgb = neutral_color
        run_right.font.italic = True

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()

        # Left: Doc reference
        run_left = footer_para.add_run(doc_ref)
        run_left.font.size = Pt(FONT_SIZES['footer'])
        run_left.font.color.rgb = neutral_color

        # Center: Page X of Y
        footer_para.add_run('\t')
        add_page_number(footer_para)

        # Right: Date
        footer_para.add_run('\t')
        run_date = footer_para.add_run(report_date)
        run_date.font.size = Pt(FONT_SIZES['footer'])
        run_date.font.color.rgb = neutral_color


# =============================================================================
# KPI SUMMARY BOX
# =============================================================================

def add_kpi_summary_box(document: Document, metrics: Dict[str, Any]):
    """
    Add a KPI summary callout box to the Executive Summary.

    Args:
        document: python-docx Document object
        metrics: dict with keys:
            - total_risks, threats, opportunities
            - initial_p80, residual_p80
            - risk_reduction_pct
            - risk_profile ('Low', 'Moderate', 'High', 'Very High')
    """
    # Create a 1-row, 1-column table as container
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)

    # Style the cell as a callout box
    set_cell_shading(cell, KVI_COLORS['kpi_box_bg'].lstrip('#'))
    set_cell_border(cell, color=KVI_COLORS['kpi_box_border'].lstrip('#'), size='12')
    set_cell_margins(cell, top=150, bottom=150, left=200, right=200)

    # Title
    p1 = cell.paragraphs[0]
    p1.clear()
    title_run = p1.add_run('KEY METRICS AT A GLANCE')
    title_run.font.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['dark']))
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metrics line 1 - Risk counts
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        f"Total Risks: {metrics.get('total_risks', 0)}  │  "
        f"Threats: {metrics.get('threats', 0)}  │  "
        f"Opportunities: {metrics.get('opportunities', 0)}"
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    # Metrics line 2 - P80 exposure
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    initial_p80 = metrics.get('initial_p80', 0)
    residual_p80 = metrics.get('residual_p80', 0)

    run3a = p3.add_run('P80 Exposure: ')
    run3a.font.size = Pt(11)
    run3a.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    run3b = p3.add_run(f"{format_currency(initial_p80)} → {format_currency(residual_p80)}")
    run3b.font.size = Pt(11)
    run3b.font.bold = True
    run3b.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['dark']))

    # Metrics line 3 - Risk reduction with color
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run4a = p4.add_run('Risk Reduction: ')
    run4a.font.size = Pt(11)
    run4a.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    risk_reduction = metrics.get('risk_reduction_pct', 0)
    run4b = p4.add_run(f"{format_percentage(risk_reduction)}")
    run4b.font.size = Pt(11)
    run4b.font.bold = True
    run4b.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['success']))

    # Risk profile indicator - using confidence level based profiles
    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    profile = metrics.get('risk_profile', 'Balanced / Moderate Risk Profile')

    # Map profile text to colors (consistent with confidence level section)
    profile_colors = {
        'Optimistic / High Risk Tolerance': ('#E67E22', '●'),  # Orange
        'Balanced / Moderate Risk Profile': (KVI_COLORS['primary'], '●'),  # Blue
        'Conservative / Low Risk Tolerance': (KVI_COLORS['success'], '●'),  # Green
        # Legacy support
        'Low': (KVI_COLORS['success'], '●'),
        'Moderate': (KVI_COLORS['warning'], '●'),
        'High': ('#E67E22', '●'),
        'Very High': (KVI_COLORS['danger'], '●'),
    }
    color, indicator = profile_colors.get(profile, (KVI_COLORS['neutral'], '●'))

    run5a = p5.add_run('Risk Profile: ')
    run5a.font.size = Pt(10)
    run5a.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    run5b = p5.add_run(f"{indicator} {profile}")
    run5b.font.size = Pt(10)
    run5b.font.bold = True
    run5b.font.color.rgb = RGBColor(*hex_to_rgb(color))

    # Spacer
    document.add_paragraph()


# =============================================================================
# FIGURE CAPTIONS
# =============================================================================

def add_captioned_image(
    document: Document,
    image_path: str,
    caption_text: str,
    width: Inches = Inches(6)
) -> int:
    """
    Add an image with a numbered caption.

    Args:
        document: python-docx Document object
        image_path: Path to the image file
        caption_text: Caption description text
        width: Image width

    Returns:
        Figure number

    Example output: "Figure 3: Risk Heatmap Comparison Before vs After Mitigation"
    """
    # Add image centered
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_path, width=width)

    # Add caption
    caption_para = document.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_num = figure_counter.next()
    caption_run = caption_para.add_run(f'Figure {fig_num}: {caption_text}')
    caption_run.font.size = Pt(FONT_SIZES['caption'])
    caption_run.font.italic = True
    caption_run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    # Add spacer
    document.add_paragraph()

    return fig_num


def get_chart_caption(chart_name: str) -> str:
    """
    Get predefined caption text for a chart.

    Args:
        chart_name: Key for the chart in CHART_CAPTIONS

    Returns:
        Caption text or the chart name if not found
    """
    return CHART_CAPTIONS.get(chart_name, chart_name)


# =============================================================================
# APPENDICES
# =============================================================================

def add_glossary_appendix(document: Document):
    """
    Add Glossary of Terms appendix.

    Args:
        document: python-docx Document object
    """
    document.add_heading('Appendix C: Glossary of Terms', level=1)

    document.add_paragraph(
        'This glossary provides definitions for key terms used throughout this report.'
    )
    document.add_paragraph()

    table = document.add_table(rows=len(GLOSSARY_TERMS) + 1, cols=2)

    # Header
    table.cell(0, 0).text = 'Term'
    table.cell(0, 1).text = 'Definition'

    # Terms
    for i, (term, definition) in enumerate(GLOSSARY_TERMS, start=1):
        table.cell(i, 0).text = term
        table.cell(i, 1).text = definition

    apply_professional_table_style(table)

    # Set column widths
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.5)


def add_assumptions_appendix(document: Document, simulation_params: Optional[Dict] = None):
    """
    Add Assumptions and Limitations appendix.

    Args:
        document: python-docx Document object
        simulation_params: dict with 'iterations', 'confidence_level', 'data_date', etc.
    """
    document.add_heading('Appendix D: Assumptions and Limitations', level=1)

    if simulation_params is None:
        simulation_params = {}

    # Simulation Parameters
    document.add_heading('Simulation Parameters', level=2)

    iterations = simulation_params.get('iterations', 15000)
    confidence_level = simulation_params.get('confidence_level', 80)
    data_date = simulation_params.get('data_date', 'As per Risk Register')

    params_para = document.add_paragraph()
    params_para.add_run(f'• Monte Carlo Iterations: {iterations:,}\n')
    params_para.add_run(f'• Selected Confidence Level: P{confidence_level}\n')
    params_para.add_run(f'• Risk Data Date: {data_date}\n')
    params_para.add_run('• Correlation Model: Independent risks (no correlation modelled)')

    # Key Assumptions
    document.add_heading('Key Assumptions', level=2)
    for assumption in DEFAULT_ASSUMPTIONS:
        document.add_paragraph(assumption, style='List Bullet')

    # Limitations
    document.add_heading('Limitations', level=2)
    for limitation in DEFAULT_LIMITATIONS:
        document.add_paragraph(limitation, style='List Bullet')

    # Disclaimer
    document.add_heading('Disclaimer', level=2)
    disclaimer_para = document.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(DEFAULT_DISCLAIMER)
    disclaimer_run.font.italic = True


def add_approval_appendix(document: Document, approval_roles: Optional[List[Tuple]] = None, config: Optional[Dict] = None):
    """
    Add Document Approval sign-off page with document control section.

    Args:
        document: python-docx Document object
        approval_roles: List of (role, title, name) tuples
        config: Optional configuration dict with revision, author, date info
    """
    document.add_heading('Appendix E: Document Approval', level=1)

    # Document Control Section
    document.add_heading('Document Control', level=2)

    if config is None:
        config = {}

    report_date = config.get('report_date', datetime.now().strftime('%d %B %Y'))
    revision = config.get('revision', 'A')
    author = config.get('author', '')

    dc_table = document.add_table(rows=2, cols=4)
    dc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    dc_headers = ['Rev', 'Date', 'Author', 'Description']
    for j, header in enumerate(dc_headers):
        dc_table.cell(0, j).text = header

    # First revision entry
    dc_data = [revision, report_date, author, 'Initial Issue']
    for j, data in enumerate(dc_data):
        dc_table.cell(1, j).text = data

    apply_professional_table_style(dc_table)

    # Set column widths
    dc_table.columns[0].width = Inches(0.8)
    dc_table.columns[1].width = Inches(1.5)
    dc_table.columns[2].width = Inches(2.0)
    dc_table.columns[3].width = Inches(2.2)

    document.add_paragraph()

    # Approval Section
    document.add_heading('Document Approval', level=2)

    document.add_paragraph(
        'This document has been reviewed and approved by the following personnel:'
    )
    document.add_paragraph()

    if approval_roles is None:
        approval_roles = DEFAULT_APPROVAL_ROLES

    table = document.add_table(rows=len(approval_roles) + 1, cols=4)

    # Header
    headers = ['Role', 'Name', 'Signature', 'Date']
    for j, header in enumerate(headers):
        table.cell(0, j).text = header

    # Data rows
    for i, (role, title, name) in enumerate(approval_roles, start=1):
        table.cell(i, 0).text = f'{role}\n({title})'
        table.cell(i, 1).text = name if name else ''
        table.cell(i, 2).text = ''  # Signature space
        table.cell(i, 3).text = ''  # Date space

    apply_professional_table_style(table)

    # Set row heights for signature space
    for row in table.rows[1:]:
        set_row_height(row, 1.5)

    # Set column widths
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(1.8)
    table.columns[2].width = Inches(1.5)
    table.columns[3].width = Inches(1.2)


def add_phase_summary_table(document: Document, risk_data: List[Dict]):
    """
    Add a summary table showing risks by crystallization phase.
    Only include if phase data is available.

    Args:
        document: python-docx Document object
        risk_data: List of risk dictionaries with phase data
    """
    # Check if phase data exists
    phases_with_data = [r.get('phase') for r in risk_data if r.get('phase') in VALID_PHASES]

    if not phases_with_data:
        return  # Skip if no phase data

    document.add_heading('Risk Distribution by Project Phase', level=2)

    # Count risks per phase
    phase_counts = {phase: 0 for phase in VALID_PHASES}
    phase_exposure = {phase: 0.0 for phase in VALID_PHASES}

    for risk in risk_data:
        phase = risk.get('phase')
        if phase in VALID_PHASES:
            phase_counts[phase] += 1
            ev = risk.get('initial_impact', 0) * risk.get('initial_likelihood', 0)
            phase_exposure[phase] += ev

    # Create table
    table = document.add_table(rows=len(VALID_PHASES) + 2, cols=3)

    # Header
    headers = ['Phase', 'Risk Count', 'Expected Exposure (M CHF)']
    for j, header in enumerate(headers):
        table.cell(0, j).text = header

    # Data rows
    total_count = 0
    total_exposure = 0.0
    for i, phase in enumerate(VALID_PHASES, start=1):
        table.cell(i, 0).text = phase
        table.cell(i, 1).text = str(phase_counts[phase])
        exposure_m = phase_exposure[phase] / 1_000_000
        table.cell(i, 2).text = f"{exposure_m:,.2f}"
        total_count += phase_counts[phase]
        total_exposure += exposure_m

    # Total row
    total_row = len(VALID_PHASES) + 1
    table.cell(total_row, 0).text = 'Total'
    table.cell(total_row, 1).text = str(total_count)
    table.cell(total_row, 2).text = f"{total_exposure:,.2f}"

    apply_professional_table_style(table, has_total_row=True)


# =============================================================================
# DOCUMENT PROPERTIES
# =============================================================================

def set_document_properties(
    document: Document,
    project_name: str,
    author: str = 'Kanadevia INOVA'
):
    """
    Set document metadata properties.

    Args:
        document: python-docx Document object
        project_name: Project name for title
        author: Document author
    """
    core_props = document.core_properties
    core_props.title = f'Risk Assessment Report - {project_name}'
    core_props.author = author
    core_props.subject = 'Monte Carlo Risk Analysis'
    core_props.keywords = f'Risk Assessment, Monte Carlo, P80, {project_name}, Kanadevia INOVA'
    core_props.category = 'Project Risk Management'
    core_props.comments = 'Generated by KVI Risk Assessment Application'


# =============================================================================
# NUMBERED HEADINGS
# =============================================================================

def add_numbered_heading(document: Document, text: str, level: int = 1, is_appendix: bool = False):
    """
    Add a numbered heading to the document.

    Args:
        document: python-docx Document object
        text: Heading text
        level: Heading level (1, 2, or 3)
        is_appendix: If True, don't add number prefix (appendices have their own numbering)

    Returns:
        The created paragraph object
    """
    if is_appendix:
        # Appendices use letter-based numbering already in the text
        heading = document.add_heading(text, level=level)
    else:
        # Add section number prefix
        if level == 1:
            num = section_counter.next_level1()
            heading = document.add_heading(f"{num}. {text}", level=level)
        elif level == 2:
            num = section_counter.next_level2()
            heading = document.add_heading(f"{num} {text}", level=level)
        else:
            num = section_counter.next_level3()
            heading = document.add_heading(f"{num} {text}", level=level)

    # Apply KVI styling
    for run in heading.runs:
        run.font.name = FONTS['heading']
        if level == 1:
            run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['table_header']))
        elif level == 2:
            run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['primary']))
        else:
            run.font.color.rgb = RGBColor(*hex_to_rgb(KVI_COLORS['neutral']))

    return heading


# =============================================================================
# RESET FUNCTION
# =============================================================================

def reset_for_new_document():
    """Reset global state for generating a new document."""
    figure_counter.reset()
    section_counter.reset()
