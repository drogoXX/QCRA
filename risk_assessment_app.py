"""
Risk Assessment Tool - Streamlit Application
Probabilistic Risk Analysis with Monte Carlo Simulation
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import io
from datetime import datetime

# Page configuration
st.set_page_config(
    page_title="Risk Assessment Tool",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    </style>
    """, unsafe_allow_html=True)

# Title
st.markdown('<p class="main-header">🎯 Risk Assessment Tool - Monte Carlo Simulation</p>', unsafe_allow_html=True)

# Helper functions
def parse_currency(value):
    """Parse currency string to float"""
    if pd.isna(value) or value == '' or value == '0 CHF':
        return 0.0
    try:
        # Remove CHF, commas, and spaces
        clean_value = str(value).replace('CHF', '').replace(',', '').replace(' ', '').strip()
        # Handle quotes
        clean_value = clean_value.replace('"', '')
        return float(clean_value)
    except:
        return 0.0

def parse_percentage(value):
    """Parse percentage string to float"""
    if pd.isna(value) or value == '':
        return 0.0
    try:
        clean_value = str(value).replace('%', '').strip()
        return float(clean_value) / 100
    except:
        return 0.0

@st.cache_data
def load_risk_data(file_path):
    """Load and process risk register data"""
    df = pd.read_csv(file_path, encoding='utf-8-sig')

    # Clean column names
    df.columns = df.columns.str.strip()

    # Normalize phase column names (support underscore format)
    phase_col_mapping = {
        'crystallization_phase': 'Crystallization Phase',
        'phase_weight_distribution': 'Phase Weight Distribution'
    }
    for old_name, new_name in phase_col_mapping.items():
        if old_name in df.columns and new_name not in df.columns:
            df.rename(columns={old_name: new_name}, inplace=True)

    # Parse currency columns
    currency_cols = ['Initial risk', 'Likely Initial Risk', 'Residual risk',
                     'Likely Residual Risk', 'Cost of Measures']
    for col in currency_cols:
        if col in df.columns:
            df[col + '_Value'] = df[col].apply(parse_currency)

    # Parse likelihood columns
    df['Initial_Likelihood'] = df['Likelihood'].apply(parse_percentage)

    # Handle the second Likelihood column (for residual)
    likelihood_cols = [col for col in df.columns if 'Likelihood' in col]
    if len(likelihood_cols) > 1:
        df['Residual_Likelihood'] = df[likelihood_cols[1]].apply(parse_percentage)

    # Parse schedule impact
    df['Schedule_Impact'] = df['Schedule Impact?'].str.strip().str.lower().isin(['yes', 'yes '])

    # Parse phase columns if present (for Time-Phased Contingency Profile)
    if 'Phase Weight Distribution' in df.columns:
        df['Phase_Weights'] = df['Phase Weight Distribution'].apply(parse_phase_weights)

    return df

def run_monte_carlo(df, n_simulations=10000, risk_type='initial', random_numbers=None):
    """
    Run Monte Carlo simulation for risk portfolio

    Parameters:
    - df: DataFrame with risk data
    - n_simulations: Number of Monte Carlo iterations
    - risk_type: 'initial' or 'residual'
    - random_numbers: Optional pre-generated random numbers for Common Random Numbers technique
                     Shape should be (n_simulations, len(df))
    """
    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'

    # Prepare data
    impacts = df[impact_col].values
    likelihoods = df[likelihood_col].values

    # Generate or use provided random numbers
    if random_numbers is None:
        random_numbers = np.random.random((n_simulations, len(df)))

    # Monte Carlo simulation
    results = np.zeros(n_simulations)
    risk_occurrences = np.zeros((n_simulations, len(df)))

    for i in range(n_simulations):
        # For each risk, determine if it occurs (Bernoulli trial)
        # Using the same random numbers ensures variance reduction
        occurred = random_numbers[i] < likelihoods
        risk_occurrences[i] = occurred

        # Sum the impacts of risks that occurred
        results[i] = np.sum(impacts * occurred)

    return results, risk_occurrences

def calculate_statistics(results):
    """Calculate statistical measures from Monte Carlo results"""
    return {
        'mean': np.mean(results),
        'median': np.median(results),
        'std': np.std(results),
        'p50': np.percentile(results, 50),
        'p80': np.percentile(results, 80),
        'p90': np.percentile(results, 90),
        'p95': np.percentile(results, 95),
        'min': np.min(results),
        'max': np.max(results)
    }

def get_confidence_value(stats, confidence_level):
    """Get the risk value for the selected confidence level

    Parameters:
    - stats: Dictionary with statistical measures from calculate_statistics()
    - confidence_level: String like 'P50', 'P80', 'P90', or 'P95'

    Returns:
    - The corresponding percentile value
    """
    percentile_map = {
        'P50': 'p50',
        'P80': 'p80',
        'P90': 'p90',
        'P95': 'p95'
    }
    return stats[percentile_map.get(confidence_level, 'p80')]

def calculate_confidence_comparison(results, mitigation_cost, selected_confidence='P80'):
    """
    Calculate confidence level comparison metrics for contingency planning.

    Args:
        results: Monte Carlo simulation results array
        mitigation_cost: Total cost of mitigation measures
        selected_confidence: Currently selected confidence level

    Returns:
        Dictionary with comparison tables and metrics
    """
    # Extract percentiles
    percentiles = {
        'P50': np.percentile(results, 50),
        'P80': np.percentile(results, 80),
        'P90': np.percentile(results, 90),
        'P95': np.percentile(results, 95),
        'Max': np.max(results)
    }

    # Calculate total contingency (residual exposure + mitigation cost)
    total_contingency = {level: value + mitigation_cost for level, value in percentiles.items()}

    # Calculate delta from P50
    p50_value = total_contingency['P50']
    delta_from_p50 = {level: value - p50_value for level, value in total_contingency.items()}

    # Calculate premium vs P50 (percentage)
    premium_vs_p50 = {}
    for level, value in total_contingency.items():
        if level == 'P50':
            premium_vs_p50[level] = 0
        else:
            premium_vs_p50[level] = ((value - p50_value) / p50_value * 100) if p50_value > 0 else 0

    # Confidence level comparison table data
    comparison_table = {
        'levels': ['P50', 'P80', 'P90', 'P95', 'Max'],
        'residual_exposure': [percentiles[l] for l in ['P50', 'P80', 'P90', 'P95', 'Max']],
        'mitigation_cost': [mitigation_cost] * 5,
        'total_contingency': [total_contingency[l] for l in ['P50', 'P80', 'P90', 'P95', 'Max']],
        'delta_from_p50': [delta_from_p50[l] for l in ['P50', 'P80', 'P90', 'P95', 'Max']],
        'premium_vs_p50': [premium_vs_p50[l] for l in ['P50', 'P80', 'P90', 'P95', 'Max']]
    }

    # Incremental cost analysis
    incremental_steps = [
        ('P50 → P80', 'P50', 'P80', 30),
        ('P80 → P90', 'P80', 'P90', 10),
        ('P90 → P95', 'P90', 'P95', 5)
    ]

    incremental_analysis = []
    for step_name, from_level, to_level, pct_increase in incremental_steps:
        additional_cost = total_contingency[to_level] - total_contingency[from_level]
        pct_change = ((total_contingency[to_level] - total_contingency[from_level]) /
                     total_contingency[from_level] * 100) if total_contingency[from_level] > 0 else 0
        cost_per_pct = additional_cost / pct_increase if pct_increase > 0 else 0

        incremental_analysis.append({
            'step': step_name,
            'confidence_increase': f'{pct_increase}%',
            'additional_cost': additional_cost,
            'pct_increase': pct_change,
            'cost_per_1pct': cost_per_pct
        })

    return {
        'percentiles': percentiles,
        'total_contingency': total_contingency,
        'comparison_table': comparison_table,
        'incremental_analysis': incremental_analysis,
        'selected_confidence': selected_confidence,
        'mitigation_cost': mitigation_cost
    }

# =============================================================================
# TIME-PHASED CONTINGENCY PROFILE FUNCTIONS
# =============================================================================

# Default project phase definitions
DEFAULT_PHASES = {
    'ENG': {'name': 'Engineering', 'order': 1, 'color': '#3498db'},
    'PROC': {'name': 'Procurement', 'order': 2, 'color': '#9b59b6'},
    'FAB': {'name': 'Fabrication', 'order': 3, 'color': '#e67e22'},
    'CONS': {'name': 'Construction', 'order': 4, 'color': '#27ae60'},
    'COMM': {'name': 'Commissioning', 'order': 5, 'color': '#e74c3c'},
    'WARR': {'name': 'Warranty', 'order': 6, 'color': '#95a5a6'}
}

def parse_phase_weights(weight_string):
    """
    Parse phase weight distribution string into dictionary.

    Args:
        weight_string: String like "ENG:0.5|PROC:0.3|FAB:0.2"

    Returns:
        Dictionary of phase -> weight
    """
    if pd.isna(weight_string) or not weight_string:
        return {}

    weights = {}
    try:
        pairs = str(weight_string).split('|')
        for pair in pairs:
            if ':' in pair:
                phase, weight = pair.split(':')
                weights[phase.strip()] = float(weight.strip())
    except:
        pass

    return weights

def load_enhanced_risk_data(file_path):
    """
    Load and process enhanced risk register data with phase information.

    Args:
        file_path: Path to enhanced risk register CSV

    Returns:
        DataFrame with processed risk and phase data
    """
    df = pd.read_csv(file_path, encoding='utf-8-sig')

    # Clean column names
    df.columns = df.columns.str.strip()

    # Parse currency columns
    currency_cols = ['Initial risk', 'Likely Initial Risk', 'Residual risk',
                     'Likely Residual Risk', 'Cost of Measures']
    for col in currency_cols:
        if col in df.columns:
            df[col + '_Value'] = df[col].apply(parse_currency)

    # Parse likelihood columns
    df['Initial_Likelihood'] = df['Likelihood'].apply(parse_percentage)

    # Handle the second Likelihood column (for residual)
    likelihood_cols = [col for col in df.columns if 'Likelihood' in col]
    if len(likelihood_cols) > 1:
        df['Residual_Likelihood'] = df[likelihood_cols[1]].apply(parse_percentage)

    # Parse schedule impact
    if 'Schedule Impact?' in df.columns:
        df['Schedule_Impact'] = df['Schedule Impact?'].str.strip().str.lower().isin(['yes', 'yes '])

    # Parse phase weights
    if 'Phase Weight Distribution' in df.columns:
        df['Phase_Weights'] = df['Phase Weight Distribution'].apply(parse_phase_weights)
    else:
        df['Phase_Weights'] = [{} for _ in range(len(df))]

    # Default crystallization phase if not specified
    if 'Crystallization Phase' not in df.columns:
        df['Crystallization Phase'] = 'ENG'

    return df

def calculate_phase_allocation(df, results, risk_occurrences, risk_type='residual',
                               confidence_level='P80', phases=None):
    """
    Calculate contingency allocation across project phases.

    Args:
        df: DataFrame with enhanced risk data
        results: Monte Carlo simulation results
        risk_occurrences: Array of risk occurrence flags from simulation
        risk_type: 'initial' or 'residual'
        confidence_level: 'P50', 'P80', 'P90', or 'P95'
        phases: Dictionary of phase definitions (uses DEFAULT_PHASES if None)

    Returns:
        Dictionary with phase allocation data
    """
    if phases is None:
        phases = DEFAULT_PHASES

    # Get impact column
    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'

    n_simulations = len(results)
    phase_results = {phase: np.zeros(n_simulations) for phase in phases.keys()}

    # Calculate expected value per risk
    df_temp = df.copy()
    df_temp['EV'] = df_temp[impact_col] * df_temp[likelihood_col]

    # For each simulation, allocate risk impacts to phases
    for sim_idx in range(n_simulations):
        occurred_mask = risk_occurrences[sim_idx].astype(bool)

        for risk_idx, (_, row) in enumerate(df.iterrows()):
            if occurred_mask[risk_idx]:
                impact = row[impact_col]
                weights = row.get('Phase_Weights', {})
                primary_phase = row.get('Crystallization Phase', 'ENG')

                if weights:
                    # Distribute impact according to weights
                    for phase, weight in weights.items():
                        if phase in phase_results:
                            phase_results[phase][sim_idx] += impact * weight
                else:
                    # Assign all impact to primary phase
                    if primary_phase in phase_results:
                        phase_results[primary_phase][sim_idx] += impact

    # Calculate statistics for each phase
    percentile_map = {'P50': 50, 'P80': 80, 'P90': 90, 'P95': 95}
    percentile_value = percentile_map.get(confidence_level, 80)

    # First, calculate total simulation results by summing all phases per simulation
    # This is needed for correct total percentile calculation
    # (Percentiles are NOT additive: sum of P80s ≠ P80 of sum)
    total_sims = np.zeros(n_simulations)
    for phase_sims in phase_results.values():
        total_sims += phase_sims

    # Calculate correct total percentiles from combined simulation results
    total_at_confidence = np.percentile(total_sims, percentile_value)
    total_p50 = np.percentile(total_sims, 50)
    total_p80 = np.percentile(total_sims, 80)
    total_p90 = np.percentile(total_sims, 90)
    total_p95 = np.percentile(total_sims, 95)

    phase_stats = {}
    total_ev = 0
    sum_of_phase_allocations = 0  # Sum of individual phase percentiles (for distribution %)

    for phase, phase_sims in phase_results.items():
        ev = np.mean(phase_sims)
        at_confidence = np.percentile(phase_sims, percentile_value)
        p50 = np.percentile(phase_sims, 50)
        p80 = np.percentile(phase_sims, 80)
        p90 = np.percentile(phase_sims, 90)
        p95 = np.percentile(phase_sims, 95)

        total_ev += ev
        sum_of_phase_allocations += at_confidence

        phase_stats[phase] = {
            'name': phases[phase]['name'],
            'order': phases[phase]['order'],
            'color': phases[phase]['color'],
            'expected_value': ev,
            'at_confidence': at_confidence,
            'p50': p50,
            'p80': p80,
            'p90': p90,
            'p95': p95,
            'std': np.std(phase_sims),
            'simulation_results': phase_sims
        }

    # Calculate percentages
    for phase in phase_stats:
        if total_ev > 0:
            phase_stats[phase]['ev_percentage'] = (phase_stats[phase]['expected_value'] / total_ev) * 100
        else:
            phase_stats[phase]['ev_percentage'] = 0

        # Use sum_of_phase_allocations for distribution % so percentages sum to 100%
        if sum_of_phase_allocations > 0:
            phase_stats[phase]['confidence_percentage'] = (phase_stats[phase]['at_confidence'] / sum_of_phase_allocations) * 100
        else:
            phase_stats[phase]['confidence_percentage'] = 0

    # Calculate cumulative allocation (for S-curve)
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])
    cumulative_ev = 0
    cumulative_confidence = 0

    for phase, stats in sorted_phases:
        cumulative_ev += stats['expected_value']
        cumulative_confidence += stats['at_confidence']
        phase_stats[phase]['cumulative_ev'] = cumulative_ev
        phase_stats[phase]['cumulative_at_confidence'] = cumulative_confidence
        phase_stats[phase]['cumulative_ev_pct'] = (cumulative_ev / total_ev * 100) if total_ev > 0 else 0
        # Use sum_of_phase_allocations so cumulative reaches 100%
        phase_stats[phase]['cumulative_confidence_pct'] = (cumulative_confidence / sum_of_phase_allocations * 100) if sum_of_phase_allocations > 0 else 0

    return {
        'phase_stats': phase_stats,
        'total_ev': total_ev,
        'total_at_confidence': total_at_confidence,  # Correct P80 of total (matches Confidence Level Comparison)
        'sum_of_phase_allocations': sum_of_phase_allocations,  # Sum of individual phase P80s (for reference)
        'confidence_level': confidence_level,
        'phases': phases
    }

def get_early_warning_status(consumed_pct, phase_complete_pct, threshold_amber=0.1, threshold_red=0.2):
    """
    Determine early warning indicator status based on consumption vs completion.

    Args:
        consumed_pct: Percentage of contingency consumed (0-100)
        phase_complete_pct: Percentage of phase completed (0-100)
        threshold_amber: Amber threshold (deviation)
        threshold_red: Red threshold (deviation)

    Returns:
        Dictionary with status ('Green', 'Amber', 'Red') and deviation
    """
    # Calculate expected consumption based on phase completion
    expected_consumption = phase_complete_pct

    # Calculate deviation (positive means over-consuming)
    deviation = consumed_pct - expected_consumption

    if deviation > threshold_red * 100:
        status = 'Red'
        message = f'Over-consuming by {deviation:.1f}% - Immediate attention required'
    elif deviation > threshold_amber * 100:
        status = 'Amber'
        message = f'Over-consuming by {deviation:.1f}% - Monitor closely'
    elif deviation < -threshold_amber * 100:
        status = 'Green'
        message = f'Under-consuming by {abs(deviation):.1f}% - On track'
    else:
        status = 'Green'
        message = f'Consumption aligned with progress'

    return {
        'status': status,
        'deviation': deviation,
        'message': message,
        'consumed_pct': consumed_pct,
        'expected_pct': expected_consumption
    }

def create_risk_heatmap(df, risk_type='initial'):
    """Create risk heatmap showing concentration of risks in each cell"""
    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
        title = 'Initial Risk Heatmap'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'
        title = 'Residual Risk Heatmap'
    
    # Create categories
    impact_bins = [0, 100000, 1000000, 10000000, 100000000, np.inf]
    impact_labels = ['Very Low', 'Low', 'Medium', 'High', 'Very High']
    
    likelihood_bins = [0, 0.1, 0.25, 0.5, 0.75, 1.0]
    likelihood_labels = ['Very Low', 'Low', 'Medium', 'High', 'Very High']
    
    # Categorize risks (use absolute values for impact to handle opportunities)
    df_temp = df.copy()
    df_temp['Impact_Cat'] = pd.cut(np.abs(df_temp[impact_col]), bins=impact_bins, labels=impact_labels)
    df_temp['Likelihood_Cat'] = pd.cut(df_temp[likelihood_col], bins=likelihood_bins, labels=likelihood_labels)
    
    # Create heatmap data
    heatmap_data = []
    for imp_cat in impact_labels:
        row_data = []
        for lik_cat in likelihood_labels:
            # Count risks in this cell
            mask = (df_temp['Impact_Cat'] == imp_cat) & (df_temp['Likelihood_Cat'] == lik_cat)
            count = mask.sum()
            
            # Calculate total value in this cell
            total_value = (df_temp.loc[mask, impact_col] * df_temp.loc[mask, likelihood_col]).sum()
            
            # Get risk IDs in this cell
            risk_ids = df_temp.loc[mask, 'Risk ID'].tolist()
            
            row_data.append({
                'count': count,
                'value': total_value,
                'risk_ids': risk_ids
            })
        heatmap_data.append(row_data)
    
    # Prepare data for plotly
    z_counts = [[cell['count'] for cell in row] for row in heatmap_data]
    z_values = [[cell['value']/1e6 for cell in row] for row in heatmap_data]  # Convert to millions
    
    # Create hover text
    hover_text = []
    for i, imp_cat in enumerate(impact_labels):
        row_text = []
        for j, lik_cat in enumerate(likelihood_labels):
            cell = heatmap_data[i][j]
            text = f"<b>{imp_cat} Impact / {lik_cat} Likelihood</b><br>"
            text += f"Number of Risks: {cell['count']}<br>"
            text += f"Total Expected Value: {cell['value']/1e6:.2f}M CHF<br>"
            if cell['risk_ids']:
                text += f"Risk IDs: {', '.join(map(str, cell['risk_ids']))}"
            row_text.append(text)
        hover_text.append(row_text)
    
    # Create heatmap
    fig = go.Figure(data=go.Heatmap(
        z=z_counts,
        x=likelihood_labels,
        y=impact_labels,
        customdata=z_values,
        hovertemplate='%{hovertext}<extra></extra>',
        hovertext=hover_text,
        colorscale='Reds',
        colorbar=dict(title="Number<br>of Risks"),
        text=[[f"{count}<br>{value:.1f}M" for count, value in zip(count_row, value_row)] 
              for count_row, value_row in zip(z_counts, z_values)],
        texttemplate="%{text}",
        textfont={"size": 10}
    ))
    
    # Add grid lines and styling
    fig.update_xaxes(side="bottom", title="Likelihood →")
    fig.update_yaxes(title="← Impact")
    
    fig.update_layout(
        title=f'{title}<br><sub>Cell shows: Risk Count / Total Expected Value (M CHF)</sub>',
        height=600,
        xaxis=dict(showgrid=True, gridcolor='white', gridwidth=2),
        yaxis=dict(showgrid=True, gridcolor='white', gridwidth=2)
    )
    
    return fig

def create_risk_bubble_chart(df, risk_type='initial'):
    """Create bubble chart with risk size representing expected value"""
    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
        title = 'Initial Risk Bubble Chart'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'
        title = 'Residual Risk Bubble Chart'
    
    df_plot = df.copy()
    
    # Calculate expected value and use for bubble size
    df_plot['Expected_Value'] = np.abs(df_plot[impact_col] * df_plot[likelihood_col])
    df_plot['Risk_Type'] = df_plot[impact_col].apply(lambda x: 'Opportunity' if x < 0 else 'Risk')
    
    # Create risk score for coloring
    df_plot['Risk_Score'] = df_plot[impact_col] * df_plot[likelihood_col]
    
    # Scale bubble sizes (cube root for better visualization)
    df_plot['Bubble_Size'] = np.power(df_plot['Expected_Value'], 1/3) / 100
    
    fig = px.scatter(df_plot,
                     x=likelihood_col,
                     y=impact_col,
                     size='Bubble_Size',
                     color='Risk_Score',
                     hover_data=['Risk ID', 'Risk Description', 'Expected_Value', 'Risk_Type'],
                     labels={
                         likelihood_col: 'Likelihood',
                         impact_col: 'Impact (CHF)',
                         'Expected_Value': 'Expected Value (CHF)'
                     },
                     color_continuous_scale='RdYlGn_r',
                     title=f'{title}<br><sub>Bubble size = Expected Value | Color = Risk Score</sub>')
    
    # Add quadrant lines
    median_likelihood = df_plot[likelihood_col].median()
    median_impact = df_plot[impact_col].median()
    
    fig.add_vline(x=median_likelihood, line_dash="dash", line_color="gray", opacity=0.5,
                  annotation_text="Median Likelihood", annotation_position="top")
    fig.add_hline(y=median_impact, line_dash="dash", line_color="gray", opacity=0.5,
                  annotation_text="Median Impact", annotation_position="right")
    
    # Add zero line for opportunities
    fig.add_hline(y=0, line_dash="solid", line_color="black", opacity=0.3)
    
    # Add quadrant labels
    max_likelihood = df_plot[likelihood_col].max()
    max_impact = df_plot[impact_col].max()
    min_impact = df_plot[impact_col].min()
    
    fig.add_annotation(x=median_likelihood/2, y=max_impact*0.9,
                      text="Low Likelihood<br>High Impact",
                      showarrow=False, font=dict(size=10, color="gray"))
    
    fig.add_annotation(x=median_likelihood + (max_likelihood-median_likelihood)/2, y=max_impact*0.9,
                      text="High Likelihood<br>High Impact",
                      showarrow=False, font=dict(size=10, color="red"))
    
    fig.add_annotation(x=median_likelihood/2, y=min_impact*0.5,
                      text="Low Likelihood<br>Low Impact",
                      showarrow=False, font=dict(size=10, color="green"))
    
    fig.add_annotation(x=median_likelihood + (max_likelihood-median_likelihood)/2, y=min_impact*0.5,
                      text="High Likelihood<br>Low Impact",
                      showarrow=False, font=dict(size=10, color="orange"))
    
    fig.update_layout(height=700, showlegend=True)

    return fig

def create_3d_risk_surface(df, risk_type='initial'):
    """
    Create 3D surface plot visualization of risk landscape.

    Shows risks positioned in 3D space with:
    - X-axis: Likelihood
    - Y-axis: Impact
    - Z-axis: Expected Value
    - Color: Risk severity
    """
    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
        ev_col = 'Initial_EV'
        title = '3D Risk Landscape - Initial Assessment'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'
        ev_col = 'Residual_EV'
        title = '3D Risk Landscape - Residual Assessment'

    df_plot = df.copy()

    # Separate threats and opportunities for different coloring
    df_threats = df_plot[df_plot[ev_col] > 0]
    df_opportunities = df_plot[df_plot[ev_col] < 0]

    fig = go.Figure()

    # Add threats as red markers
    if len(df_threats) > 0:
        fig.add_trace(go.Scatter3d(
            x=df_threats[likelihood_col] * 100,
            y=np.abs(df_threats[impact_col]) / 1e6,
            z=df_threats[ev_col] / 1e6,
            mode='markers+text',
            marker=dict(
                size=np.clip(np.abs(df_threats[ev_col]) / 1e6 * 3 + 5, 8, 30),
                color=df_threats[ev_col] / 1e6,
                colorscale='Reds',
                opacity=0.8,
                line=dict(width=1, color='darkred')
            ),
            text=df_threats['Risk ID'].astype(str),
            textposition='top center',
            textfont=dict(size=9, color='darkred'),
            name='Threats',
            hovertemplate=(
                '<b>Risk %{text}</b><br>' +
                'Likelihood: %{x:.1f}%<br>' +
                'Impact: %{y:.2f}M CHF<br>' +
                'Expected Value: %{z:.2f}M CHF<br>' +
                '<extra></extra>'
            )
        ))

    # Add opportunities as green markers
    if len(df_opportunities) > 0:
        fig.add_trace(go.Scatter3d(
            x=df_opportunities[likelihood_col] * 100,
            y=np.abs(df_opportunities[impact_col]) / 1e6,
            z=df_opportunities[ev_col] / 1e6,
            mode='markers+text',
            marker=dict(
                size=np.clip(np.abs(df_opportunities[ev_col]) / 1e6 * 3 + 5, 8, 30),
                color='green',
                opacity=0.8,
                line=dict(width=1, color='darkgreen'),
                symbol='diamond'
            ),
            text=df_opportunities['Risk ID'].astype(str),
            textposition='top center',
            textfont=dict(size=9, color='darkgreen'),
            name='Opportunities',
            hovertemplate=(
                '<b>Risk %{text}</b><br>' +
                'Likelihood: %{x:.1f}%<br>' +
                'Impact: %{y:.2f}M CHF<br>' +
                'Expected Value: %{z:.2f}M CHF<br>' +
                '<extra></extra>'
            )
        ))

    # Add a zero plane to show the threshold between threats and opportunities
    max_likelihood = 100
    max_impact = np.abs(df_plot[impact_col]).max() / 1e6 * 1.1

    # Create mesh for zero plane
    xx = np.linspace(0, max_likelihood, 10)
    yy = np.linspace(0, max_impact, 10)
    xx, yy = np.meshgrid(xx, yy)
    zz = np.zeros_like(xx)

    fig.add_trace(go.Surface(
        x=xx, y=yy, z=zz,
        opacity=0.3,
        colorscale=[[0, 'lightgray'], [1, 'lightgray']],
        showscale=False,
        name='Zero Plane',
        hoverinfo='skip'
    ))

    # Add risk zone indicators (vertical planes)
    # High risk zone boundary
    fig.add_trace(go.Scatter3d(
        x=[50, 50, 50, 50, 50],
        y=[0, max_impact, max_impact, 0, 0],
        z=[df_plot[ev_col].min()/1e6, df_plot[ev_col].min()/1e6,
           df_plot[ev_col].max()/1e6, df_plot[ev_col].max()/1e6, df_plot[ev_col].min()/1e6],
        mode='lines',
        line=dict(color='orange', width=3, dash='dash'),
        name='50% Likelihood Threshold',
        showlegend=True
    ))

    fig.update_layout(
        title=dict(
            text=f'{title}<br><sub>Bubble size represents Expected Value magnitude</sub>',
            font=dict(size=16, color='#1F4E78')
        ),
        scene=dict(
            xaxis=dict(
                title='Likelihood (%)',
                range=[0, 100],
                gridcolor='lightgray',
                showbackground=True,
                backgroundcolor='rgba(230,230,250,0.3)'
            ),
            yaxis=dict(
                title='Impact (M CHF)',
                gridcolor='lightgray',
                showbackground=True,
                backgroundcolor='rgba(230,250,230,0.3)'
            ),
            zaxis=dict(
                title='Expected Value (M CHF)',
                gridcolor='lightgray',
                showbackground=True,
                backgroundcolor='rgba(250,230,230,0.3)',
                zeroline=True,
                zerolinecolor='black',
                zerolinewidth=2
            ),
            camera=dict(
                eye=dict(x=1.5, y=1.5, z=1.2)
            ),
            aspectmode='manual',
            aspectratio=dict(x=1, y=1, z=0.8)
        ),
        height=700,
        legend=dict(
            yanchor="top",
            y=0.99,
            xanchor="left",
            x=0.01,
            bgcolor='rgba(255,255,255,0.8)'
        ),
        margin=dict(l=0, r=0, t=80, b=0)
    )

    return fig

def create_3d_risk_comparison(df):
    """
    Create side-by-side 3D comparison of initial vs residual risk landscape.
    Shows risk movement through mitigation in 3D space.
    """
    from plotly.subplots import make_subplots

    fig = make_subplots(
        rows=1, cols=2,
        specs=[[{'type': 'scatter3d'}, {'type': 'scatter3d'}]],
        subplot_titles=('Initial Risk Landscape', 'Residual Risk Landscape'),
        horizontal_spacing=0.05
    )

    # Prepare data
    df_plot = df.copy()

    # Separate threats and opportunities
    df_threats = df_plot[df_plot['Initial_EV'] > 0]
    df_opportunities = df_plot[df_plot['Initial_EV'] < 0]

    # Initial risks - Threats
    if len(df_threats) > 0:
        fig.add_trace(go.Scatter3d(
            x=df_threats['Initial_Likelihood'] * 100,
            y=np.abs(df_threats['Initial risk_Value']) / 1e6,
            z=df_threats['Initial_EV'] / 1e6,
            mode='markers',
            marker=dict(
                size=np.clip(np.abs(df_threats['Initial_EV']) / 1e6 * 2 + 4, 6, 20),
                color='red',
                opacity=0.7
            ),
            name='Initial Threats',
            hovertemplate='<b>Risk %{text}</b><br>EV: %{z:.2f}M<extra></extra>',
            text=df_threats['Risk ID'].astype(str)
        ), row=1, col=1)

    # Initial risks - Opportunities
    if len(df_opportunities) > 0:
        fig.add_trace(go.Scatter3d(
            x=df_opportunities['Initial_Likelihood'] * 100,
            y=np.abs(df_opportunities['Initial risk_Value']) / 1e6,
            z=df_opportunities['Initial_EV'] / 1e6,
            mode='markers',
            marker=dict(
                size=np.clip(np.abs(df_opportunities['Initial_EV']) / 1e6 * 2 + 4, 6, 20),
                color='green',
                opacity=0.7,
                symbol='diamond'
            ),
            name='Initial Opportunities',
            hovertemplate='<b>Risk %{text}</b><br>EV: %{z:.2f}M<extra></extra>',
            text=df_opportunities['Risk ID'].astype(str)
        ), row=1, col=1)

    # Residual risks - Threats
    df_threats_res = df_plot[df_plot['Residual_EV'] > 0]
    df_opportunities_res = df_plot[df_plot['Residual_EV'] < 0]

    if len(df_threats_res) > 0:
        fig.add_trace(go.Scatter3d(
            x=df_threats_res['Residual_Likelihood'] * 100,
            y=np.abs(df_threats_res['Residual risk_Value']) / 1e6,
            z=df_threats_res['Residual_EV'] / 1e6,
            mode='markers',
            marker=dict(
                size=np.clip(np.abs(df_threats_res['Residual_EV']) / 1e6 * 2 + 4, 6, 20),
                color='salmon',
                opacity=0.7
            ),
            name='Residual Threats',
            showlegend=True,
            hovertemplate='<b>Risk %{text}</b><br>EV: %{z:.2f}M<extra></extra>',
            text=df_threats_res['Risk ID'].astype(str)
        ), row=1, col=2)

    if len(df_opportunities_res) > 0:
        fig.add_trace(go.Scatter3d(
            x=df_opportunities_res['Residual_Likelihood'] * 100,
            y=np.abs(df_opportunities_res['Residual risk_Value']) / 1e6,
            z=df_opportunities_res['Residual_EV'] / 1e6,
            mode='markers',
            marker=dict(
                size=np.clip(np.abs(df_opportunities_res['Residual_EV']) / 1e6 * 2 + 4, 6, 20),
                color='lightgreen',
                opacity=0.7,
                symbol='diamond'
            ),
            name='Residual Opportunities',
            showlegend=True,
            hovertemplate='<b>Risk %{text}</b><br>EV: %{z:.2f}M<extra></extra>',
            text=df_opportunities_res['Risk ID'].astype(str)
        ), row=1, col=2)

    # Calculate common axis ranges
    max_likelihood = 100
    max_impact = max(
        np.abs(df_plot['Initial risk_Value']).max(),
        np.abs(df_plot['Residual risk_Value']).max()
    ) / 1e6 * 1.1
    ev_range = [
        min(df_plot['Initial_EV'].min(), df_plot['Residual_EV'].min()) / 1e6 * 1.1,
        max(df_plot['Initial_EV'].max(), df_plot['Residual_EV'].max()) / 1e6 * 1.1
    ]

    # Update scene for both subplots
    scene_config = dict(
        xaxis=dict(title='Likelihood (%)', range=[0, max_likelihood]),
        yaxis=dict(title='Impact (M CHF)', range=[0, max_impact]),
        zaxis=dict(title='EV (M CHF)', range=ev_range),
        camera=dict(eye=dict(x=1.5, y=1.5, z=1.0))
    )

    fig.update_layout(
        scene=scene_config,
        scene2=scene_config,
        title=dict(
            text='3D Risk Landscape Comparison: Before vs After Mitigation',
            font=dict(size=16, color='#1F4E78')
        ),
        height=600,
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.15,
            xanchor="center",
            x=0.5
        )
    )

    return fig

def create_confidence_bar_chart(confidence_data, selected_confidence='P80'):
    """
    Create horizontal bar chart showing contingency at different confidence levels.

    Args:
        confidence_data: Dictionary from calculate_confidence_comparison()
        selected_confidence: Currently selected confidence level to highlight
    """
    levels = ['P50', 'P80', 'P90', 'P95']
    contingencies = [confidence_data['total_contingency'][l] / 1e6 for l in levels]

    # Create colors - highlight selected level
    colors = []
    for level in levels:
        if level == selected_confidence:
            colors.append('#1F4E78')  # Dark blue for selected
        elif level == 'P50':
            colors.append('#2ECC71')  # Green for P50
        elif level == 'P95':
            colors.append('#E74C3C')  # Red for P95
        else:
            colors.append('#3498DB')  # Light blue for others

    fig = go.Figure()

    fig.add_trace(go.Bar(
        y=levels,
        x=contingencies,
        orientation='h',
        marker=dict(
            color=colors,
            line=dict(color='black', width=1)
        ),
        text=[f'{v:.1f}M CHF' for v in contingencies],
        textposition='outside',
        hovertemplate='<b>%{y}</b><br>Total Contingency: %{x:.2f}M CHF<extra></extra>'
    ))

    # Add star marker for selected confidence
    selected_idx = levels.index(selected_confidence) if selected_confidence in levels else 1
    fig.add_annotation(
        x=contingencies[selected_idx],
        y=selected_confidence,
        text='★ SELECTED',
        showarrow=True,
        arrowhead=2,
        arrowsize=1,
        arrowwidth=2,
        ax=50,
        ay=0,
        font=dict(size=12, color='#1F4E78', family='Arial Black')
    )

    fig.update_layout(
        title=dict(
            text='Total Contingency by Confidence Level<br><sub>Residual Exposure + Mitigation Cost</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Total Contingency (M CHF)',
        yaxis_title='Confidence Level',
        height=350,
        margin=dict(l=80, r=120, t=80, b=50),
        showlegend=False
    )

    return fig

def create_confidence_cdf_chart(results, confidence_data, selected_confidence='P80'):
    """
    Create CDF chart with multiple confidence level markers.

    Args:
        results: Monte Carlo simulation results array
        confidence_data: Dictionary from calculate_confidence_comparison()
        selected_confidence: Currently selected confidence level to highlight
    """
    sorted_results = np.sort(results)
    cdf = np.arange(1, len(sorted_results) + 1) / len(sorted_results) * 100

    fig = go.Figure()

    # Add CDF line
    fig.add_trace(go.Scatter(
        x=sorted_results / 1e6,
        y=cdf,
        mode='lines',
        name='CDF',
        line=dict(color='#3498DB', width=3),
        hovertemplate='Value: %{x:.2f}M CHF<br>Percentile: %{y:.1f}%<extra></extra>'
    ))

    # Add percentile markers
    percentile_configs = [
        ('P50', 50, '#2ECC71', 'solid'),
        ('P80', 80, '#F39C12', 'solid'),
        ('P90', 90, '#E67E22', 'dash'),
        ('P95', 95, '#E74C3C', 'solid')
    ]

    for level, pct, color, dash in percentile_configs:
        value = confidence_data['percentiles'][level] / 1e6
        line_width = 4 if level == selected_confidence else 2

        # Vertical line
        fig.add_vline(
            x=value,
            line=dict(color=color, width=line_width, dash=dash),
            annotation_text=f'{level}: {value:.1f}M' + (' ★' if level == selected_confidence else ''),
            annotation_position='top',
            annotation_font=dict(size=10, color=color)
        )

        # Horizontal line to y-axis
        fig.add_hline(
            y=pct,
            line=dict(color=color, width=1, dash='dot'),
        )

    fig.update_layout(
        title=dict(
            text='Cumulative Distribution Function with Confidence Levels<br><sub>Vertical lines show contingency at each percentile</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Risk Exposure (M CHF)',
        yaxis_title='Cumulative Probability (%)',
        height=400,
        yaxis=dict(range=[0, 105]),
        showlegend=False
    )

    return fig

def create_cost_confidence_curve(confidence_data):
    """
    Create line chart showing cost-confidence relationship (diminishing returns).

    Args:
        confidence_data: Dictionary from calculate_confidence_comparison()
    """
    # Data points
    levels = ['P50', 'P80', 'P90', 'P95']
    percentiles = [50, 80, 90, 95]
    contingencies = [confidence_data['total_contingency'][l] / 1e6 for l in levels]

    fig = go.Figure()

    # Main line
    fig.add_trace(go.Scatter(
        x=percentiles,
        y=contingencies,
        mode='lines+markers',
        name='Total Contingency',
        line=dict(color='#1F4E78', width=3),
        marker=dict(size=12, color='#1F4E78', line=dict(width=2, color='white')),
        hovertemplate='<b>P%{x}</b><br>Contingency: %{y:.2f}M CHF<extra></extra>'
    ))

    # Add incremental cost annotations
    for i in range(1, len(levels)):
        delta = contingencies[i] - contingencies[i-1]
        mid_x = (percentiles[i] + percentiles[i-1]) / 2
        mid_y = (contingencies[i] + contingencies[i-1]) / 2

        fig.add_annotation(
            x=mid_x,
            y=mid_y + (contingencies[-1] - contingencies[0]) * 0.08,
            text=f'+{delta:.1f}M',
            showarrow=False,
            font=dict(size=10, color='#E74C3C'),
            bgcolor='rgba(255,255,255,0.8)'
        )

    # Add diminishing returns shading
    fig.add_trace(go.Scatter(
        x=percentiles + percentiles[::-1],
        y=contingencies + [contingencies[0]] * len(contingencies),
        fill='toself',
        fillcolor='rgba(52, 152, 219, 0.1)',
        line=dict(color='rgba(0,0,0,0)'),
        showlegend=False,
        hoverinfo='skip'
    ))

    # Add reference line for P50 baseline
    fig.add_hline(
        y=contingencies[0],
        line=dict(color='#2ECC71', width=2, dash='dash'),
        annotation_text=f'P50 Baseline: {contingencies[0]:.1f}M',
        annotation_position='bottom right',
        annotation_font=dict(size=10, color='#2ECC71')
    )

    fig.update_layout(
        title=dict(
            text='Cost-Confidence Trade-off Curve<br><sub>Shows diminishing returns as confidence increases</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Confidence Level (%)',
        yaxis_title='Total Contingency (M CHF)',
        xaxis=dict(
            tickmode='array',
            tickvals=percentiles,
            ticktext=['P50\n(Median)', 'P80\n(Standard)', 'P90\n(Conservative)', 'P95\n(Highly\nConservative)']
        ),
        height=400,
        showlegend=False
    )

    return fig

def create_incremental_cost_chart(confidence_data):
    """
    Create bar chart showing incremental cost for each confidence step.

    Args:
        confidence_data: Dictionary from calculate_confidence_comparison()
    """
    incremental = confidence_data['incremental_analysis']

    steps = [item['step'] for item in incremental]
    costs = [item['additional_cost'] / 1e6 for item in incremental]
    cost_per_pct = [item['cost_per_1pct'] / 1e6 for item in incremental]

    fig = go.Figure()

    # Additional cost bars
    fig.add_trace(go.Bar(
        x=steps,
        y=costs,
        name='Additional Cost',
        marker=dict(
            color=['#3498DB', '#F39C12', '#E74C3C'],
            line=dict(color='black', width=1)
        ),
        text=[f'+{c:.1f}M' for c in costs],
        textposition='outside',
        hovertemplate='<b>%{x}</b><br>Additional Cost: %{y:.2f}M CHF<extra></extra>'
    ))

    # Add cost per % annotation
    for i, (step, cpc) in enumerate(zip(steps, cost_per_pct)):
        fig.add_annotation(
            x=step,
            y=costs[i] * 0.5,
            text=f'{cpc:.2f}M/1%',
            showarrow=False,
            font=dict(size=10, color='white', family='Arial Black'),
            bgcolor='rgba(0,0,0,0.5)'
        )

    fig.update_layout(
        title=dict(
            text='Incremental Cost per Confidence Step<br><sub>Cost increases accelerate at higher confidence levels</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Confidence Step',
        yaxis_title='Additional Cost (M CHF)',
        height=350,
        showlegend=False
    )

    return fig

# =============================================================================
# TIME-PHASED CONTINGENCY VISUALIZATION FUNCTIONS
# =============================================================================

def create_phase_allocation_bar_chart(phase_allocation_data, show_cumulative=False):
    """
    Create stacked/grouped bar chart showing contingency allocation by phase.

    Args:
        phase_allocation_data: Dictionary from calculate_phase_allocation()
        show_cumulative: If True, show cumulative values instead of individual

    Returns:
        Plotly figure
    """
    phase_stats = phase_allocation_data['phase_stats']
    confidence_level = phase_allocation_data['confidence_level']

    # Sort phases by order
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

    phases = [stats['name'] for _, stats in sorted_phases]
    phase_codes = [code for code, _ in sorted_phases]
    colors = [stats['color'] for _, stats in sorted_phases]

    if show_cumulative:
        ev_values = [stats['cumulative_ev'] / 1e6 for _, stats in sorted_phases]
        conf_values = [stats['cumulative_at_confidence'] / 1e6 for _, stats in sorted_phases]
        title = f'Cumulative Contingency Allocation by Phase ({confidence_level})'
    else:
        ev_values = [stats['expected_value'] / 1e6 for _, stats in sorted_phases]
        conf_values = [stats['at_confidence'] / 1e6 for _, stats in sorted_phases]
        title = f'Contingency Allocation by Phase ({confidence_level})'

    fig = go.Figure()

    # Expected Value bars
    fig.add_trace(go.Bar(
        x=phases,
        y=ev_values,
        name='Expected Value',
        marker_color='rgba(52, 152, 219, 0.7)',
        text=[f'{v:.2f}M' for v in ev_values],
        textposition='outside',
        hovertemplate='<b>%{x}</b><br>Expected Value: %{y:.2f}M CHF<extra></extra>'
    ))

    # At Confidence bars
    fig.add_trace(go.Bar(
        x=phases,
        y=conf_values,
        name=f'{confidence_level} Value',
        marker_color='rgba(231, 76, 60, 0.7)',
        text=[f'{v:.2f}M' for v in conf_values],
        textposition='outside',
        hovertemplate=f'<b>%{{x}}</b><br>{confidence_level}: %{{y:.2f}}M CHF<extra></extra>'
    ))

    fig.update_layout(
        title=dict(
            text=f'{title}<br><sub>Comparison of Expected Value vs {confidence_level} Contingency</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Project Phase',
        yaxis_title='Contingency Amount (M CHF)',
        barmode='group',
        height=450,
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=1.02,
            xanchor='right',
            x=1
        )
    )

    return fig

def create_phase_scurve_chart(phase_allocation_data, show_burndown=True):
    """
    Create S-Curve showing cumulative contingency allocation/consumption.

    Args:
        phase_allocation_data: Dictionary from calculate_phase_allocation()
        show_burndown: If True, show remaining contingency (burn-down)

    Returns:
        Plotly figure
    """
    phase_stats = phase_allocation_data['phase_stats']
    confidence_level = phase_allocation_data['confidence_level']
    total_at_confidence = phase_allocation_data['total_at_confidence']

    # Sort phases by order
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

    # Build data points (including start point)
    phases = ['Start'] + [stats['name'] for _, stats in sorted_phases]
    phase_codes = [''] + [code for code, _ in sorted_phases]

    # Cumulative allocation
    cumulative_values = [0]
    for _, stats in sorted_phases:
        cumulative_values.append(stats['cumulative_at_confidence'])

    # Remaining contingency (burn-down perspective)
    remaining_values = [total_at_confidence - cv for cv in cumulative_values]

    # Percentages
    cumulative_pct = [(cv / total_at_confidence * 100) if total_at_confidence > 0 else 0 for cv in cumulative_values]
    remaining_pct = [100 - cp for cp in cumulative_pct]

    fig = make_subplots(specs=[[{"secondary_y": True}]])

    if show_burndown:
        # Burn-down line (remaining contingency)
        fig.add_trace(
            go.Scatter(
                x=phases,
                y=[rv / 1e6 for rv in remaining_values],
                mode='lines+markers+text',
                name='Remaining Contingency',
                line=dict(color='#E74C3C', width=3),
                marker=dict(size=10),
                text=[f'{rp:.0f}%' for rp in remaining_pct],
                textposition='top center',
                hovertemplate='<b>%{x}</b><br>Remaining: %{y:.2f}M CHF (%{text})<extra></extra>'
            ),
            secondary_y=False
        )
    else:
        # Cumulative allocation line
        fig.add_trace(
            go.Scatter(
                x=phases,
                y=[cv / 1e6 for cv in cumulative_values],
                mode='lines+markers+text',
                name='Cumulative Allocation',
                line=dict(color='#3498DB', width=3),
                marker=dict(size=10),
                text=[f'{cp:.0f}%' for cp in cumulative_pct],
                textposition='top center',
                hovertemplate='<b>%{x}</b><br>Allocated: %{y:.2f}M CHF (%{text})<extra></extra>'
            ),
            secondary_y=False
        )

    # Add phase-specific allocation as bars
    phase_allocations = [0] + [stats['at_confidence'] / 1e6 for _, stats in sorted_phases]
    fig.add_trace(
        go.Bar(
            x=phases,
            y=phase_allocations,
            name='Phase Allocation',
            marker_color='rgba(52, 152, 219, 0.3)',
            hovertemplate='<b>%{x}</b><br>Phase Amount: %{y:.2f}M CHF<extra></extra>'
        ),
        secondary_y=False
    )

    title = 'Contingency Burn-Down Curve' if show_burndown else 'Cumulative Contingency S-Curve'

    fig.update_layout(
        title=dict(
            text=f'{title} ({confidence_level})<br><sub>Shows contingency {"depletion" if show_burndown else "allocation"} across project phases</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Project Phase',
        height=450,
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=1.02,
            xanchor='right',
            x=1
        ),
        hovermode='x unified'
    )

    fig.update_yaxes(title_text='Contingency Amount (M CHF)', secondary_y=False)

    return fig

def create_phase_waterfall_chart(phase_allocation_data):
    """
    Create waterfall chart showing phase-by-phase contribution to total contingency.

    Args:
        phase_allocation_data: Dictionary from calculate_phase_allocation()

    Returns:
        Plotly figure
    """
    phase_stats = phase_allocation_data['phase_stats']
    confidence_level = phase_allocation_data['confidence_level']
    total_at_confidence = phase_allocation_data['total_at_confidence']

    # Sort phases by order
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

    # Prepare waterfall data
    labels = [stats['name'] for _, stats in sorted_phases] + ['Total']
    values = [stats['at_confidence'] / 1e6 for _, stats in sorted_phases]
    values.append(total_at_confidence / 1e6)

    measures = ['relative'] * len(sorted_phases) + ['total']
    colors = [stats['color'] for _, stats in sorted_phases] + ['#2C3E50']

    fig = go.Figure(go.Waterfall(
        name='Phase Contribution',
        orientation='v',
        measure=measures,
        x=labels,
        y=values[:-1] + [0],  # Total is calculated automatically
        textposition='outside',
        text=[f'{v:.2f}M' for v in values[:-1]] + [f'{total_at_confidence/1e6:.2f}M'],
        connector=dict(line=dict(color='rgb(63, 63, 63)')),
        increasing=dict(marker=dict(color='rgba(52, 152, 219, 0.7)')),
        totals=dict(marker=dict(color='#2C3E50')),
        hovertemplate='<b>%{x}</b><br>Contribution: %{y:.2f}M CHF<extra></extra>'
    ))

    # Add percentage annotations
    cumulative = 0
    for i, (_, stats) in enumerate(sorted_phases):
        pct = stats['confidence_percentage']
        cumulative += stats['at_confidence']
        fig.add_annotation(
            x=i,
            y=cumulative / 1e6,
            text=f'{pct:.1f}%',
            showarrow=False,
            font=dict(size=9, color='#666'),
            yshift=5
        )

    fig.update_layout(
        title=dict(
            text=f'Phase Contribution to Total Contingency ({confidence_level})<br><sub>Waterfall showing each phase\'s addition to contingency reserve</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Project Phase',
        yaxis_title='Contingency Contribution (M CHF)',
        height=450,
        showlegend=False
    )

    return fig

def create_early_warning_gauge(warning_status, phase_name):
    """
    Create gauge chart for early warning indicator.

    Args:
        warning_status: Dictionary from get_early_warning_status()
        phase_name: Name of the phase

    Returns:
        Plotly figure
    """
    status = warning_status['status']
    consumed_pct = warning_status['consumed_pct']
    expected_pct = warning_status['expected_pct']
    deviation = warning_status['deviation']

    # Color based on status
    status_colors = {
        'Green': '#27AE60',
        'Amber': '#F39C12',
        'Red': '#E74C3C'
    }
    color = status_colors.get(status, '#95A5A6')

    fig = go.Figure(go.Indicator(
        mode='gauge+number+delta',
        value=consumed_pct,
        delta={'reference': expected_pct, 'relative': False, 'valueformat': '.1f'},
        title={'text': f'{phase_name}<br><sub>Contingency Consumption</sub>'},
        gauge={
            'axis': {'range': [0, 100], 'tickwidth': 1},
            'bar': {'color': color},
            'bgcolor': 'white',
            'borderwidth': 2,
            'bordercolor': '#666',
            'steps': [
                {'range': [0, expected_pct * 0.9], 'color': '#E8F6E8'},
                {'range': [expected_pct * 0.9, expected_pct * 1.1], 'color': '#FFF3CD'},
                {'range': [expected_pct * 1.1, 100], 'color': '#F8D7DA'}
            ],
            'threshold': {
                'line': {'color': '#2C3E50', 'width': 4},
                'thickness': 0.75,
                'value': expected_pct
            }
        },
        number={'suffix': '%', 'valueformat': '.1f'}
    ))

    fig.update_layout(
        height=250,
        margin=dict(l=20, r=20, t=50, b=20)
    )

    return fig

def create_phase_risk_distribution_chart(df, phase_allocation_data):
    """
    Create chart showing number of risks per phase and their distribution.

    Args:
        df: DataFrame with risk data including phase information
        phase_allocation_data: Dictionary from calculate_phase_allocation()

    Returns:
        Plotly figure
    """
    phase_stats = phase_allocation_data['phase_stats']
    phases = phase_allocation_data['phases']

    # Count risks per crystallization phase
    if 'Crystallization Phase' in df.columns:
        phase_counts = df['Crystallization Phase'].value_counts()
    else:
        phase_counts = pd.Series(dtype=int)

    # Sort phases by order
    sorted_phases = sorted(phases.items(), key=lambda x: x[1]['order'])

    phase_names = []
    risk_counts = []
    colors = []
    ev_values = []

    for code, phase_info in sorted_phases:
        phase_names.append(phase_info['name'])
        risk_counts.append(phase_counts.get(code, 0))
        colors.append(phase_info['color'])
        ev_values.append(phase_stats.get(code, {}).get('expected_value', 0) / 1e6)

    fig = make_subplots(specs=[[{"secondary_y": True}]])

    # Risk count bars
    fig.add_trace(
        go.Bar(
            x=phase_names,
            y=risk_counts,
            name='Number of Risks',
            marker_color=colors,
            text=risk_counts,
            textposition='outside',
            hovertemplate='<b>%{x}</b><br>Risks: %{y}<extra></extra>'
        ),
        secondary_y=False
    )

    # Expected value line
    fig.add_trace(
        go.Scatter(
            x=phase_names,
            y=ev_values,
            mode='lines+markers',
            name='Expected Value (M CHF)',
            line=dict(color='#E74C3C', width=2),
            marker=dict(size=8),
            hovertemplate='<b>%{x}</b><br>EV: %{y:.2f}M CHF<extra></extra>'
        ),
        secondary_y=True
    )

    fig.update_layout(
        title=dict(
            text='Risk Distribution by Project Phase<br><sub>Number of risks and their expected value per phase</sub>',
            font=dict(size=14, color='#1F4E78')
        ),
        xaxis_title='Project Phase',
        height=400,
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=1.02,
            xanchor='right',
            x=1
        )
    )

    fig.update_yaxes(title_text='Number of Risks', secondary_y=False)
    fig.update_yaxes(title_text='Expected Value (M CHF)', secondary_y=True)

    return fig

def create_risk_matrix(df, risk_type='initial'):
    """Create interactive risk matrix"""
    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
        title = 'Initial Risk Matrix'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'
        title = 'Residual Risk Matrix'

    # Create a copy to avoid modifying original dataframe
    df_plot = df.copy()
    
    # Create categories for impact (using absolute values for categorization)
    df_plot['Impact_Category'] = pd.cut(np.abs(df_plot[impact_col]), 
                                    bins=[0, 100000, 1000000, 10000000, np.inf],
                                    labels=['Low', 'Medium', 'High', 'Very High'])
    
    # Create categories for likelihood
    df_plot['Likelihood_Category'] = pd.cut(df_plot[likelihood_col],
                                        bins=[0, 0.2, 0.4, 0.6, 1.0],
                                        labels=['Low', 'Medium', 'High', 'Very High'])
    
    # Create risk score
    df_plot['Risk_Score'] = df_plot[impact_col] * df_plot[likelihood_col]
    
    # Create size for markers (must be positive, use absolute value + small offset)
    df_plot['Marker_Size'] = np.abs(df_plot['Risk_Score']) + 1000  # Add offset to ensure visibility
    
    # Create hover text showing if risk is negative (opportunity)
    df_plot['Risk_Type'] = df_plot[impact_col].apply(lambda x: 'Opportunity' if x < 0 else 'Risk')
    
    fig = px.scatter(df_plot, 
                     x=likelihood_col, 
                     y=impact_col,
                     size='Marker_Size',
                     color='Risk_Score',
                     hover_data=['Risk ID', 'Risk Description', 'Risk_Type'],
                     labels={likelihood_col: 'Likelihood', 
                            impact_col: 'Impact (CHF)'},
                     title=title,
                     color_continuous_scale='RdYlGn_r')
    
    # Add grid lines for zones
    fig.add_hline(y=0, line_dash="solid", line_color="black", opacity=0.3)  # Zero line
    fig.add_hline(y=1000000, line_dash="dash", line_color="gray", opacity=0.5)
    fig.add_hline(y=10000000, line_dash="dash", line_color="gray", opacity=0.5)
    fig.add_hline(y=-1000000, line_dash="dash", line_color="gray", opacity=0.5)
    fig.add_vline(x=0.2, line_dash="dash", line_color="gray", opacity=0.5)
    fig.add_vline(x=0.4, line_dash="dash", line_color="gray", opacity=0.5)
    fig.add_vline(x=0.6, line_dash="dash", line_color="gray", opacity=0.5)
    
    fig.update_layout(height=600, showlegend=True)
    
    return fig

def create_cdf_plot(results, stats, risk_type, selected_confidence='P95'):
    """Create cumulative distribution function plot with highlighted confidence level"""
    sorted_results = np.sort(results)
    cdf = np.arange(1, len(sorted_results) + 1) / len(sorted_results)

    fig = go.Figure()

    # CDF line
    fig.add_trace(go.Scatter(
        x=sorted_results,
        y=cdf * 100,
        mode='lines',
        name='CDF',
        line=dict(color='blue', width=2)
    ))

    # Add percentile lines with highlighting for selected confidence level
    percentiles = [('P50', stats['p50'], 'green'),
                   ('P80', stats['p80'], 'orange'),
                   ('P95', stats['p95'], 'red')]

    for label, value, color in percentiles:
        # Make selected confidence level more prominent
        if label == selected_confidence:
            line_width = 4
            line_dash = "solid"
            annotation_text = f"★ {label}: {value/1e6:.2f}M CHF (ACTIVE)"
        else:
            line_width = 2
            line_dash = "dash"
            annotation_text = f"{label}: {value/1e6:.2f}M CHF"

        fig.add_vline(x=value, line_dash=line_dash, line_color=color, line_width=line_width,
                     annotation_text=annotation_text,
                     annotation_position="top")

    fig.update_layout(
        title=f'Cumulative Distribution Function - {risk_type.title()} Risk Exposure',
        xaxis_title='Total Risk Exposure (CHF)',
        yaxis_title='Cumulative Probability (%)',
        height=500,
        hovermode='x unified'
    )

    return fig

def create_box_plot(initial_results, residual_results):
    """Create box plot comparing initial vs residual risk"""
    fig = go.Figure()
    
    fig.add_trace(go.Box(
        y=initial_results / 1e6,
        name='Initial Risk',
        marker_color='indianred'
    ))
    
    fig.add_trace(go.Box(
        y=residual_results / 1e6,
        name='Residual Risk',
        marker_color='lightseagreen'
    ))
    
    fig.update_layout(
        title='Risk Exposure Distribution: Initial vs Residual',
        yaxis_title='Total Risk Exposure (Million CHF)',
        height=500,
        showlegend=True
    )
    
    return fig

def create_histogram(results, stats, risk_type):
    """Create histogram with statistics overlay"""
    fig = go.Figure()
    
    fig.add_trace(go.Histogram(
        x=results / 1e6,
        nbinsx=50,
        name='Frequency',
        marker_color='lightblue',
        opacity=0.7
    ))
    
    # Add mean line
    fig.add_vline(x=stats['mean']/1e6, line_dash="dash", line_color="blue",
                 annotation_text=f"Mean: {stats['mean']/1e6:.2f}M",
                 annotation_position="top right")
    
    # Add median line
    fig.add_vline(x=stats['median']/1e6, line_dash="dash", line_color="green",
                 annotation_text=f"Median: {stats['median']/1e6:.2f}M",
                 annotation_position="top left")
    
    fig.update_layout(
        title=f'Distribution of Total Risk Exposure - {risk_type.title()}',
        xaxis_title='Total Risk Exposure (Million CHF)',
        yaxis_title='Frequency',
        height=500,
        showlegend=True
    )
    
    return fig

def calculate_mitigation_roi(df):
    """Calculate ROI for risk mitigation measures"""
    # Calculate initial expected value
    df['Initial_EV'] = df['Initial risk_Value'] * df['Initial_Likelihood']

    # Calculate residual expected value
    df['Residual_EV'] = df['Residual risk_Value'] * df['Residual_Likelihood']

    # Calculate risk reduction
    df['Risk_Reduction'] = df['Initial_EV'] - df['Residual_EV']

    # Classify as Threat or Opportunity based on Initial_EV
    # Threats have positive EV (costs), Opportunities have negative EV (benefits)
    df['Risk_Type'] = df['Initial_EV'].apply(lambda x: 'Opportunity' if x < 0 else 'Threat')

    # Calculate ROI (only for risks with mitigation cost > 0)
    df['ROI'] = np.where(
        df['Cost of Measures_Value'] > 0,
        (df['Risk_Reduction'] - df['Cost of Measures_Value']) / df['Cost of Measures_Value'] * 100,
        0
    )

    # Calculate benefit-cost ratio
    df['BC_Ratio'] = np.where(
        df['Cost of Measures_Value'] > 0,
        df['Risk_Reduction'] / df['Cost of Measures_Value'],
        0
    )

    return df

def calculate_threat_opportunity_metrics(df):
    """
    Calculate separate metrics for threats and opportunities.

    In quantitative risk analysis:
    - Threats: Positive EV (potential costs/losses)
    - Opportunities: Negative EV (potential benefits/savings)
    - Net Exposure: Threat EV + Opportunity EV (where Opportunity EV is negative)

    Returns:
        dict with threat_ev, opportunity_ev, net_exposure, and counts
    """
    # Separate threats and opportunities
    df_threats = df[df['Initial_EV'] > 0]
    df_opportunities = df[df['Initial_EV'] < 0]

    # Calculate EVs (Opportunity EV will be negative)
    threat_initial_ev = df_threats['Initial_EV'].sum()
    opportunity_initial_ev = df_opportunities['Initial_EV'].sum()  # This is negative

    threat_residual_ev = df_threats['Residual_EV'].sum()
    opportunity_residual_ev = df_opportunities['Residual_EV'].sum()  # This is negative

    # Net exposure = Threats + Opportunities (opportunities reduce the total)
    net_initial_exposure = threat_initial_ev + opportunity_initial_ev
    net_residual_exposure = threat_residual_ev + opportunity_residual_ev

    # Calculate risk reduction for each category
    threat_reduction = threat_initial_ev - threat_residual_ev
    # For opportunities: "improvement" means opportunity EV becomes more negative (larger benefit)
    opportunity_change = opportunity_initial_ev - opportunity_residual_ev

    return {
        'threat_count': len(df_threats),
        'opportunity_count': len(df_opportunities),
        'threat_initial_ev': threat_initial_ev,
        'opportunity_initial_ev': opportunity_initial_ev,  # Negative value
        'threat_residual_ev': threat_residual_ev,
        'opportunity_residual_ev': opportunity_residual_ev,  # Negative value
        'net_initial_exposure': net_initial_exposure,
        'net_residual_exposure': net_residual_exposure,
        'threat_reduction': threat_reduction,
        'opportunity_change': opportunity_change,
        'net_reduction': net_initial_exposure - net_residual_exposure
    }

def generate_risk_narrative(df, initial_stats, residual_stats, confidence_level, sensitivity_df=None):
    """
    Generate executive risk narrative using templates.

    Args:
        df: Risk register DataFrame
        initial_stats: Statistics dictionary for initial risk
        residual_stats: Statistics dictionary for residual risk
        confidence_level: Selected confidence level (P50/P80/P95)
        sensitivity_df: Optional sensitivity analysis DataFrame

    Returns:
        Dictionary with narrative sections
    """
    # Get threat/opportunity metrics
    to_metrics = calculate_threat_opportunity_metrics(df)

    # Calculate key metrics
    total_risks = len(df)
    threat_count = to_metrics['threat_count']
    opportunity_count = to_metrics['opportunity_count']

    # Get confidence level values
    conf_key = confidence_level.lower()
    initial_conf_value = initial_stats.get(conf_key, initial_stats.get('p80', 0))
    residual_conf_value = residual_stats.get(conf_key, residual_stats.get('p80', 0))

    # Calculate risk reduction
    if initial_conf_value > 0:
        risk_reduction_pct = ((initial_conf_value - residual_conf_value) / initial_conf_value) * 100
    else:
        risk_reduction_pct = 0

    # Identify top risks (by absolute EV)
    df_sorted = df.copy()
    df_sorted['Abs_EV'] = np.abs(df_sorted['Initial_EV'])
    top_risks = df_sorted.nlargest(5, 'Abs_EV')

    # Identify high-priority risks (high likelihood AND high impact)
    high_priority = df[(df['Initial_Likelihood'] >= 0.5) & (np.abs(df['Initial risk_Value']) >= 5000000)]

    # Build executive summary
    executive_summary = f"""The risk portfolio comprises {total_risks} identified risks, consisting of {threat_count} threats and {opportunity_count} opportunities. At the {confidence_level} confidence level, the total risk exposure is {initial_conf_value/1e6:.1f}M CHF before mitigation and {residual_conf_value/1e6:.1f}M CHF after mitigation measures are applied.

The net risk exposure (threats minus opportunities) stands at {to_metrics['net_initial_exposure']/1e6:.1f}M CHF initially, reducing to {to_metrics['net_residual_exposure']/1e6:.1f}M CHF after mitigation—representing an overall reduction of {risk_reduction_pct:.1f}%."""

    # Build critical findings
    top_risk = top_risks.iloc[0] if len(top_risks) > 0 else None
    if top_risk is not None:
        risk_type = "threat" if top_risk['Initial_EV'] > 0 else "opportunity"
        critical_findings = f"""**Highest-Impact Risk**: "{top_risk['Risk Description']}" (Risk ID: {top_risk['Risk ID']}) is the most significant {risk_type} with an expected value of {top_risk['Initial_EV']/1e6:.2f}M CHF.

**High-Priority Risks**: {len(high_priority)} risks have been identified as high-priority (likelihood ≥50% AND impact ≥5M CHF), requiring immediate attention.

**Threat Exposure**: Total threat exposure is {to_metrics['threat_initial_ev']/1e6:.1f}M CHF, with planned mitigations reducing this by {to_metrics['threat_reduction']/1e6:.1f}M CHF ({(to_metrics['threat_reduction']/to_metrics['threat_initial_ev']*100) if to_metrics['threat_initial_ev'] > 0 else 0:.1f}%).

**Opportunity Potential**: {opportunity_count} opportunities have been identified with a combined potential benefit of {abs(to_metrics['opportunity_initial_ev'])/1e6:.1f}M CHF."""
    else:
        critical_findings = "No risks identified in the portfolio."

    # Build top 5 risks summary
    top_risks_summary = "**Top 5 Risks by Expected Value:**\n\n"
    for i, (_, risk) in enumerate(top_risks.iterrows(), 1):
        risk_type_icon = "⚠️" if risk['Initial_EV'] > 0 else "✅"
        top_risks_summary += f"{i}. {risk_type_icon} **Risk {risk['Risk ID']}**: {risk['Risk Description'][:80]}{'...' if len(risk['Risk Description']) > 80 else ''}\n"
        top_risks_summary += f"   - Expected Value: {risk['Initial_EV']/1e6:.2f}M CHF | Likelihood: {risk['Initial_Likelihood']*100:.0f}%\n\n"

    # Build sensitivity insights if available
    if sensitivity_df is not None and len(sensitivity_df) > 0:
        top_driver = sensitivity_df.iloc[0]
        risks_80_pct = (sensitivity_df['Cumulative %'] <= 80).sum()
        sensitivity_insights = f"""**Key Risk Drivers**: The top risk driver (Risk {top_driver['Risk ID']}) accounts for {top_driver['Variance %']:.1f}% of total portfolio variance. Just {risks_80_pct} risks drive 80% of the overall uncertainty, following the Pareto principle."""
    else:
        sensitivity_insights = ""

    # Build recommendations
    recommendations = []

    if len(high_priority) >= 3:
        recommendations.append("**Prioritize High-Impact Risks**: Focus mitigation efforts on the " +
                              f"{len(high_priority)} high-priority risks that pose the greatest threat to project objectives.")

    if risk_reduction_pct < 30:
        recommendations.append("**Enhance Mitigation Strategies**: Current mitigation measures achieve only " +
                              f"{risk_reduction_pct:.1f}% risk reduction. Consider strengthening mitigation plans for top risks.")
    elif risk_reduction_pct >= 50:
        recommendations.append("**Maintain Mitigation Momentum**: Current strategies achieve " +
                              f"{risk_reduction_pct:.1f}% risk reduction. Continue monitoring and executing planned mitigations.")

    if opportunity_count > 0:
        recommendations.append(f"**Capitalize on Opportunities**: {opportunity_count} opportunities have been identified. " +
                              "Develop action plans to maximize potential benefits.")

    if sensitivity_df is not None and len(sensitivity_df) > 0:
        risks_80_pct = (sensitivity_df['Cumulative %'] <= 80).sum()
        recommendations.append(f"**Focus on Key Drivers**: {risks_80_pct} risks drive 80% of uncertainty. " +
                              "Concentrate resources on these high-impact items.")

    recommendations_text = "\n\n".join(recommendations) if recommendations else "No specific recommendations at this time."

    # Build mitigation effectiveness summary
    mitigation_summary = f"""**Mitigation Effectiveness Analysis**:

| Metric | Initial | Residual | Change |
|--------|---------|----------|--------|
| Threat Exposure | {to_metrics['threat_initial_ev']/1e6:.1f}M | {to_metrics['threat_residual_ev']/1e6:.1f}M | -{to_metrics['threat_reduction']/1e6:.1f}M |
| Opportunity Value | {to_metrics['opportunity_initial_ev']/1e6:.1f}M | {to_metrics['opportunity_residual_ev']/1e6:.1f}M | {to_metrics['opportunity_change']/1e6:+.1f}M |
| Net Exposure | {to_metrics['net_initial_exposure']/1e6:.1f}M | {to_metrics['net_residual_exposure']/1e6:.1f}M | -{to_metrics['net_reduction']/1e6:.1f}M |

Overall risk reduction effectiveness: **{risk_reduction_pct:.1f}%**"""

    return {
        'executive_summary': executive_summary,
        'critical_findings': critical_findings,
        'top_risks_summary': top_risks_summary,
        'sensitivity_insights': sensitivity_insights,
        'recommendations': recommendations_text,
        'mitigation_summary': mitigation_summary,
        'risk_reduction_pct': risk_reduction_pct,
        'total_risks': total_risks,
        'threat_count': threat_count,
        'opportunity_count': opportunity_count
    }

def create_tornado_chart(df, top_n=15):
    """Create tornado chart for sensitivity analysis"""
    # Calculate risk range (impact * likelihood) - use absolute value for sorting
    df_temp = df.copy()
    df_temp['Risk_Range'] = df_temp['Initial risk_Value'] * df_temp['Initial_Likelihood']
    df_temp['Risk_Range_Abs'] = np.abs(df_temp['Risk_Range'])
    
    # Sort by absolute risk range and take top N
    df_sorted = df_temp.nlargest(top_n, 'Risk_Range_Abs')
    
    # Sort by actual risk range for display
    df_sorted = df_sorted.sort_values('Risk_Range')
    
    # Determine color based on whether it's positive or negative
    colors = ['green' if x < 0 else 'red' for x in df_sorted['Risk_Range']]
    
    fig = go.Figure()
    
    fig.add_trace(go.Bar(
        y=df_sorted['Risk Description'],
        x=df_sorted['Risk_Range'] / 1e6,
        orientation='h',
        marker=dict(
            color=colors,
            line=dict(color='rgba(0,0,0,0.3)', width=1)
        ),
        text=df_sorted['Risk_Range'].apply(lambda x: f'{x/1e6:.2f}M'),
        textposition='outside',
        hovertemplate='<b>%{y}</b><br>Expected Value: %{x:.2f}M CHF<extra></extra>'
    ))
    
    fig.update_layout(
        title=f'Top {top_n} Risks by Expected Value (Tornado Chart)<br><sub>Red = Risk (Cost), Green = Opportunity (Benefit)</sub>',
        xaxis_title='Expected Value (Million CHF)',
        yaxis_title='',
        height=600,
        showlegend=False,
        xaxis=dict(zeroline=True, zerolinewidth=2, zerolinecolor='black')
    )
    
    return fig

def perform_sensitivity_analysis(df, n_simulations=10000):
    """
    Enhanced sensitivity analysis with variance contribution
    Shows which risks drive the most uncertainty in total exposure

    Uses Common Random Numbers (CRN) technique to reduce sampling noise
    and ensure accurate variance decomposition (no negative contributions)
    """
    # Generate random numbers ONCE for all simulations (Common Random Numbers)
    # This eliminates Monte Carlo sampling noise in variance comparisons
    random_numbers = np.random.random((n_simulations, len(df)))

    # Baseline simulation using the common random numbers
    baseline_results, _ = run_monte_carlo(df, n_simulations, 'initial', random_numbers)
    baseline_variance = np.var(baseline_results)
    baseline_mean = np.mean(baseline_results)

    # Calculate contribution of each risk to total variance
    risk_contributions = []

    for idx, risk in df.iterrows():
        # Create modified dataframe with this risk set to zero probability
        df_modified = df.copy()
        df_modified.loc[idx, 'Initial_Likelihood'] = 0

        # Run simulation without this risk using SAME random numbers
        # This is the key: same random draws, only the likelihood changes
        modified_results, _ = run_monte_carlo(df_modified, n_simulations, 'initial', random_numbers)
        modified_variance = np.var(modified_results)
        modified_mean = np.mean(modified_results)

        # Calculate variance reduction and mean impact
        # With CRN, this should always be >= 0 (or very close due to numerical precision)
        variance_contribution = baseline_variance - modified_variance
        mean_contribution = baseline_mean - modified_mean

        risk_contributions.append({
            'Risk ID': risk['Risk ID'],
            'Risk Description': risk['Risk Description'],
            'Variance Contribution': variance_contribution,
            'Variance %': (variance_contribution / baseline_variance * 100) if baseline_variance > 0 else 0,
            'Mean Contribution': mean_contribution,
            'Expected Value': risk['Initial risk_Value'] * risk['Initial_Likelihood']
        })

    # Create DataFrame and sort by variance contribution
    sensitivity_df = pd.DataFrame(risk_contributions)
    sensitivity_df = sensitivity_df.sort_values('Variance Contribution', ascending=False)

    # Calculate cumulative percentage (Pareto)
    sensitivity_df['Cumulative %'] = sensitivity_df['Variance %'].cumsum()

    return sensitivity_df

def create_enhanced_tornado_chart(sensitivity_df, top_n=15):
    """Create enhanced tornado chart showing variance contribution"""
    # Take top N risks
    df_plot = sensitivity_df.head(top_n).copy()
    
    # Sort by variance contribution for display
    df_plot = df_plot.sort_values('Variance Contribution')
    
    # Create figure with secondary y-axis for cumulative %
    fig = make_subplots(
        rows=1, cols=1,
        specs=[[{"secondary_y": False}]]
    )
    
    # Add variance contribution bars
    fig.add_trace(
        go.Bar(
            y=df_plot['Risk Description'],
            x=df_plot['Variance %'],
            orientation='h',
            name='Variance Contribution',
            marker=dict(
                color=df_plot['Variance %'],
                colorscale='Reds',
                showscale=True,
                colorbar=dict(title="Variance %")
            ),
            text=df_plot['Variance %'].apply(lambda x: f'{x:.1f}%'),
            textposition='outside',
            hovertemplate='<b>%{y}</b><br>Variance Contribution: %{x:.2f}%<br>Mean Impact: %{customdata:.2f}M CHF<extra></extra>',
            customdata=df_plot['Mean Contribution'] / 1e6
        )
    )
    
    fig.update_layout(
        title=f'Sensitivity Analysis - Top {top_n} Risks by Variance Contribution<br><sub>Shows which risks drive the most uncertainty in total exposure</sub>',
        xaxis_title='Contribution to Total Variance (%)',
        yaxis_title='',
        height=600,
        showlegend=False
    )
    
    return fig

def create_pareto_chart(sensitivity_df, top_n=20):
    """Create Pareto chart showing cumulative variance contribution"""
    # Take top N risks - already sorted by variance contribution descending
    df_plot = sensitivity_df.head(top_n).copy()

    # IMPORTANT: Ensure data is sorted by Variance % descending for proper Pareto chart
    # This ensures bars descend from left to right and cumulative curve ascends smoothly
    df_plot = df_plot.sort_values('Variance %', ascending=False).reset_index(drop=True)

    # CRITICAL: Recalculate cumulative % after sorting to ensure correct Pareto curve
    # The cumulative % from sensitivity_df was based on a different order
    df_plot['Cumulative %'] = df_plot['Variance %'].cumsum()

    # CRITICAL FIX: Remove any duplicate Risk IDs to prevent multiple line segments
    # Keep only the first occurrence (highest variance) if duplicates exist
    df_plot = df_plot.drop_duplicates(subset=['Risk ID'], keep='first').reset_index(drop=True)

    # Create truncated risk names for x-axis (max 35 chars for readability)
    df_plot['Risk_Name_Short'] = df_plot['Risk Description'].apply(
        lambda x: x[:35] + '...' if len(x) > 35 else x
    )

    # Prepare hover data with full risk description
    df_plot['hover_text'] = df_plot.apply(
        lambda row: f"<b>{row['Risk ID']}</b>: {row['Risk Description']}", axis=1
    )

    # Create figure WITH secondary y-axis (bars auto-scaled, cumulative 0-100%)
    fig = make_subplots(specs=[[{"secondary_y": True}]])

    # Add bars for individual variance contribution on PRIMARY y-axis (auto-scaled)
    fig.add_trace(
        go.Bar(
            x=df_plot['Risk_Name_Short'],  # Use truncated risk names for readability
            y=df_plot['Variance %'],
            name='Individual Variance %',
            marker_color='#E74C3C',  # Vibrant red
            text=df_plot['Variance %'].apply(lambda x: f'{x:.1f}%'),  # Bar labels
            textposition='outside',  # Labels above bars
            textfont=dict(size=10, color='#E74C3C'),
            customdata=df_plot['hover_text'],
            hovertemplate='%{customdata}<br>Individual Variance: %{y:.2f}%<extra></extra>'
        ),
        secondary_y=False
    )

    # CRITICAL FIX: Add cumulative line as SINGLE trace with explicit connectgaps
    # This ensures ONE continuous line instead of multiple segments
    fig.add_trace(
        go.Scatter(
            x=df_plot['Risk_Name_Short'],  # Same x-axis as bars
            y=df_plot['Cumulative %'],
            name='Cumulative %',
            mode='lines+markers',  # Both lines and markers
            line=dict(color='#2C3E50', width=4),  # Dark slate, thick line
            marker=dict(size=10, symbol='diamond', color='#2C3E50'),
            connectgaps=True,  # Ensure continuous line even if gaps exist
            customdata=df_plot['hover_text'],
            hovertemplate='%{customdata}<br>Cumulative: %{y:.1f}%<extra></extra>'
        ),
        secondary_y=True
    )

    # Add 80% reference line (on secondary y-axis)
    fig.add_hline(
        y=80,
        line_dash="dash",
        line_color="#27AE60",  # Bright green
        line_width=4,  # Thicker for prominence
        annotation_text="80% Threshold",
        annotation_position="right",
        annotation=dict(
            font=dict(size=12, color="#27AE60"),
            bgcolor="rgba(255, 255, 255, 0.9)",
            bordercolor="#27AE60",
            borderwidth=2,
            borderpad=4
        ),
        secondary_y=True
    )

    # Update x-axis with 45-degree rotation and preserve order
    fig.update_xaxes(
        title_text="Risk Description (ordered by variance contribution)",
        tickangle=-45,  # Rotate labels 45 degrees
        tickfont=dict(size=10),
        title_font=dict(size=12),
        categoryorder='array',  # Preserve the exact order from data
        categoryarray=df_plot['Risk_Name_Short'].tolist()  # Explicit order: highest to lowest variance
    )

    # Update PRIMARY y-axis (left) - Individual variance (auto-scaled)
    fig.update_yaxes(
        title_text="Individual Variance Contribution (%)",
        secondary_y=False,
        showgrid=True,  # Subtle gridlines
        gridwidth=1,
        gridcolor='rgba(128, 128, 128, 0.2)',
        tickfont=dict(size=11),
        title_font=dict(size=12)
    )

    # Update SECONDARY y-axis (right) - Cumulative % (0-100% fixed scale)
    fig.update_yaxes(
        title_text="Cumulative Variance Contribution (%)",
        secondary_y=True,
        range=[0, 100],  # Fixed 0-100% scale for cumulative
        showgrid=False,  # No gridlines on secondary axis
        tickfont=dict(size=11),
        title_font=dict(size=12)
    )

    # Update layout with 16x8 inch figure size and proper dimensions
    fig.update_layout(
        title={
            'text': 'Pareto Analysis - Risk Variance Contribution<br><sub>Identify the vital few risks driving most uncertainty (80/20 rule)</sub>',
            'font': {'size': 17}
        },
        width=1536,  # 16 inches * 96 DPI = 1536 pixels
        height=768,  # 8 inches * 96 DPI = 768 pixels
        hovermode='x unified',
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="center",
            x=0.5,
            font=dict(size=11)
        ),
        margin=dict(l=100, r=100, t=140, b=200),  # Larger bottom margin for rotated risk names
        font=dict(size=11),
        plot_bgcolor='white',
        paper_bgcolor='white'
    )

    return fig

# ============================================================================
# WHAT-IF SCENARIO ANALYSIS FUNCTIONS
# ============================================================================

def run_scenario_monte_carlo(df, scenario_adjustments, n_simulations=10000, risk_type='initial'):
    """
    Run Monte Carlo simulation with scenario adjustments applied

    Parameters:
    - df: Original DataFrame with risk data
    - scenario_adjustments: Dict of {risk_id: {'likelihood': float, 'impact': float}}
    - n_simulations: Number of Monte Carlo iterations
    - risk_type: 'initial' or 'residual'

    Returns:
    - results: Array of simulation results
    - modified_df: DataFrame with adjustments applied
    """
    # Create a copy of the dataframe to apply adjustments
    modified_df = df.copy()

    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'

    # Apply scenario adjustments
    for risk_id, adjustments in scenario_adjustments.items():
        mask = modified_df['Risk ID'] == risk_id
        if mask.any():
            if 'likelihood' in adjustments:
                modified_df.loc[mask, likelihood_col] = adjustments['likelihood']
            if 'impact' in adjustments:
                modified_df.loc[mask, impact_col] = adjustments['impact']

    # Prepare data for simulation
    impacts = modified_df[impact_col].values
    likelihoods = modified_df[likelihood_col].values

    # Generate random numbers
    random_numbers = np.random.random((n_simulations, len(modified_df)))

    # Monte Carlo simulation
    results = np.zeros(n_simulations)

    for i in range(n_simulations):
        occurred = random_numbers[i] < likelihoods
        results[i] = np.sum(impacts * occurred)

    return results, modified_df

def create_scenario_comparison_chart(baseline_stats, scenario_stats, scenario_name, confidence_level='P95'):
    """Create a comparison chart between baseline and scenario results"""

    metrics = ['Mean', 'P50', 'P80', 'P95']
    baseline_values = [
        baseline_stats['mean'] / 1e6,
        baseline_stats['p50'] / 1e6,
        baseline_stats['p80'] / 1e6,
        baseline_stats['p95'] / 1e6
    ]
    scenario_values = [
        scenario_stats['mean'] / 1e6,
        scenario_stats['p50'] / 1e6,
        scenario_stats['p80'] / 1e6,
        scenario_stats['p95'] / 1e6
    ]

    fig = go.Figure()

    # Baseline bars
    fig.add_trace(go.Bar(
        name='Baseline',
        x=metrics,
        y=baseline_values,
        marker_color='#3498DB',
        text=[f'{v:.2f}M' for v in baseline_values],
        textposition='outside'
    ))

    # Scenario bars
    fig.add_trace(go.Bar(
        name=scenario_name,
        x=metrics,
        y=scenario_values,
        marker_color='#E74C3C',
        text=[f'{v:.2f}M' for v in scenario_values],
        textposition='outside'
    ))

    # Highlight the selected confidence level
    confidence_idx = metrics.index(confidence_level)

    fig.update_layout(
        title=f'Scenario Comparison: Baseline vs {scenario_name}',
        xaxis_title='Metric',
        yaxis_title='Risk Exposure (Million CHF)',
        barmode='group',
        height=500,
        showlegend=True,
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=1.02,
            xanchor='center',
            x=0.5
        )
    )

    return fig

def create_scenario_cdf_comparison(baseline_results, scenario_results, baseline_stats, scenario_stats,
                                    scenario_name, confidence_level='P95'):
    """Create overlaid CDF curves for baseline vs scenario"""

    # Sort results for CDF
    baseline_sorted = np.sort(baseline_results)
    scenario_sorted = np.sort(scenario_results)

    # Create probability array (0 to 1)
    n = len(baseline_sorted)
    probabilities = np.arange(1, n + 1) / n

    fig = go.Figure()

    # Baseline CDF
    fig.add_trace(go.Scatter(
        x=baseline_sorted / 1e6,
        y=probabilities * 100,
        mode='lines',
        name='Baseline',
        line=dict(color='#3498DB', width=3)
    ))

    # Scenario CDF
    fig.add_trace(go.Scatter(
        x=scenario_sorted / 1e6,
        y=probabilities * 100,
        mode='lines',
        name=scenario_name,
        line=dict(color='#E74C3C', width=3)
    ))

    # Add confidence level lines
    percentile_map = {'P50': 50, 'P80': 80, 'P95': 95}
    conf_percentile = percentile_map[confidence_level]

    # Horizontal line at confidence level
    fig.add_hline(y=conf_percentile, line_dash='dash', line_color='#2ECC71', line_width=2,
                  annotation_text=f'{confidence_level} Level', annotation_position='right')

    # Vertical lines at baseline and scenario values
    baseline_val = get_confidence_value(baseline_stats, confidence_level) / 1e6
    scenario_val = get_confidence_value(scenario_stats, confidence_level) / 1e6

    fig.add_vline(x=baseline_val, line_dash='dot', line_color='#3498DB', line_width=2)
    fig.add_vline(x=scenario_val, line_dash='dot', line_color='#E74C3C', line_width=2)

    fig.update_layout(
        title=f'CDF Comparison: Baseline vs {scenario_name}',
        xaxis_title='Total Risk Exposure (Million CHF)',
        yaxis_title='Cumulative Probability (%)',
        height=500,
        showlegend=True,
        legend=dict(
            orientation='h',
            yanchor='bottom',
            y=1.02,
            xanchor='center',
            x=0.5
        ),
        yaxis=dict(range=[0, 100])
    )

    return fig

def create_scenario_impact_table(df, scenario_adjustments, risk_type='initial'):
    """Create a table showing which risks were modified in the scenario"""

    if risk_type == 'initial':
        impact_col = 'Initial risk_Value'
        likelihood_col = 'Initial_Likelihood'
    else:
        impact_col = 'Residual risk_Value'
        likelihood_col = 'Residual_Likelihood'

    modified_risks = []

    for risk_id, adjustments in scenario_adjustments.items():
        original_row = df[df['Risk ID'] == risk_id]
        if len(original_row) > 0:
            original_row = original_row.iloc[0]

            orig_likelihood = original_row[likelihood_col]
            orig_impact = original_row[impact_col]

            new_likelihood = adjustments.get('likelihood', orig_likelihood)
            new_impact = adjustments.get('impact', orig_impact)

            orig_ev = orig_likelihood * orig_impact
            new_ev = new_likelihood * new_impact
            ev_change = new_ev - orig_ev

            modified_risks.append({
                'Risk ID': risk_id,
                'Description': original_row['Risk Description'][:50] + '...' if len(str(original_row['Risk Description'])) > 50 else original_row['Risk Description'],
                'Original Likelihood': f'{orig_likelihood*100:.1f}%',
                'New Likelihood': f'{new_likelihood*100:.1f}%',
                'Original Impact (M CHF)': f'{orig_impact/1e6:.2f}',
                'New Impact (M CHF)': f'{new_impact/1e6:.2f}',
                'EV Change (M CHF)': f'{ev_change/1e6:+.2f}'
            })

    return pd.DataFrame(modified_risks) if modified_risks else None

# ============================================================================
# SANKEY DIAGRAM FOR RISK FLOW
# ============================================================================

def create_risk_sankey(df, initial_stats, residual_stats, confidence_level='P95'):
    """
    Create a Sankey diagram showing the flow of both threats and opportunities.

    Shows how initial risk exposure (threats + opportunities) flows through
    mitigation to result in residual exposure, with threats and opportunities
    clearly distinguished.

    Parameters:
    - df: DataFrame with risk data including Initial_EV, Residual_EV, Risk_Reduction
    - initial_stats: Statistics dict for initial risk (from Monte Carlo)
    - residual_stats: Statistics dict for residual risk (from Monte Carlo)
    - confidence_level: Selected confidence level (P50, P80, P95)
    """

    # Separate threats (positive EV) and opportunities (negative EV)
    df_threats = df[df['Initial_EV'] > 0].copy()
    df_opportunities = df[df['Initial_EV'] < 0].copy()

    # Calculate threat values
    threat_initial = df_threats['Initial_EV'].sum()
    threat_residual = df_threats['Residual_EV'].sum()
    threat_reduced = max(0, threat_initial - threat_residual)

    # Calculate opportunity values (these are negative, so we use abs for display)
    opp_initial = abs(df_opportunities['Initial_EV'].sum()) if len(df_opportunities) > 0 else 0
    opp_residual = abs(df_opportunities['Residual_EV'].sum()) if len(df_opportunities) > 0 else 0
    opp_change = opp_residual - opp_initial  # Positive if opportunity improved

    # Net exposure calculations
    net_initial = threat_initial - opp_initial  # Threats minus opportunities
    net_residual = threat_residual - opp_residual

    # Node labels - Threats flow (red tones), Opportunities flow (green tones)
    # Node indices:
    # 0: Initial Threats
    # 1: Threat Reduced
    # 2: Residual Threats
    # 3: Initial Opportunities (shown as benefit)
    # 4: Opportunity Enhanced (if improved)
    # 5: Residual Opportunities
    # 6: Net Residual Exposure

    labels = [
        f"Initial Threats\n{threat_initial/1e6:.1f}M CHF",                    # 0
        f"Threat Reduced\n{threat_reduced/1e6:.1f}M CHF",                     # 1
        f"Residual Threats\n{threat_residual/1e6:.1f}M CHF",                  # 2
        f"Initial Opportunities\n-{opp_initial/1e6:.1f}M CHF",                # 3
        f"Opp. Change\n{opp_change/1e6:+.1f}M CHF",                           # 4
        f"Residual Opportunities\n-{opp_residual/1e6:.1f}M CHF",              # 5
        f"Net Residual\n{net_residual/1e6:.1f}M CHF"                          # 6
    ]

    # Colors: Red tones for threats, Green/Blue tones for opportunities
    colors = [
        "#C0392B",  # 0: Initial Threats - Dark Red
        "#27AE60",  # 1: Threat Reduced - Green (good!)
        "#E74C3C",  # 2: Residual Threats - Red
        "#1ABC9C",  # 3: Initial Opportunities - Teal
        "#16A085",  # 4: Opportunity Change - Dark Teal
        "#2ECC71",  # 5: Residual Opportunities - Green
        "#8E44AD"   # 6: Net Residual - Purple
    ]

    links_source = []
    links_target = []
    links_value = []
    links_color = []

    # Threat flows (warm colors)
    # Flow: Initial Threats -> Threat Reduced
    if threat_reduced > 0:
        links_source.append(0)
        links_target.append(1)
        links_value.append(threat_reduced)
        links_color.append("rgba(39, 174, 96, 0.6)")  # Green - reduction is good

    # Flow: Initial Threats -> Residual Threats
    if threat_residual > 0:
        links_source.append(0)
        links_target.append(2)
        links_value.append(threat_residual)
        links_color.append("rgba(231, 76, 60, 0.6)")  # Red

    # Flow: Residual Threats -> Net Residual
    if threat_residual > 0:
        links_source.append(2)
        links_target.append(6)
        links_value.append(threat_residual)
        links_color.append("rgba(142, 68, 173, 0.5)")  # Purple

    # Opportunity flows (cool colors) - only if opportunities exist
    if len(df_opportunities) > 0:
        # Flow: Initial Opportunities -> Opportunity Change (if changed)
        if abs(opp_change) > 0:
            links_source.append(3)
            links_target.append(4)
            links_value.append(abs(opp_change))
            links_color.append("rgba(22, 160, 133, 0.6)")  # Teal

        # Flow: Initial Opportunities -> Residual Opportunities
        # The flow represents the portion that remains
        min_opp = min(opp_initial, opp_residual)
        if min_opp > 0:
            links_source.append(3)
            links_target.append(5)
            links_value.append(min_opp)
            links_color.append("rgba(46, 204, 113, 0.6)")  # Green

        # If opportunity improved (residual > initial), add flow from change to residual
        if opp_change > 0:
            links_source.append(4)
            links_target.append(5)
            links_value.append(opp_change)
            links_color.append("rgba(46, 204, 113, 0.6)")  # Green

    fig = go.Figure(data=[go.Sankey(
        node=dict(
            pad=30,
            thickness=35,
            line=dict(color="black", width=2),
            label=labels,
            color=colors,
            hovertemplate='%{label}<extra></extra>'
        ),
        link=dict(
            source=links_source,
            target=links_target,
            value=links_value,
            color=links_color,
            hovertemplate='%{source.label} → %{target.label}<br>Value: %{value:,.0f} CHF<extra></extra>'
        ),
        textfont=dict(size=14, color='black', family='Arial Black')
    )])

    # Calculate metrics for title
    threat_reduction_pct = (threat_reduced / threat_initial * 100) if threat_initial > 0 else 0
    opp_count = len(df_opportunities)
    threat_count = len(df_threats)

    fig.update_layout(
        title={
            'text': f'<b>Risk Mitigation Flow: Threats & Opportunities</b><br>'
                   f'<sup>{threat_count} Threats (Initial: {threat_initial/1e6:.1f}M, Reduced: {threat_reduction_pct:.0f}%) | '
                   f'{opp_count} Opportunities (Benefit: {opp_residual/1e6:.1f}M) | '
                   f'Net Exposure: {net_residual/1e6:.1f}M CHF</sup>',
            'font': {'size': 16, 'color': '#2C3E50'},
            'x': 0.5,
            'xanchor': 'center'
        },
        font=dict(size=14, color='#2C3E50', family='Arial'),
        height=550,
        paper_bgcolor='#FAFAFA',
        margin=dict(l=30, r=30, t=100, b=30)
    )

    return fig

def create_risk_sankey_detailed(df):
    """
    Create a more detailed Sankey diagram with risk categories.

    Shows risk flow broken down by schedule impact status.
    """

    # Filter to only include risks (positive EV), excluding opportunities
    df_risks = df[df['Initial_EV'] > 0].copy()

    # Separate by schedule impact
    df_schedule = df_risks[df_risks['Schedule_Impact'] == True]
    df_no_schedule = df_risks[df_risks['Schedule_Impact'] == False]

    # Calculate values
    total_initial = df_risks['Initial_EV'].sum()
    schedule_initial = df_schedule['Initial_EV'].sum()
    no_schedule_initial = df_no_schedule['Initial_EV'].sum()

    schedule_residual = df_schedule['Residual_EV'].abs().sum()
    no_schedule_residual = df_no_schedule['Residual_EV'].abs().sum()

    schedule_reduced = max(0, schedule_initial - schedule_residual)
    no_schedule_reduced = max(0, no_schedule_initial - no_schedule_residual)

    total_residual = df_risks['Residual_EV'].abs().sum()
    total_reduced = schedule_reduced + no_schedule_reduced

    labels = [
        f"Total Initial Risk\n{total_initial/1e6:.1f}M CHF",        # 0
        f"Schedule Impact\n{schedule_initial/1e6:.1f}M CHF",         # 1
        f"No Schedule Impact\n{no_schedule_initial/1e6:.1f}M CHF",   # 2
        f"Reduced (Schedule)\n{schedule_reduced/1e6:.1f}M CHF",      # 3
        f"Reduced (No Schedule)\n{no_schedule_reduced/1e6:.1f}M CHF",# 4
        f"Residual (Schedule)\n{schedule_residual/1e6:.1f}M CHF",    # 5
        f"Residual (No Schedule)\n{no_schedule_residual/1e6:.1f}M CHF", # 6
        f"Total Risk Reduced\n{total_reduced/1e6:.1f}M CHF",         # 7
        f"Total Residual\n{total_residual/1e6:.1f}M CHF"             # 8
    ]

    colors = [
        "#C0392B",  # Total Initial - Dark Red
        "#922B21",  # Schedule Impact - Darker Red
        "#D35400",  # No Schedule Impact - Dark Orange
        "#1E8449",  # Reduced (Schedule) - Dark Green
        "#28B463",  # Reduced (No Schedule) - Green
        "#2471A3",  # Residual (Schedule) - Dark Blue
        "#5DADE2",  # Residual (No Schedule) - Light Blue
        "#117A65",  # Total Reduced - Teal
        "#B9770E"   # Total Residual - Dark Gold
    ]

    # Build links dynamically (skip zero values)
    links_source = []
    links_target = []
    links_value = []
    links_color = []

    link_definitions = [
        (0, 1, schedule_initial, "rgba(146, 43, 33, 0.5)"),
        (0, 2, no_schedule_initial, "rgba(211, 84, 0, 0.5)"),
        (1, 3, schedule_reduced, "rgba(30, 132, 73, 0.5)"),
        (1, 5, schedule_residual, "rgba(36, 113, 163, 0.5)"),
        (2, 4, no_schedule_reduced, "rgba(40, 180, 99, 0.5)"),
        (2, 6, no_schedule_residual, "rgba(93, 173, 226, 0.5)"),
        (3, 7, schedule_reduced, "rgba(17, 122, 101, 0.5)"),
        (4, 7, no_schedule_reduced, "rgba(17, 122, 101, 0.5)"),
        (5, 8, schedule_residual, "rgba(185, 119, 14, 0.5)"),
        (6, 8, no_schedule_residual, "rgba(185, 119, 14, 0.5)")
    ]

    for source, target, value, color in link_definitions:
        if value > 0:
            links_source.append(source)
            links_target.append(target)
            links_value.append(value)
            links_color.append(color)

    fig = go.Figure(data=[go.Sankey(
        node=dict(
            pad=25,
            thickness=30,
            line=dict(color="black", width=2),
            label=labels,
            color=colors
        ),
        link=dict(
            source=links_source,
            target=links_target,
            value=links_value,
            color=links_color
        ),
        textfont=dict(size=14, color='black', family='Arial Black')
    )])

    # Count opportunities excluded
    opportunities_count = len(df[df['Initial_EV'] <= 0])
    opportunities_note = f" | Opportunities excluded: {opportunities_count}" if opportunities_count > 0 else ""

    fig.update_layout(
        title={
            'text': f'<b>Risk Flow by Schedule Impact</b><br><sup>Breakdown by schedule impact category{opportunities_note}</sup>',
            'font': {'size': 18, 'color': '#2C3E50'},
            'x': 0.5,
            'xanchor': 'center'
        },
        font=dict(size=14, color='#2C3E50', family='Arial'),
        height=600,
        paper_bgcolor='#FAFAFA',
        margin=dict(l=20, r=20, t=80, b=20)
    )

    return fig

# ============================================================================
# DOCX Report Generation Functions
# ============================================================================
def set_cell_background(cell, fill_color):
    """Set background color for table cell"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    kwargs: top, bottom, left, right, insideH, insideV
    values: {'sz': 24, 'val': 'single', 'color': '#000000'}
    """
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)

def format_table_professional(table, has_header=True):
    """Apply professional formatting to table"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Set column widths to auto
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1  # Center vertically

    if has_header and len(table.rows) > 0:
        # Format header row
        for cell in table.rows[0].cells:
            # Dark blue background
            set_cell_background(cell, "1F4E78")
            # White text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = 'Calibri'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Format data rows with alternating colors
    for i, row in enumerate(table.rows[1:], 1):
        # Alternating row colors
        if i % 2 == 0:
            bg_color = "F2F2F2"  # Light gray
        else:
            bg_color = "FFFFFF"  # White

        for cell in row.cells:
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)

def format_table_executive(table, has_header=True, highlight_rows=None):
    """Apply executive-level professional formatting to summary tables with enhanced styling"""
    from docx.shared import Pt, RGBColor, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    highlight_rows = highlight_rows or []

    # Set table alignment to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '1F4E78')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Format cells
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1  # Center vertically

    if has_header and len(table.rows) > 0:
        # Format header row with gradient-style professional look
        for cell in table.rows[0].cells:
            set_cell_background(cell, "1F4E78")  # Dark blue
            # Add bottom border emphasis
            set_cell_border(cell, bottom={'sz': 12, 'val': 'single', 'color': '0D2E4D'})
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = 'Calibri'

    # Format data rows
    for i, row in enumerate(table.rows[1:], 1):
        is_highlight = i in highlight_rows

        if is_highlight:
            bg_color = "D4EDFC"  # Light blue highlight
        elif i % 2 == 0:
            bg_color = "F8F9FA"  # Very light gray
        else:
            bg_color = "FFFFFF"  # White

        for j, cell in enumerate(row.cells):
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                # Right-align numeric columns (typically not the first column)
                if j > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if is_highlight:
                        run.font.bold = True

def format_table_contingency(table):
    """Apply special formatting for contingency allocation summary table"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Set table alignment
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply thick table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Thicker outer border
        border.set(qn('w:color'), '1F4E78')
        tblBorders.append(border)

    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'B4C6E7')  # Lighter internal borders
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1

    # Format header row
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_background(cell, "1F4E78")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = 'Calibri'

    # Format data rows with special styling
    row_styles = [
        ("E8F4FD", False),   # Row 1: Light blue - Residual Risk Reserve
        ("E8F4FD", False),   # Row 2: Light blue - Mitigation Investment
        ("1F4E78", True),    # Row 3: Dark blue - Total Contingency (header-style)
    ]

    for i, row in enumerate(table.rows[1:], 0):
        if i < len(row_styles):
            bg_color, is_bold = row_styles[i]
        else:
            bg_color, is_bold = ("FFFFFF", False)

        for j, cell in enumerate(row.cells):
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                if j > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    run.font.bold = is_bold
                    # White text for dark row
                    if bg_color == "1F4E78":
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(0, 0, 0)

def plotly_to_image_bytes(fig, width=1600, height=800):
    """Convert Plotly figure to PNG image bytes using kaleido"""
    try:
        import io
        # Use kaleido for server-side rendering (lightweight, no browser needed)
        img_bytes = fig.to_image(format="png", width=width, height=height, scale=2)
        return io.BytesIO(img_bytes)
    except Exception as e:
        # Fall back gracefully if kaleido is not available
        return None

# ============= MATPLOTLIB CHART FUNCTIONS FOR DOCX EXPORT =============

def create_matplotlib_risk_matrix_combined(df):
    """Create side-by-side risk matrix comparison (initial vs residual) for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, axes = plt.subplots(1, 2, figsize=(16, 7))

    # Determine common axis limits for both charts
    max_likelihood = max(df['Initial_Likelihood'].max(), df['Residual_Likelihood'].max()) * 100 * 1.1
    max_impact = max(df['Initial risk_Value'].max(), df['Residual risk_Value'].max()) / 1e6 * 1.1
    max_ev = max(
        (df['Initial_Likelihood'] * df['Initial risk_Value']).max(),
        (df['Residual_Likelihood'] * df['Residual risk_Value']).max()
    ) / 1e6

    for idx, (ax, risk_type) in enumerate(zip(axes, ['initial', 'residual'])):
        if risk_type == 'initial':
            x = df['Initial_Likelihood'] * 100
            y = df['Initial risk_Value'] / 1e6
            title = 'Initial Risk Assessment'
            color_label = 'Before Mitigation'
        else:
            x = df['Residual_Likelihood'] * 100
            y = df['Residual risk_Value'] / 1e6
            title = 'Residual Risk Assessment'
            color_label = 'After Mitigation'

        # Calculate expected values for color
        ev = x * y / 100

        # Create scatter plot with consistent color scale
        scatter = ax.scatter(x, y, c=ev, cmap='RdYlGn_r', s=120, alpha=0.7,
                            edgecolors='black', linewidth=0.5, vmin=0, vmax=max_ev)

        # Add colorbar
        cbar = plt.colorbar(scatter, ax=ax)
        cbar.set_label('Expected Value (M CHF)', fontsize=9)

        # Set consistent axis limits
        ax.set_xlim(0, max_likelihood)
        ax.set_ylim(0, max_impact)

        # Add quadrant lines using overall median
        ax.axhline(y=df['Initial risk_Value'].median() / 1e6, color='gray', linestyle='--', alpha=0.5)
        ax.axvline(x=df['Initial_Likelihood'].median() * 100, color='gray', linestyle='--', alpha=0.5)

        # Labels
        ax.set_xlabel('Likelihood (%)', fontsize=11)
        ax.set_ylabel('Impact (Million CHF)', fontsize=11)
        ax.set_title(title, fontsize=13, fontweight='bold', color='#1F4E78')

        # Add risk labels
        for _, row in df.iterrows():
            if risk_type == 'initial':
                xi, yi = row['Initial_Likelihood'] * 100, row['Initial risk_Value'] / 1e6
            else:
                xi, yi = row['Residual_Likelihood'] * 100, row['Residual risk_Value'] / 1e6
            ax.annotate(row['Risk ID'], (xi, yi), fontsize=6, alpha=0.7,
                       xytext=(2, 2), textcoords='offset points')

        ax.grid(True, alpha=0.3)

    # Add overall title
    fig.suptitle('Risk Matrix Comparison: Before vs After Mitigation', fontsize=14, fontweight='bold', color='#1F4E78', y=1.02)
    plt.tight_layout()

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_heatmap_combined(df):
    """Create side-by-side risk heatmap comparison (initial vs residual) for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.colors as mcolors
    from matplotlib.patches import Rectangle

    fig, axes = plt.subplots(1, 2, figsize=(16, 8))

    # Define bins and labels
    impact_bins = [0, 100000, 1000000, 10000000, 100000000, np.inf]
    impact_labels = ['Very Low', 'Low', 'Medium', 'High', 'Very High']

    likelihood_bins = [0, 0.1, 0.25, 0.5, 0.75, 1.0]
    likelihood_labels = ['Very Low', 'Low', 'Medium', 'High', 'Very High']

    # Risk severity color scale (green to red)
    risk_colors = ['#2ECC71', '#F1C40F', '#E67E22', '#E74C3C', '#8E44AD']

    def create_heatmap_data(df_temp, impact_col, likelihood_col):
        """Helper to create heatmap data for a given risk type"""
        df_work = df_temp.copy()
        df_work['Impact_Cat'] = pd.cut(np.abs(df_work[impact_col]), bins=impact_bins, labels=impact_labels)
        df_work['Likelihood_Cat'] = pd.cut(df_work[likelihood_col], bins=likelihood_bins, labels=likelihood_labels)

        heatmap_data = []
        for imp_cat in impact_labels:
            row_data = []
            for lik_cat in likelihood_labels:
                mask = (df_work['Impact_Cat'] == imp_cat) & (df_work['Likelihood_Cat'] == lik_cat)
                count = mask.sum()
                total_ev = (df_work.loc[mask, impact_col] * df_work.loc[mask, likelihood_col]).sum()
                risk_ids = df_work.loc[mask, 'Risk ID'].tolist()
                row_data.append({'count': count, 'ev': total_ev, 'risk_ids': risk_ids})
            heatmap_data.append(row_data)
        return heatmap_data

    # Calculate max count for consistent color scale
    initial_data = create_heatmap_data(df, 'Initial risk_Value', 'Initial_Likelihood')
    residual_data = create_heatmap_data(df, 'Residual risk_Value', 'Residual_Likelihood')

    max_count = max(
        max(cell['count'] for row in initial_data for cell in row),
        max(cell['count'] for row in residual_data for cell in row)
    )
    if max_count == 0:
        max_count = 1

    for ax_idx, (ax, heatmap_data, title) in enumerate(zip(
        axes,
        [initial_data, residual_data],
        ['Initial Risk Heatmap (Before Mitigation)', 'Residual Risk Heatmap (After Mitigation)']
    )):
        # Create base risk severity grid (background colors)
        for i, imp_cat in enumerate(impact_labels):
            for j, lik_cat in enumerate(likelihood_labels):
                # Risk severity based on position (higher impact + higher likelihood = higher risk)
                severity = (i + j) / 8  # Normalized 0-1
                bg_color = plt.cm.RdYlGn_r(severity * 0.7 + 0.15)  # Scale to avoid extremes
                rect = Rectangle((j, i), 1, 1, facecolor=bg_color, alpha=0.3, edgecolor='white', linewidth=2)
                ax.add_patch(rect)

        # Create count-based overlay
        for i, imp_cat in enumerate(impact_labels):
            for j, lik_cat in enumerate(likelihood_labels):
                cell = heatmap_data[i][j]
                count = cell['count']
                ev = cell['ev'] / 1e6  # Convert to millions

                if count > 0:
                    # Add darker overlay based on count
                    intensity = count / max_count
                    overlay = Rectangle((j, i), 1, 1, facecolor='darkred', alpha=intensity * 0.5,
                                        edgecolor='black', linewidth=1)
                    ax.add_patch(overlay)

                    # Add text
                    text_color = 'white' if intensity > 0.3 else 'black'
                    ax.text(j + 0.5, i + 0.6, f'{count}', ha='center', va='center',
                           fontsize=14, fontweight='bold', color=text_color)
                    ax.text(j + 0.5, i + 0.35, f'{ev:.1f}M', ha='center', va='center',
                           fontsize=9, color=text_color)

                    # Add risk IDs if not too many
                    if len(cell['risk_ids']) <= 3 and len(cell['risk_ids']) > 0:
                        ids_text = ','.join(map(str, cell['risk_ids']))
                        ax.text(j + 0.5, i + 0.15, f'({ids_text})', ha='center', va='center',
                               fontsize=7, color=text_color, alpha=0.8)
                else:
                    ax.text(j + 0.5, i + 0.5, '-', ha='center', va='center',
                           fontsize=12, color='gray', alpha=0.5)

        # Set axis properties
        ax.set_xlim(0, 5)
        ax.set_ylim(0, 5)
        ax.set_xticks([0.5, 1.5, 2.5, 3.5, 4.5])
        ax.set_xticklabels(likelihood_labels, fontsize=9)
        ax.set_yticks([0.5, 1.5, 2.5, 3.5, 4.5])
        ax.set_yticklabels(impact_labels, fontsize=9)
        ax.set_xlabel('Likelihood →', fontsize=11, fontweight='bold')
        ax.set_ylabel('← Impact', fontsize=11, fontweight='bold')
        ax.set_title(title, fontsize=13, fontweight='bold', color='#1F4E78', pad=15)
        ax.set_aspect('equal')

        # Add grid
        for i in range(6):
            ax.axhline(y=i, color='white', linewidth=2)
            ax.axvline(x=i, color='white', linewidth=2)

    # Add legend
    fig.text(0.5, 0.02, 'Cell shows: Risk Count / Expected Value (M CHF) / (Risk IDs)',
             ha='center', fontsize=10, style='italic', color='#666666')

    # Add overall title
    fig.suptitle('Risk Heatmap Comparison: Before vs After Mitigation',
                fontsize=14, fontweight='bold', color='#1F4E78', y=0.98)
    plt.tight_layout(rect=[0, 0.05, 1, 0.95])

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_3d_comparison(df):
    """Create side-by-side 3D risk landscape comparison for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from mpl_toolkits.mplot3d import Axes3D

    fig = plt.figure(figsize=(16, 7))

    # Common calculations
    max_impact = max(
        np.abs(df['Initial risk_Value']).max(),
        np.abs(df['Residual risk_Value']).max()
    ) / 1e6 * 1.1

    max_ev = max(
        df['Initial_EV'].max(),
        df['Residual_EV'].max()
    ) / 1e6 * 1.1

    min_ev = min(
        df['Initial_EV'].min(),
        df['Residual_EV'].min()
    ) / 1e6 * 1.1

    # Separate threats and opportunities
    df_threats = df[df['Initial_EV'] > 0]
    df_opportunities = df[df['Initial_EV'] < 0]

    for idx, (risk_type, title) in enumerate([('initial', 'Initial Risk Landscape'),
                                               ('residual', 'Residual Risk Landscape')]):
        ax = fig.add_subplot(1, 2, idx + 1, projection='3d')

        if risk_type == 'initial':
            likelihood_col = 'Initial_Likelihood'
            impact_col = 'Initial risk_Value'
            ev_col = 'Initial_EV'
        else:
            likelihood_col = 'Residual_Likelihood'
            impact_col = 'Residual risk_Value'
            ev_col = 'Residual_EV'

        # Get current data
        df_t = df[df[ev_col] > 0]  # Threats
        df_o = df[df[ev_col] < 0]  # Opportunities

        # Plot threats (red)
        if len(df_t) > 0:
            x_t = df_t[likelihood_col] * 100
            y_t = np.abs(df_t[impact_col]) / 1e6
            z_t = df_t[ev_col] / 1e6
            sizes_t = np.clip(np.abs(df_t[ev_col]) / 1e6 * 30 + 50, 50, 300)
            ax.scatter(x_t, y_t, z_t, c='red', s=sizes_t, alpha=0.7,
                      edgecolors='darkred', linewidth=0.5, label='Threats', depthshade=True)
            # Add risk ID labels for top 5 threats
            top_threats = df_t.nlargest(5, ev_col)
            for _, row in top_threats.iterrows():
                ax.text(row[likelihood_col] * 100, np.abs(row[impact_col]) / 1e6,
                       row[ev_col] / 1e6, str(row['Risk ID']), fontsize=7, color='darkred')

        # Plot opportunities (green)
        if len(df_o) > 0:
            x_o = df_o[likelihood_col] * 100
            y_o = np.abs(df_o[impact_col]) / 1e6
            z_o = df_o[ev_col] / 1e6
            sizes_o = np.clip(np.abs(df_o[ev_col]) / 1e6 * 30 + 50, 50, 300)
            ax.scatter(x_o, y_o, z_o, c='green', s=sizes_o, alpha=0.7,
                      edgecolors='darkgreen', linewidth=0.5, label='Opportunities',
                      marker='^', depthshade=True)

        # Add zero plane
        xx, yy = np.meshgrid(np.linspace(0, 100, 10), np.linspace(0, max_impact, 10))
        zz = np.zeros_like(xx)
        ax.plot_surface(xx, yy, zz, alpha=0.2, color='gray')

        # Add 50% likelihood line
        y_line = np.linspace(0, max_impact, 10)
        z_line = np.linspace(min_ev, max_ev, 10)
        ax.plot([50]*10, y_line, [0]*10, 'orange', linestyle='--', linewidth=2, alpha=0.7)

        # Labels and formatting
        ax.set_xlabel('Likelihood (%)', fontsize=9, labelpad=5)
        ax.set_ylabel('Impact (M CHF)', fontsize=9, labelpad=5)
        ax.set_zlabel('Expected Value (M CHF)', fontsize=9, labelpad=5)
        ax.set_title(title, fontsize=12, fontweight='bold', color='#1F4E78', pad=10)

        # Set consistent axis limits
        ax.set_xlim(0, 100)
        ax.set_ylim(0, max_impact)
        ax.set_zlim(min_ev, max_ev)

        # Set viewing angle
        ax.view_init(elev=20, azim=45)

        # Add legend
        ax.legend(loc='upper left', fontsize=8)

    # Overall title
    fig.suptitle('3D Risk Landscape: Before vs After Mitigation\nBubble size = Expected Value magnitude | Gray plane = Zero EV threshold',
                fontsize=13, fontweight='bold', color='#1F4E78', y=1.02)
    plt.tight_layout()

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_confidence_comparison(confidence_data, selected_confidence='P80'):
    """Create confidence level comparison charts for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, axes = plt.subplots(1, 2, figsize=(16, 6))

    # Data preparation
    levels = ['P50', 'P80', 'P90', 'P95']
    percentiles = [50, 80, 90, 95]
    contingencies = [confidence_data['total_contingency'][l] / 1e6 for l in levels]
    premiums = [confidence_data['comparison_table']['premium_vs_p50'][i] for i in range(4)]

    # Colors
    colors = []
    for level in levels:
        if level == selected_confidence:
            colors.append('#1F4E78')
        elif level == 'P50':
            colors.append('#2ECC71')
        elif level == 'P95':
            colors.append('#E74C3C')
        else:
            colors.append('#3498DB')

    # Left chart: Bar chart
    ax1 = axes[0]
    bars = ax1.barh(levels, contingencies, color=colors, edgecolor='black', linewidth=1)

    # Add value labels
    for i, (bar, cont, prem) in enumerate(zip(bars, contingencies, premiums)):
        width = bar.get_width()
        label = f'{cont:.1f}M CHF'
        if levels[i] == selected_confidence:
            label += ' ★'
        ax1.text(width + 0.5, bar.get_y() + bar.get_height()/2,
                label, va='center', fontsize=10, fontweight='bold')
        if prem > 0:
            ax1.text(width - 1, bar.get_y() + bar.get_height()/2,
                    f'+{prem:.0f}%', va='center', ha='right', fontsize=9,
                    color='white', fontweight='bold')

    ax1.set_xlabel('Total Contingency (M CHF)', fontsize=11)
    ax1.set_ylabel('Confidence Level', fontsize=11)
    ax1.set_title('Total Contingency by Confidence Level', fontsize=13, fontweight='bold', color='#1F4E78')
    ax1.set_xlim(0, max(contingencies) * 1.25)

    # Right chart: Cost-Confidence curve
    ax2 = axes[1]
    ax2.plot(percentiles, contingencies, 'o-', color='#1F4E78', linewidth=3, markersize=12)

    # Fill area under curve
    ax2.fill_between(percentiles, contingencies, contingencies[0], alpha=0.2, color='#3498DB')

    # Add incremental cost annotations
    for i in range(1, len(levels)):
        delta = contingencies[i] - contingencies[i-1]
        mid_x = (percentiles[i] + percentiles[i-1]) / 2
        mid_y = (contingencies[i] + contingencies[i-1]) / 2
        ax2.annotate(f'+{delta:.1f}M', (mid_x, mid_y),
                    fontsize=9, color='#E74C3C', fontweight='bold',
                    bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))

    # Add P50 baseline
    ax2.axhline(y=contingencies[0], color='#2ECC71', linestyle='--', linewidth=2, alpha=0.7)
    ax2.text(52, contingencies[0] - 1, f'P50 Baseline: {contingencies[0]:.1f}M',
            fontsize=9, color='#2ECC71')

    ax2.set_xlabel('Confidence Level (%)', fontsize=11)
    ax2.set_ylabel('Total Contingency (M CHF)', fontsize=11)
    ax2.set_title('Cost-Confidence Trade-off Curve', fontsize=13, fontweight='bold', color='#1F4E78')
    ax2.set_xticks(percentiles)
    ax2.set_xticklabels(['P50\n(Median)', 'P80\n(Standard)', 'P90\n(Conservative)', 'P95\n(Highly Cons.)'])
    ax2.grid(True, alpha=0.3)

    # Overall title
    fig.suptitle('Confidence Level Comparison: Cost-Confidence Trade-off Analysis',
                fontsize=14, fontweight='bold', color='#1F4E78', y=1.02)
    plt.tight_layout()

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_phase_allocation(phase_allocation_data):
    """Create time-phased contingency allocation chart for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    phase_stats = phase_allocation_data['phase_stats']
    confidence_level = phase_allocation_data['confidence_level']
    total_at_confidence = phase_allocation_data['total_at_confidence']

    # Sort phases by order
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

    phases = [stats['name'] for _, stats in sorted_phases]
    phase_codes = [code for code, _ in sorted_phases]
    colors = [stats['color'] for _, stats in sorted_phases]
    ev_values = [stats['expected_value'] / 1e6 for _, stats in sorted_phases]
    conf_values = [stats['at_confidence'] / 1e6 for _, stats in sorted_phases]
    cumulative_values = [0] + [stats['cumulative_at_confidence'] / 1e6 for _, stats in sorted_phases]

    fig, axes = plt.subplots(1, 2, figsize=(16, 6))

    # Left chart: Grouped bar chart
    ax1 = axes[0]
    x = np.arange(len(phases))
    width = 0.35

    bars1 = ax1.bar(x - width/2, ev_values, width, label='Expected Value', color='#3498DB', alpha=0.7, edgecolor='black')
    bars2 = ax1.bar(x + width/2, conf_values, width, label=f'{confidence_level} Value', color='#E74C3C', alpha=0.7, edgecolor='black')

    # Add value labels
    for bar, val in zip(bars1, ev_values):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f'{val:.1f}M', ha='center', va='bottom', fontsize=8)
    for bar, val in zip(bars2, conf_values):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f'{val:.1f}M', ha='center', va='bottom', fontsize=8)

    ax1.set_xlabel('Project Phase', fontsize=11)
    ax1.set_ylabel('Contingency Amount (M CHF)', fontsize=11)
    ax1.set_title(f'Phase Contingency Allocation ({confidence_level})', fontsize=13, fontweight='bold', color='#1F4E78')
    ax1.set_xticks(x)
    ax1.set_xticklabels(phases, rotation=45, ha='right')
    ax1.legend(loc='upper right')
    ax1.grid(True, axis='y', alpha=0.3)

    # Right chart: S-Curve / Burn-Down
    ax2 = axes[1]

    # Cumulative allocation line
    phase_labels = ['Start'] + phases
    cumulative_pct = [(cv / (total_at_confidence/1e6) * 100) if total_at_confidence > 0 else 0 for cv in cumulative_values]

    ax2.plot(range(len(phase_labels)), cumulative_values, 'o-', color='#3498DB', linewidth=3, markersize=10, label='Cumulative Allocation')

    # Fill under curve
    ax2.fill_between(range(len(phase_labels)), cumulative_values, alpha=0.2, color='#3498DB')

    # Burn-down line (remaining)
    remaining_values = [total_at_confidence/1e6 - cv for cv in cumulative_values]
    ax2.plot(range(len(phase_labels)), remaining_values, 'o--', color='#E74C3C', linewidth=2, markersize=8, label='Remaining Contingency')

    # Add percentage labels
    for i, (cv, cp) in enumerate(zip(cumulative_values, cumulative_pct)):
        ax2.annotate(f'{cp:.0f}%', (i, cv), textcoords='offset points', xytext=(0, 10),
                    ha='center', fontsize=9, fontweight='bold', color='#1F4E78')

    ax2.set_xlabel('Project Phase', fontsize=11)
    ax2.set_ylabel('Contingency Amount (M CHF)', fontsize=11)
    ax2.set_title(f'Contingency S-Curve & Burn-Down ({confidence_level})', fontsize=13, fontweight='bold', color='#1F4E78')
    ax2.set_xticks(range(len(phase_labels)))
    ax2.set_xticklabels(phase_labels, rotation=45, ha='right')
    ax2.legend(loc='upper right')
    ax2.grid(True, alpha=0.3)

    # Overall title
    fig.suptitle('Time-Phased Contingency Profile: Phase Allocation Analysis',
                fontsize=14, fontweight='bold', color='#1F4E78', y=1.02)
    plt.tight_layout()

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_phase_waterfall(phase_allocation_data):
    """Create waterfall chart for phase contributions for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    phase_stats = phase_allocation_data['phase_stats']
    confidence_level = phase_allocation_data['confidence_level']
    total_at_confidence = phase_allocation_data['total_at_confidence']

    # Sort phases by order
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

    phases = [stats['name'] for _, stats in sorted_phases]
    colors = [stats['color'] for _, stats in sorted_phases]
    values = [stats['at_confidence'] / 1e6 for _, stats in sorted_phases]
    percentages = [stats['confidence_percentage'] for _, stats in sorted_phases]

    fig, ax = plt.subplots(figsize=(12, 6))

    # Calculate waterfall positions
    cumulative = 0
    for i, (phase, val, color, pct) in enumerate(zip(phases, values, colors, percentages)):
        # Draw bar
        ax.bar(i, val, bottom=cumulative, color=color, edgecolor='black', linewidth=1)

        # Add value label inside bar
        ax.text(i, cumulative + val/2, f'{val:.1f}M\n({pct:.0f}%)',
               ha='center', va='center', fontsize=9, fontweight='bold', color='white')

        cumulative += val

    # Add total bar
    ax.bar(len(phases), total_at_confidence/1e6, color='#2C3E50', edgecolor='black', linewidth=2)
    ax.text(len(phases), total_at_confidence/1e6/2, f'Total\n{total_at_confidence/1e6:.1f}M',
           ha='center', va='center', fontsize=10, fontweight='bold', color='white')

    # Add connecting lines
    cumulative = 0
    for i in range(len(phases)):
        cumulative += values[i]
        if i < len(phases) - 1:
            ax.hlines(y=cumulative, xmin=i+0.4, xmax=i+0.6, color='gray', linestyle='-', linewidth=1)

    ax.set_xlabel('Project Phase', fontsize=11)
    ax.set_ylabel('Contingency Amount (M CHF)', fontsize=11)
    ax.set_title(f'Waterfall: Phase Contributions to Total Contingency ({confidence_level})',
                fontsize=13, fontweight='bold', color='#1F4E78')
    ax.set_xticks(range(len(phases) + 1))
    ax.set_xticklabels(phases + ['Total'], rotation=45, ha='right')
    ax.grid(True, axis='y', alpha=0.3)

    plt.tight_layout()

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_risk_matrix(df, risk_type='initial'):
    """Create risk matrix using matplotlib for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 8))

    if risk_type == 'initial':
        x = df['Initial_Likelihood'] * 100
        y = df['Initial risk_Value'] / 1e6
        title = 'Risk Matrix - Initial Risk Assessment'
    else:
        x = df['Residual_Likelihood'] * 100
        y = df['Residual risk_Value'] / 1e6
        title = 'Risk Matrix - Residual Risk Assessment'

    # Calculate expected values for color
    ev = x * y / 100

    # Create scatter plot
    scatter = ax.scatter(x, y, c=ev, cmap='RdYlGn_r', s=150, alpha=0.7, edgecolors='black', linewidth=0.5)

    # Add colorbar
    cbar = plt.colorbar(scatter, ax=ax)
    cbar.set_label('Expected Value (M CHF)', fontsize=10)

    # Add quadrant lines
    ax.axhline(y=y.median(), color='gray', linestyle='--', alpha=0.5)
    ax.axvline(x=x.median(), color='gray', linestyle='--', alpha=0.5)

    # Labels
    ax.set_xlabel('Likelihood (%)', fontsize=12)
    ax.set_ylabel('Impact (Million CHF)', fontsize=12)
    ax.set_title(title, fontsize=14, fontweight='bold', color='#1F4E78')

    # Add risk labels
    for idx, row in df.iterrows():
        if risk_type == 'initial':
            xi, yi = row['Initial_Likelihood'] * 100, row['Initial risk_Value'] / 1e6
        else:
            xi, yi = row['Residual_Likelihood'] * 100, row['Residual risk_Value'] / 1e6
        ax.annotate(row['Risk ID'], (xi, yi), fontsize=7, alpha=0.7,
                   xytext=(3, 3), textcoords='offset points')

    ax.grid(True, alpha=0.3)
    plt.tight_layout()

    # Save to bytes
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_cdf_combined(initial_results, initial_stats, residual_results, residual_stats, confidence_level='P95'):
    """Create side-by-side CDF comparison (initial vs residual) for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, axes = plt.subplots(1, 2, figsize=(16, 6))

    # Determine common x-axis limit
    max_value = max(initial_results.max(), residual_results.max()) / 1e6 * 1.05

    for idx, (ax, results, stats, title) in enumerate(zip(
        axes,
        [initial_results, residual_results],
        [initial_stats, residual_stats],
        ['Initial Risk Exposure', 'Residual Risk Exposure']
    )):
        # Sort results for CDF
        sorted_results = np.sort(results)
        cumulative = np.arange(1, len(sorted_results) + 1) / len(sorted_results)

        # Plot CDF
        ax.plot(sorted_results / 1e6, cumulative * 100, color='#1F4E78', linewidth=2)
        ax.fill_between(sorted_results / 1e6, cumulative * 100, alpha=0.3, color='#1F4E78')

        # Add percentile lines
        percentiles = {'P50': (stats['p50'], '#2ecc71'), 'P80': (stats['p80'], '#f39c12'), 'P95': (stats['p95'], '#e74c3c')}
        for label, (value, color) in percentiles.items():
            ax.axvline(x=value / 1e6, color=color, linestyle='--', linewidth=2, alpha=0.8)
            pct = int(label[1:])
            ax.axhline(y=pct, color=color, linestyle=':', linewidth=1, alpha=0.5)

            # Highlight selected confidence level
            if label == confidence_level:
                ax.annotate(f'{label}: {value/1e6:.2f}M CHF',
                           xy=(value / 1e6, pct), xytext=(10, 10),
                           textcoords='offset points', fontsize=9, fontweight='bold',
                           color=color, bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor=color))
            else:
                ax.annotate(f'{label}: {value/1e6:.2f}M',
                           xy=(value / 1e6, pct), xytext=(5, 5),
                           textcoords='offset points', fontsize=8, color=color)

        ax.set_title(title, fontsize=13, fontweight='bold', color='#1F4E78')
        ax.set_xlabel('Total Risk Exposure (Million CHF)', fontsize=11)
        ax.set_ylabel('Cumulative Probability (%)', fontsize=11)
        ax.grid(True, alpha=0.3)
        ax.set_ylim(0, 100)
        ax.set_xlim(0, max_value)

    # Add overall title
    fig.suptitle('CDF Comparison: Before vs After Mitigation', fontsize=14, fontweight='bold', color='#1F4E78', y=1.02)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_cdf(results, stats, risk_type='initial', confidence_level='P95'):
    """Create CDF plot using matplotlib for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 6))

    # Sort results for CDF
    sorted_results = np.sort(results)
    cumulative = np.arange(1, len(sorted_results) + 1) / len(sorted_results)

    # Plot CDF
    ax.plot(sorted_results / 1e6, cumulative * 100, color='#1F4E78', linewidth=2)
    ax.fill_between(sorted_results / 1e6, cumulative * 100, alpha=0.3, color='#1F4E78')

    # Add percentile lines
    percentiles = {'P50': (stats['p50'], '#2ecc71'), 'P80': (stats['p80'], '#f39c12'), 'P95': (stats['p95'], '#e74c3c')}
    for label, (value, color) in percentiles.items():
        ax.axvline(x=value / 1e6, color=color, linestyle='--', linewidth=2, alpha=0.8)
        pct = int(label[1:])
        ax.axhline(y=pct, color=color, linestyle=':', linewidth=1, alpha=0.5)

        # Highlight selected confidence level
        if label == confidence_level:
            ax.annotate(f'{label}: {value/1e6:.2f}M CHF',
                       xy=(value / 1e6, pct), xytext=(10, 10),
                       textcoords='offset points', fontsize=10, fontweight='bold',
                       color=color, bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor=color))
        else:
            ax.annotate(f'{label}: {value/1e6:.2f}M',
                       xy=(value / 1e6, pct), xytext=(5, 5),
                       textcoords='offset points', fontsize=9, color=color)

    title = 'Cumulative Distribution Function - ' + ('Initial' if risk_type == 'initial' else 'Residual') + ' Risk'
    ax.set_title(title, fontsize=14, fontweight='bold', color='#1F4E78')
    ax.set_xlabel('Total Risk Exposure (Million CHF)', fontsize=12)
    ax.set_ylabel('Cumulative Probability (%)', fontsize=12)
    ax.grid(True, alpha=0.3)
    ax.set_ylim(0, 100)

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_histogram_combined(initial_results, initial_stats, residual_results, residual_stats):
    """Create side-by-side histogram comparison (initial vs residual) for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, axes = plt.subplots(1, 2, figsize=(16, 6))

    # Determine common axis limits
    all_results = np.concatenate([initial_results, residual_results])
    min_val = all_results.min() / 1e6
    max_val = all_results.max() / 1e6
    bins = np.linspace(min_val, max_val, 51)

    for idx, (ax, results, stats, title) in enumerate(zip(
        axes,
        [initial_results, residual_results],
        [initial_stats, residual_stats],
        ['Initial Risk Distribution', 'Residual Risk Distribution']
    )):
        # Create histogram
        n, _, patches = ax.hist(results / 1e6, bins=bins, color='#1F4E78', alpha=0.7, edgecolor='white')

        # Add mean and median lines
        ax.axvline(x=stats['mean'] / 1e6, color='#e74c3c', linestyle='-', linewidth=2, label=f'Mean: {stats["mean"]/1e6:.2f}M')
        ax.axvline(x=stats['p50'] / 1e6, color='#2ecc71', linestyle='--', linewidth=2, label=f'Median: {stats["p50"]/1e6:.2f}M')
        ax.axvline(x=stats['p95'] / 1e6, color='#f39c12', linestyle=':', linewidth=2, label=f'P95: {stats["p95"]/1e6:.2f}M')

        ax.set_title(title, fontsize=13, fontweight='bold', color='#1F4E78')
        ax.set_xlabel('Total Risk Exposure (Million CHF)', fontsize=11)
        ax.set_ylabel('Frequency', fontsize=11)
        ax.legend(loc='upper right', fontsize=9)
        ax.grid(True, alpha=0.3, axis='y')
        ax.set_xlim(min_val, max_val)

    # Add overall title
    fig.suptitle('Risk Distribution Comparison: Before vs After Mitigation', fontsize=14, fontweight='bold', color='#1F4E78', y=1.02)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_histogram(results, stats, risk_type='initial'):
    """Create histogram using matplotlib for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 6))

    # Create histogram
    n, bins, patches = ax.hist(results / 1e6, bins=50, color='#1F4E78', alpha=0.7, edgecolor='white')

    # Add mean and median lines
    ax.axvline(x=stats['mean'] / 1e6, color='#e74c3c', linestyle='-', linewidth=2, label=f'Mean: {stats["mean"]/1e6:.2f}M')
    ax.axvline(x=stats['p50'] / 1e6, color='#2ecc71', linestyle='--', linewidth=2, label=f'Median: {stats["p50"]/1e6:.2f}M')
    ax.axvline(x=stats['p95'] / 1e6, color='#f39c12', linestyle=':', linewidth=2, label=f'P95: {stats["p95"]/1e6:.2f}M')

    title = 'Risk Exposure Distribution - ' + ('Initial' if risk_type == 'initial' else 'Residual') + ' Risk'
    ax.set_title(title, fontsize=14, fontweight='bold', color='#1F4E78')
    ax.set_xlabel('Total Risk Exposure (Million CHF)', fontsize=12)
    ax.set_ylabel('Frequency', fontsize=12)
    ax.legend(loc='upper right')
    ax.grid(True, alpha=0.3, axis='y')

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_pareto(sensitivity_df, top_n=20):
    """Create Pareto chart using matplotlib for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, ax1 = plt.subplots(figsize=(12, 8))

    # Get top N risks
    df_plot = sensitivity_df.head(top_n).copy()

    # Bar chart for variance contribution
    bars = ax1.barh(range(len(df_plot)), df_plot['Variance %'], color='#1F4E78', alpha=0.8)
    ax1.set_yticks(range(len(df_plot)))
    ax1.set_yticklabels([f"{row['Risk ID']}: {row['Risk Description'][:30]}..."
                         if len(row['Risk Description']) > 30 else f"{row['Risk ID']}: {row['Risk Description']}"
                         for _, row in df_plot.iterrows()], fontsize=9)
    ax1.invert_yaxis()
    ax1.set_xlabel('Variance Contribution (%)', fontsize=12, color='#1F4E78')
    ax1.tick_params(axis='x', labelcolor='#1F4E78')

    # Cumulative line on secondary axis
    ax2 = ax1.twiny()
    ax2.plot(df_plot['Cumulative %'], range(len(df_plot)), color='#e74c3c', linewidth=2, marker='o', markersize=5)
    ax2.set_xlabel('Cumulative %', fontsize=12, color='#e74c3c')
    ax2.tick_params(axis='x', labelcolor='#e74c3c')

    # Add 80% threshold line
    ax2.axvline(x=80, color='#e74c3c', linestyle='--', alpha=0.5)
    ax2.annotate('80% threshold', xy=(80, len(df_plot)-1), fontsize=9, color='#e74c3c')

    ax1.set_title('Pareto Analysis - Risk Variance Contribution', fontsize=14, fontweight='bold', color='#1F4E78', pad=20)
    ax1.grid(True, alpha=0.3, axis='x')

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def create_matplotlib_roi_chart(df_with_roi, top_n=20):
    """Create ROI bar chart using matplotlib for DOCX export"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    df_with_measures = df_with_roi[df_with_roi['Cost of Measures_Value'] > 0].copy()
    if len(df_with_measures) == 0:
        return None

    df_plot = df_with_measures.nlargest(top_n, 'ROI')

    fig, ax = plt.subplots(figsize=(12, 8))

    # Color based on ROI value (green for positive, red for negative)
    colors = ['#27ae60' if roi > 0 else '#e74c3c' for roi in df_plot['ROI']]

    bars = ax.barh(range(len(df_plot)), df_plot['ROI'], color=colors, alpha=0.8)
    ax.set_yticks(range(len(df_plot)))
    ax.set_yticklabels([f"{row['Risk ID']}: {row['Risk Description'][:35]}..."
                        if len(row['Risk Description']) > 35 else f"{row['Risk ID']}: {row['Risk Description']}"
                        for _, row in df_plot.iterrows()], fontsize=9)
    ax.invert_yaxis()

    # Add value labels on bars
    for i, (idx, row) in enumerate(df_plot.iterrows()):
        ax.annotate(f'{row["ROI"]:.0f}%', xy=(row['ROI'], i), va='center',
                   fontsize=9, fontweight='bold',
                   xytext=(5 if row['ROI'] >= 0 else -5, 0), textcoords='offset points',
                   ha='left' if row['ROI'] >= 0 else 'right')

    ax.axvline(x=0, color='black', linewidth=0.5)
    ax.set_xlabel('Return on Investment (%)', fontsize=12)
    ax.set_title('Top Mitigation Measures by ROI', fontsize=14, fontweight='bold', color='#1F4E78')
    ax.grid(True, alpha=0.3, axis='x')

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def add_docx_cover_page(doc, confidence_level):
    """Add professional cover page to DOCX report"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add spacing at top
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Risk Assessment Report')
    title_run.font.size = Pt(44)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 120)  # Professional dark blue
    title_run.font.name = 'Calibri'

    # Spacing
    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Monte Carlo Simulation & Probabilistic Risk Analysis')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)  # Gray
    subtitle_run.font.name = 'Calibri'

    # More spacing
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Horizontal line
    line = doc.add_paragraph('_' * 80)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Report metadata in a nice box
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_run = metadata.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}\n')
    meta_run.font.size = Pt(12)
    meta_run.font.name = 'Calibri'
    meta_run.font.color.rgb = RGBColor(89, 89, 89)

    conf_run = metadata.add_run(f'Confidence Level: {confidence_level}')
    conf_run.font.size = Pt(14)
    conf_run.font.bold = True
    conf_run.font.name = 'Calibri'
    conf_run.font.color.rgb = RGBColor(31, 78, 120)

    # Page break
    doc.add_page_break()

def add_docx_executive_summary(doc, initial_stats, residual_stats, df, confidence_level):
    """Add executive summary section to DOCX report"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Section heading
    heading = doc.add_heading('Executive Summary', 1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Get selected confidence values (probabilistic)
    initial_selected = get_confidence_value(initial_stats, confidence_level)
    residual_selected = get_confidence_value(residual_stats, confidence_level)
    risk_reduction_selected = initial_selected - residual_selected
    risk_reduction_pct = (risk_reduction_selected / initial_selected * 100) if initial_selected > 0 else 0

    # Calculate deterministic values (sum of expected values)
    initial_ev_total = df['Initial_EV'].sum()
    residual_ev_total = df['Residual_EV'].sum()
    deterministic_reduction = initial_ev_total - residual_ev_total
    deterministic_reduction_pct = (deterministic_reduction / initial_ev_total * 100) if initial_ev_total > 0 else 0
    total_mitigation_cost = df['Cost of Measures_Value'].sum()

    # Key findings paragraph
    intro = doc.add_paragraph()
    intro_bold = intro.add_run('Overview: ')
    intro_bold.bold = True
    intro_bold.font.name = 'Calibri'
    intro_bold.font.size = Pt(11)

    intro_text = intro.add_run(
        f'This risk assessment analyzes {len(df)} identified risks using Monte Carlo simulation '
        f'with {confidence_level} confidence level. The analysis provides both deterministic (expected value) '
        f'and probabilistic estimates of total risk exposure, evaluating the effectiveness of planned mitigation measures.'
    )
    intro_text.font.name = 'Calibri'
    intro_text.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # ========== THREAT VS OPPORTUNITY ANALYSIS SECTION ==========
    # Calculate threat/opportunity metrics
    to_metrics = calculate_threat_opportunity_metrics(df)

    to_heading = doc.add_heading('Threat vs Opportunity Analysis (Expected Values)', 2)
    for run in to_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    to_intro = doc.add_paragraph()
    to_intro_run = to_intro.add_run(
        'In quantitative risk analysis, threats and opportunities are calculated separately. '
        'Threats represent potential costs (positive EV), while opportunities represent potential benefits (negative EV). '
        'Net Exposure = Threat EV + Opportunity EV.'
    )
    to_intro_run.font.name = 'Calibri'
    to_intro_run.font.size = Pt(10)
    to_intro_run.font.italic = True
    to_intro_run.font.color.rgb = RGBColor(89, 89, 89)

    # Create table for threat/opportunity metrics
    to_table = doc.add_table(rows=7, cols=3)

    to_table.rows[0].cells[0].text = 'Category'
    to_table.rows[0].cells[1].text = 'Initial (M CHF)'
    to_table.rows[0].cells[2].text = 'Residual (M CHF)'

    to_table_data = [
        (f'Threats ({to_metrics["threat_count"]} risks)',
         f'{to_metrics["threat_initial_ev"]/1e6:.2f}',
         f'{to_metrics["threat_residual_ev"]/1e6:.2f}'),
        (f'Opportunities ({to_metrics["opportunity_count"]} risks)',
         f'{to_metrics["opportunity_initial_ev"]/1e6:.2f}',
         f'{to_metrics["opportunity_residual_ev"]/1e6:.2f}'),
        ('Net Exposure',
         f'{to_metrics["net_initial_exposure"]/1e6:.2f}',
         f'{to_metrics["net_residual_exposure"]/1e6:.2f}'),
        ('Threat Reduction',
         f'{to_metrics["threat_reduction"]/1e6:.2f}',
         f'{(to_metrics["threat_reduction"]/to_metrics["threat_initial_ev"]*100) if to_metrics["threat_initial_ev"] > 0 else 0:.1f}%'),
        ('Opportunity Change',
         f'{to_metrics["opportunity_change"]/1e6:.2f}',
         'Enhanced' if to_metrics["opportunity_change"] < 0 else 'Unchanged'),
        ('Net Reduction',
         f'{to_metrics["net_reduction"]/1e6:.2f}',
         f'{(to_metrics["net_reduction"]/to_metrics["net_initial_exposure"]*100) if to_metrics["net_initial_exposure"] > 0 else 0:.1f}%'),
    ]

    for i, (cat, initial, residual) in enumerate(to_table_data, 1):
        to_table.rows[i].cells[0].text = cat
        to_table.rows[i].cells[1].text = initial
        to_table.rows[i].cells[2].text = residual

    format_table_executive(to_table, has_header=True, highlight_rows=[3, 6])

    doc.add_paragraph()  # Spacing

    # ========== DETERMINISTIC METRICS SECTION ==========
    det_heading = doc.add_heading('Overall Risk Metrics (Expected Values)', 2)
    for run in det_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    det_intro = doc.add_paragraph()
    det_intro_run = det_intro.add_run(
        'The deterministic analysis calculates risk exposure as the sum of individual expected values '
        '(Impact × Likelihood) for each risk. This provides a baseline estimate assuming average outcomes.'
    )
    det_intro_run.font.name = 'Calibri'
    det_intro_run.font.size = Pt(10)
    det_intro_run.font.italic = True
    det_intro_run.font.color.rgb = RGBColor(89, 89, 89)

    # Create table for deterministic metrics
    det_table = doc.add_table(rows=5, cols=2)

    det_table.rows[0].cells[0].text = 'Deterministic Metric'
    det_table.rows[0].cells[1].text = 'Value (Million CHF)'

    det_metrics = [
        ('Initial Risk Exposure (Σ EV)', f'{initial_ev_total/1e6:.2f}'),
        ('Residual Risk Exposure (Σ EV)', f'{residual_ev_total/1e6:.2f}'),
        ('Risk Reduction (Σ EV)', f'{deterministic_reduction/1e6:.2f} ({deterministic_reduction_pct:.1f}%)'),
        ('Total Mitigation Investment', f'{total_mitigation_cost/1e6:.2f}'),
    ]

    for i, (metric, value) in enumerate(det_metrics, 1):
        det_table.rows[i].cells[0].text = metric
        det_table.rows[i].cells[1].text = value

    format_table_executive(det_table, has_header=True, highlight_rows=[3])

    doc.add_paragraph()  # Spacing

    # ========== PROBABILISTIC METRICS SECTION ==========
    prob_heading = doc.add_heading(f'Probabilistic Risk Metrics ({confidence_level} Confidence)', 2)
    for run in prob_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    prob_intro = doc.add_paragraph()
    prob_intro_run = prob_intro.add_run(
        f'The probabilistic analysis uses Monte Carlo simulation to model portfolio-level risk exposure. '
        f'The {confidence_level} percentile indicates {confidence_level[1:]}% probability that actual exposure '
        f'will not exceed this value, accounting for risk correlations and uncertainty.'
    )
    prob_intro_run.font.name = 'Calibri'
    prob_intro_run.font.size = Pt(10)
    prob_intro_run.font.italic = True
    prob_intro_run.font.color.rgb = RGBColor(89, 89, 89)

    # Create table for probabilistic metrics
    prob_table = doc.add_table(rows=6, cols=2)

    prob_table.rows[0].cells[0].text = f'Probabilistic Metric ({confidence_level})'
    prob_table.rows[0].cells[1].text = 'Value (Million CHF)'

    prob_metrics = [
        (f'Initial Risk Exposure ({confidence_level})', f'{initial_selected/1e6:.2f}'),
        (f'Residual Risk Exposure ({confidence_level})', f'{residual_selected/1e6:.2f}'),
        (f'Portfolio Risk Reduction', f'{risk_reduction_selected/1e6:.2f} ({risk_reduction_pct:.1f}%)'),
        ('Total Mitigation Investment', f'{total_mitigation_cost/1e6:.2f}'),
        (f'Net Benefit ({confidence_level})', f'{(risk_reduction_selected - total_mitigation_cost)/1e6:.2f}'),
    ]

    for i, (metric, value) in enumerate(prob_metrics, 1):
        prob_table.rows[i].cells[0].text = metric
        prob_table.rows[i].cells[1].text = value

    format_table_executive(prob_table, has_header=True, highlight_rows=[3, 5])

    doc.add_paragraph()  # Spacing

    # Top risks section
    top_risks_heading = doc.add_heading('Critical Risks', 2)
    for run in top_risks_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    top_risks = df.nlargest(5, 'Initial_EV')
    top_risks_para = doc.add_paragraph()
    top_risks_para.add_run(
        f'The top 5 risks by expected value account for '
        f'{top_risks["Initial_EV"].sum() / df["Initial_EV"].sum() * 100:.1f}% '
        f'of total portfolio exposure:\n'
    )

    for idx, (_, risk) in enumerate(top_risks.iterrows(), 1):
        risk_para = doc.add_paragraph(style='List Number')
        risk_para.add_run(f'{risk["Risk ID"]}: ').bold = True
        risk_para.add_run(
            f'{risk["Risk Description"]} '
            f'(Expected Value: {risk["Initial_EV"]/1e6:.2f}M CHF)'
        )

    doc.add_paragraph()  # Spacing

    # ========== MANAGEMENT RECOMMENDATIONS SECTION ==========
    # Generate narrative for recommendations
    narrative = generate_risk_narrative(df, initial_stats, residual_stats, confidence_level)

    rec_heading = doc.add_heading('Management Recommendations', 2)
    for run in rec_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Add recommendations (clean markdown formatting for Word)
    recommendations_text = narrative['recommendations']
    # Split by double newline to get individual recommendations
    rec_items = [r.strip() for r in recommendations_text.split('\n\n') if r.strip()]

    for rec in rec_items:
        # Clean markdown bold markers
        clean_rec = rec.replace('**', '')
        rec_para = doc.add_paragraph(style='List Bullet')
        # Split on first colon to get title and content
        if ':' in clean_rec:
            parts = clean_rec.split(':', 1)
            rec_para.add_run(parts[0] + ': ').bold = True
            if len(parts) > 1:
                rec_para.add_run(parts[1].strip())
        else:
            rec_para.add_run(clean_rec)

    doc.add_page_break()

def add_docx_contingency_section(doc, initial_stats, residual_stats, df, confidence_level):
    """Add contingency allocation section with professional summary table and narrative"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Section heading
    heading = doc.add_heading('Project Risk Contingency Allocation', 1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Get values
    residual_selected = get_confidence_value(residual_stats, confidence_level)
    total_mitigation_cost = df['Cost of Measures_Value'].sum()
    total_contingency = residual_selected + total_mitigation_cost

    # Introductory narrative
    intro = doc.add_paragraph()
    intro_bold = intro.add_run('Contingency Planning Framework: ')
    intro_bold.bold = True
    intro_bold.font.name = 'Calibri'
    intro_bold.font.size = Pt(11)

    intro_text = intro.add_run(
        f'Based on the {confidence_level} confidence level analysis, the following contingency allocation '
        f'is recommended to adequately cover identified project risks. This allocation should be reviewed '
        f'and approved by the Board of Directors before project commitment.'
    )
    intro_text.font.name = 'Calibri'
    intro_text.font.size = Pt(11)

    doc.add_paragraph()

    # ========== CONTINGENCY ALLOCATION TABLE ==========
    table_heading = doc.add_heading('Recommended Contingency Allocation', 2)
    for run in table_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Create the professional contingency table
    table = doc.add_table(rows=4, cols=3)

    # Headers
    table.rows[0].cells[0].text = 'Component'
    table.rows[0].cells[1].text = 'Amount (M CHF)'
    table.rows[0].cells[2].text = 'Description'

    # Data rows
    table_data = [
        ('Residual Risk Reserve', f'{residual_selected/1e6:.2f}',
         f'{confidence_level} probabilistic exposure after mitigation'),
        ('Mitigation Investment', f'{total_mitigation_cost/1e6:.2f}',
         'Total cost of planned risk mitigation measures'),
        ('TOTAL PROJECT CONTINGENCY', f'{total_contingency/1e6:.2f}',
         'Recommended allocation for BoD approval'),
    ]

    for i, (component, amount, description) in enumerate(table_data, 1):
        table.rows[i].cells[0].text = component
        table.rows[i].cells[1].text = amount
        table.rows[i].cells[2].text = description

    # Apply special contingency formatting
    format_table_contingency(table)

    doc.add_paragraph()

    # ========== PERCENTILE SELECTION NARRATIVE ==========
    narrative_heading = doc.add_heading('Understanding the Selected Confidence Level', 2)
    for run in narrative_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Confidence level explanation
    if confidence_level == 'P50':
        conf_explanation = (
            'The P50 (median) confidence level represents a 50% probability that actual risk exposure '
            'will not exceed this value. This is considered an optimistic estimate suitable for projects '
            'with high risk tolerance, strong risk management capabilities, or where cost constraints '
            'require accepting higher uncertainty. Organizations selecting P50 should maintain robust '
            'contingency access mechanisms and be prepared for potential cost overruns.'
        )
        profile_assessment = 'Optimistic / High Risk Tolerance'
        profile_color = RGBColor(230, 126, 34)  # Orange
    elif confidence_level == 'P80':
        conf_explanation = (
            'The P80 confidence level represents an 80% probability that actual risk exposure '
            'will not exceed this value. This is considered a balanced, moderately conservative estimate '
            'suitable for most commercial and infrastructure projects. P80 provides reasonable protection '
            'against adverse outcomes while avoiding excessive contingency that could impact project viability. '
            'This level is commonly adopted as industry best practice for major capital projects.'
        )
        profile_assessment = 'Balanced / Moderate Risk Profile'
        profile_color = RGBColor(41, 128, 185)  # Blue
    else:  # P95
        conf_explanation = (
            'The P95 confidence level represents a 95% probability that actual risk exposure '
            'will not exceed this value. This is a conservative estimate suitable for projects with '
            'low risk tolerance, regulatory constraints, fixed-price contracts, or critical infrastructure. '
            'While P95 provides high confidence in budget adequacy, it may result in higher contingency '
            'allocation that could affect project economics or competitiveness.'
        )
        profile_assessment = 'Conservative / Low Risk Tolerance'
        profile_color = RGBColor(39, 174, 96)  # Green

    conf_para = doc.add_paragraph()
    conf_para.add_run(conf_explanation).font.name = 'Calibri'

    doc.add_paragraph()

    # Risk profile box
    profile_para = doc.add_paragraph()
    profile_label = profile_para.add_run('Project Risk Profile Assessment: ')
    profile_label.bold = True
    profile_label.font.name = 'Calibri'
    profile_label.font.size = Pt(11)

    profile_value = profile_para.add_run(profile_assessment)
    profile_value.bold = True
    profile_value.font.name = 'Calibri'
    profile_value.font.size = Pt(12)
    profile_value.font.color.rgb = profile_color

    doc.add_paragraph()

    # ========== CONTINGENCY ALLOCATION GUIDANCE ==========
    guidance_heading = doc.add_heading('Contingency Allocation Guidance', 2)
    for run in guidance_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    guidance_intro = doc.add_paragraph()
    guidance_intro.add_run(
        'The recommended contingency should be allocated and managed according to the following principles:'
    ).font.name = 'Calibri'

    guidance_items = [
        ('Governance: ', 'Contingency release should require formal change control approval, with clear '
         'thresholds for project manager, steering committee, and Board authorization levels.'),
        ('Monitoring: ', 'Track contingency drawdown against risk realization. If drawdown exceeds '
         'forecast early in the project, escalate for review and potential replenishment.'),
        ('Risk-Based Release: ', 'Link contingency release to specific risk events rather than '
         'general cost growth. Maintain traceability between risks and contingency utilization.'),
        ('Periodic Review: ', 'Reassess risk exposure quarterly. As risks are retired or new risks '
         'emerge, adjust the contingency forecast and communicate changes to stakeholders.'),
        ('Residual Contingency: ', 'Unused contingency at project completion may be returned to '
         'corporate reserves or reallocated per organizational policy.')
    ]

    for title, content in guidance_items:
        item_para = doc.add_paragraph(style='List Bullet')
        item_title = item_para.add_run(title)
        item_title.bold = True
        item_title.font.name = 'Calibri'
        item_content = item_para.add_run(content)
        item_content.font.name = 'Calibri'

    doc.add_page_break()

def add_docx_risk_portfolio_section(doc, initial_stats, residual_stats, confidence_level, risk_matrix_img=None):
    """Add risk portfolio overview section with statistics and charts"""
    from docx.shared import Inches

    doc.add_heading('Risk Portfolio Overview', 1)

    # Statistical summary
    doc.add_heading('Statistical Summary', 2)

    # Create statistics table
    table = doc.add_table(rows=8, cols=3)

    # Headers
    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Initial Risk (M CHF)'
    table.rows[0].cells[2].text = 'Residual Risk (M CHF)'

    # Data rows
    metrics = ['Mean', 'Median (P50)', 'P80', 'P95', 'Std Dev', 'Min', 'Max']
    stat_keys = ['mean', 'p50', 'p80', 'p95', 'std', 'min', 'max']

    highlight_row = None
    for i, (metric, key) in enumerate(zip(metrics, stat_keys), 1):
        # Highlight selected confidence level
        if (key == 'p50' and confidence_level == 'P50') or \
           (key == 'p80' and confidence_level == 'P80') or \
           (key == 'p95' and confidence_level == 'P95'):
            table.rows[i].cells[0].text = f'★ {metric} (SELECTED)'
            highlight_row = i
        else:
            table.rows[i].cells[0].text = metric

        table.rows[i].cells[1].text = f'{initial_stats[key]/1e6:.2f}'
        table.rows[i].cells[2].text = f'{residual_stats[key]/1e6:.2f}'

    # Apply professional formatting
    format_table_executive(table, has_header=True, highlight_rows=[highlight_row] if highlight_row else [])

    # Add risk heatmap chart if provided
    if risk_matrix_img:
        doc.add_heading('Risk Heatmap Visualization', 2)
        doc.add_picture(risk_matrix_img, width=Inches(6.5))

    doc.add_page_break()

def add_docx_monte_carlo_section(doc, n_simulations, cdf_img=None, histogram_img=None):
    """Add Monte Carlo simulation results section"""
    from docx.shared import Inches

    doc.add_heading('Monte Carlo Simulation Results', 1)

    intro = doc.add_paragraph()
    intro.add_run(
        f'The risk assessment used Monte Carlo simulation with {n_simulations:,} iterations '
        f'to model the probabilistic distribution of total risk exposure. This approach accounts '
        f'for the uncertainty in both risk occurrence and impact.'
    )

    # CDF plot
    if cdf_img:
        doc.add_heading('Cumulative Distribution Function', 2)
        desc = doc.add_paragraph()
        desc.add_run(
            'The CDF shows the probability that total risk exposure will be at or below a given value. '
            'The selected confidence level is highlighted in the chart.'
        )
        doc.add_picture(cdf_img, width=Inches(6))

    # Histogram
    if histogram_img:
        doc.add_heading('Risk Distribution', 2)
        desc = doc.add_paragraph()
        desc.add_run(
            'The histogram shows the frequency distribution of simulated total risk exposures, '
            'providing insight into the most likely outcomes.'
        )
        doc.add_picture(histogram_img, width=Inches(6))

    doc.add_page_break()

def add_docx_sensitivity_section(doc, sensitivity_df, pareto_img=None):
    """Add sensitivity analysis section"""
    from docx.shared import Inches

    doc.add_heading('Sensitivity Analysis', 1)

    intro = doc.add_paragraph()
    intro.add_run(
        'Sensitivity analysis identifies which risks contribute most to overall portfolio uncertainty. '
        'This helps prioritize risk management efforts on the risks that matter most.'
    )

    # Key findings
    top_risk = sensitivity_df.iloc[0]
    risks_80 = (sensitivity_df['Cumulative %'] <= 80).sum()

    findings = doc.add_paragraph()
    findings.add_run('Key Findings:\n').bold = True
    findings.add_run(
        f'• The top risk driver (Risk {top_risk["Risk ID"]}) accounts for {top_risk["Variance %"]:.1f}% of total variance\n'
        f'• Just {risks_80} risks drive 80% of the portfolio uncertainty\n'
        f'• This follows the Pareto principle (80/20 rule) for risk concentration\n'
    )

    # Pareto chart
    if pareto_img:
        doc.add_heading('Pareto Chart', 2)
        desc = doc.add_paragraph()
        desc.add_run(
            'The Pareto chart shows the cumulative contribution of each risk to total variance, '
            'helping identify the "vital few" risks that require focused attention.'
        )
        doc.add_picture(pareto_img, width=Inches(6.5))

    # Top 10 risks table
    doc.add_heading('Top 10 Risk Drivers', 2)

    table = doc.add_table(rows=11, cols=4)

    # Headers
    headers = ['Risk ID', 'Description', 'Variance %', 'Cumulative %']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data
    for i, (_, risk) in enumerate(sensitivity_df.head(10).iterrows(), 1):
        table.rows[i].cells[0].text = str(risk['Risk ID'])
        table.rows[i].cells[1].text = risk['Risk Description'][:50] + ('...' if len(risk['Risk Description']) > 50 else '')
        table.rows[i].cells[2].text = f'{risk["Variance %"]:.2f}%'
        table.rows[i].cells[3].text = f'{risk["Cumulative %"]:.1f}%'

    # Apply professional formatting
    format_table_executive(table, has_header=True)

    doc.add_page_break()

def add_docx_mitigation_section(doc, df_with_roi, roi_chart_img=None):
    """Add mitigation cost-benefit analysis section"""
    from docx.shared import Inches, Pt, RGBColor

    doc.add_heading('Mitigation Cost-Benefit Analysis', 1)

    # Calculate overall portfolio metrics (all risks where reduction occurred)
    df_with_reduction = df_with_roi[df_with_roi['Risk_Reduction'] > 0].copy()

    # Separate into risks with and without mitigation costs
    df_with_measures = df_with_reduction[df_with_reduction['Cost of Measures_Value'] > 0]
    df_without_measures = df_with_reduction[df_with_reduction['Cost of Measures_Value'] == 0]

    if len(df_with_reduction) > 0:
        # Overall metrics (all risks with reduction)
        total_reduction_all = df_with_reduction['Risk_Reduction'].sum()

        # Metrics for risks with mitigation measures
        total_reduction_with_cost = df_with_measures['Risk_Reduction'].sum() if len(df_with_measures) > 0 else 0
        total_cost = df_with_measures['Cost of Measures_Value'].sum() if len(df_with_measures) > 0 else 0

        # Metrics for risks without mitigation cost but still reduced
        total_reduction_no_cost = df_without_measures['Risk_Reduction'].sum() if len(df_without_measures) > 0 else 0

        net_benefit = total_reduction_all - total_cost

        summary = doc.add_paragraph()
        summary_title = summary.add_run('Overall Risk Reduction Summary:\n')
        summary_title.bold = True
        summary_title.font.name = 'Calibri'
        summary_title.font.size = Pt(11)

        summary_text = summary.add_run(
            f'• Total Risk Reduction (All Risks): {total_reduction_all/1e6:.2f} Million CHF\n'
            f'  - With Mitigation Measures ({len(df_with_measures)} risks): {total_reduction_with_cost/1e6:.2f} Million CHF\n'
            f'  - Natural Reduction ({len(df_without_measures)} risks): {total_reduction_no_cost/1e6:.2f} Million CHF\n'
            f'• Total Mitigation Cost: {total_cost/1e6:.2f} Million CHF\n'
            f'• Net Benefit: {net_benefit/1e6:.2f} Million CHF\n'
        )
        summary_text.font.name = 'Calibri'
        summary_text.font.size = Pt(10)

        if total_cost > 0:
            bc_ratio_text = summary.add_run(f'• Benefit/Cost Ratio: {total_reduction_with_cost/total_cost:.2f}\n')
            bc_ratio_text.font.name = 'Calibri'
            bc_ratio_text.font.size = Pt(10)

        # Add explanation note
        doc.add_paragraph()
        note = doc.add_paragraph()
        note_run = note.add_run(
            'Note: "Total Risk Reduction" includes all risks where reduction occurred. '
            '"Natural Reduction" refers to risks that reduced without explicit mitigation costs '
            '(e.g., due to changed circumstances, passive controls, or re-assessment). '
            'The Benefit/Cost Ratio only considers risks with active mitigation measures. '
            'These values represent the sum of individual risk Expected Values (Impact × Likelihood), '
            'which differs from portfolio-level Monte Carlo percentiles in the Executive Summary.'
        )
        note_run.font.size = Pt(9)
        note_run.font.name = 'Calibri'
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(89, 89, 89)

        # ROI chart
        if roi_chart_img:
            doc.add_heading('Return on Investment', 2)
            doc.add_picture(roi_chart_img, width=Inches(6))

        # Top ROI opportunities table (only for risks with mitigation costs)
        if len(df_with_measures) > 0:
            doc.add_heading('Top 10 Mitigation Opportunities (Active Measures)', 2)

            # Add clarifying note
            roi_note = doc.add_paragraph()
            roi_note_run = roi_note.add_run('This table shows risks with active mitigation measures (Cost > 0). ')
            roi_note_run.font.size = Pt(9)
            roi_note_run.font.name = 'Calibri'
            roi_note_run.font.italic = True
            roi_note_run.font.color.rgb = RGBColor(89, 89, 89)

            top_roi = df_with_measures.nlargest(10, 'ROI')

            table = doc.add_table(rows=min(len(top_roi) + 1, 11), cols=5)

            # Headers
            headers = ['Risk ID', 'Description', 'Risk Reduction (M CHF)', 'Cost (M CHF)', 'ROI %']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header

            # Data
            for i, (_, risk) in enumerate(top_roi.iterrows(), 1):
                table.rows[i].cells[0].text = str(risk['Risk ID'])
                table.rows[i].cells[1].text = risk['Risk Description'][:40] + ('...' if len(risk['Risk Description']) > 40 else '')
                table.rows[i].cells[2].text = f'{risk["Risk_Reduction"]/1e6:.2f}'
                table.rows[i].cells[3].text = f'{risk["Cost of Measures_Value"]/1e6:.2f}'
                table.rows[i].cells[4].text = f'{risk["ROI"]:.1f}%'

            # Apply professional formatting
            format_table_executive(table, has_header=True)
    else:
        no_reduction_para = doc.add_paragraph()
        no_reduction_text = no_reduction_para.add_run('No risk reduction occurred in the portfolio. All risks maintained the same expected values.')
        no_reduction_text.font.name = 'Calibri'
        no_reduction_text.font.size = Pt(11)
        no_reduction_text.font.italic = True

    doc.add_page_break()

def add_docx_3d_visualization_section(doc, df, chart_3d_img=None):
    """Add 3D Risk Landscape visualization section to DOCX report"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Section heading
    heading = doc.add_heading('3D Risk Landscape Visualization', 1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Introduction
    intro = doc.add_paragraph()
    intro_text = intro.add_run(
        'The 3D Risk Landscape provides an immersive visualization of the risk portfolio, '
        'showing how risks are distributed across three dimensions: Likelihood (X-axis), '
        'Impact (Y-axis), and Expected Value (Z-axis). This view clearly separates threats '
        '(above the zero plane) from opportunities (below the zero plane).'
    )
    intro_text.font.name = 'Calibri'
    intro_text.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Key for reading the chart
    key_heading = doc.add_heading('Reading the 3D Landscape', 2)
    for run in key_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    key_items = [
        ('Red spheres', 'Threats - risks with positive Expected Value (potential costs)'),
        ('Green triangles', 'Opportunities - risks with negative Expected Value (potential benefits)'),
        ('Gray plane', 'Zero EV threshold - separates threats from opportunities'),
        ('Bubble size', 'Proportional to the magnitude of Expected Value'),
        ('Orange dashed line', '50% likelihood threshold - risks to the right require attention')
    ]

    for symbol, description in key_items:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(f'{symbol}: ').bold = True
        para.add_run(description)

    doc.add_paragraph()  # Spacing

    # Add chart image if provided
    if chart_3d_img:
        chart_heading = doc.add_heading('Risk Landscape Comparison', 2)
        for run in chart_heading.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(31, 78, 120)

        desc = doc.add_paragraph()
        desc_run = desc.add_run(
            'The following visualization compares the initial risk landscape (before mitigation) '
            'with the residual risk landscape (after mitigation). Notice how risks move closer '
            'to the zero plane and reduce in size as mitigation measures take effect.'
        )
        desc_run.font.name = 'Calibri'
        desc_run.font.size = Pt(10)
        desc_run.font.italic = True

        doc.add_picture(chart_3d_img, width=Inches(6.5))

    doc.add_paragraph()  # Spacing

    # 3D Insights table
    insights_heading = doc.add_heading('3D Landscape Insights', 2)
    for run in insights_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Calculate metrics for insights
    to_metrics = calculate_threat_opportunity_metrics(df)

    # Count high-risk items
    high_risk_initial = len(df[(df['Initial_EV'] > 1e6) & (df['Initial_Likelihood'] > 0.5)])
    high_risk_residual = len(df[(df['Residual_EV'] > 1e6) & (df['Residual_Likelihood'] > 0.5)])

    # Average threat height
    avg_threat_initial = to_metrics['threat_initial_ev'] / to_metrics['threat_count'] if to_metrics['threat_count'] > 0 else 0
    avg_threat_residual = to_metrics['threat_residual_ev'] / to_metrics['threat_count'] if to_metrics['threat_count'] > 0 else 0

    # Create insights table
    insights_table = doc.add_table(rows=5, cols=3)
    insights_table.rows[0].cells[0].text = 'Metric'
    insights_table.rows[0].cells[1].text = 'Initial'
    insights_table.rows[0].cells[2].text = 'Residual'

    insights_data = [
        ('Threats (above zero plane)', str(to_metrics['threat_count']), str(len(df[df['Residual_EV'] > 0]))),
        ('Opportunities (below zero plane)', str(to_metrics['opportunity_count']), str(len(df[df['Residual_EV'] < 0]))),
        ('High-Risk Zone (>50% likelihood, >1M EV)', str(high_risk_initial), str(high_risk_residual)),
        ('Average Threat Height (EV)', f'{avg_threat_initial/1e6:.2f}M CHF', f'{avg_threat_residual/1e6:.2f}M CHF'),
    ]

    for i, (metric, initial, residual) in enumerate(insights_data, 1):
        insights_table.rows[i].cells[0].text = metric
        insights_table.rows[i].cells[1].text = initial
        insights_table.rows[i].cells[2].text = residual

    format_table_executive(insights_table, has_header=True)

    doc.add_page_break()

def add_docx_narrative_section(doc, df, initial_stats, residual_stats, confidence_level, sensitivity_df=None):
    """Add executive risk narrative section to DOCX report"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Generate the narrative
    narrative = generate_risk_narrative(df, initial_stats, residual_stats, confidence_level, sensitivity_df)

    # Section heading
    heading = doc.add_heading('Executive Risk Narrative', 1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Executive Summary subsection
    summary_heading = doc.add_heading('Portfolio Overview', 2)
    for run in summary_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    summary_para = doc.add_paragraph()
    summary_text = summary_para.add_run(narrative['executive_summary'])
    summary_text.font.name = 'Calibri'
    summary_text.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Critical Findings subsection
    findings_heading = doc.add_heading('Critical Findings', 2)
    for run in findings_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Parse and add critical findings (clean markdown)
    findings_text = narrative['critical_findings'].replace('**', '')
    findings_lines = [line.strip() for line in findings_text.split('\n\n') if line.strip()]

    for line in findings_lines:
        if ':' in line:
            parts = line.split(':', 1)
            para = doc.add_paragraph()
            para.add_run(parts[0] + ': ').bold = True
            if len(parts) > 1:
                para.add_run(parts[1].strip())
        else:
            para = doc.add_paragraph()
            para.add_run(line)

    doc.add_paragraph()  # Spacing

    # Top Risks subsection
    top_risks_heading = doc.add_heading('Top Risk Drivers', 2)
    for run in top_risks_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Get top 5 risks for detailed listing
    df_sorted = df.copy()
    df_sorted['Abs_EV'] = np.abs(df_sorted['Initial_EV'])
    top_risks = df_sorted.nlargest(5, 'Abs_EV')

    for idx, (_, risk) in enumerate(top_risks.iterrows(), 1):
        risk_type = "Threat" if risk['Initial_EV'] > 0 else "Opportunity"
        risk_para = doc.add_paragraph(style='List Number')
        risk_para.add_run(f'Risk {risk["Risk ID"]} ({risk_type}): ').bold = True
        risk_para.add_run(
            f'{risk["Risk Description"][:100]}{"..." if len(risk["Risk Description"]) > 100 else ""}\n'
        )
        detail_para = doc.add_paragraph()
        detail_para.paragraph_format.left_indent = Pt(36)
        detail_run = detail_para.add_run(
            f'Expected Value: {risk["Initial_EV"]/1e6:.2f}M CHF | '
            f'Likelihood: {risk["Initial_Likelihood"]*100:.0f}% | '
            f'Impact: {risk["Initial risk_Value"]/1e6:.2f}M CHF'
        )
        detail_run.font.size = Pt(10)
        detail_run.font.italic = True
        detail_run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()  # Spacing

    # Sensitivity Insights (if available)
    if sensitivity_df is not None and len(sensitivity_df) > 0 and narrative['sensitivity_insights']:
        sens_heading = doc.add_heading('Sensitivity Analysis Insights', 2)
        for run in sens_heading.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(31, 78, 120)

        sens_text = narrative['sensitivity_insights'].replace('**', '')
        if ':' in sens_text:
            parts = sens_text.split(':', 1)
            sens_para = doc.add_paragraph()
            sens_para.add_run(parts[0] + ': ').bold = True
            if len(parts) > 1:
                sens_para.add_run(parts[1].strip())
        else:
            sens_para = doc.add_paragraph()
            sens_para.add_run(sens_text)

        doc.add_paragraph()  # Spacing

    # Mitigation Effectiveness Summary
    mit_heading = doc.add_heading('Mitigation Effectiveness Summary', 2)
    for run in mit_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Get threat/opportunity metrics
    to_metrics = calculate_threat_opportunity_metrics(df)

    # Create mitigation effectiveness table
    mit_table = doc.add_table(rows=5, cols=4)
    mit_table.rows[0].cells[0].text = 'Metric'
    mit_table.rows[0].cells[1].text = 'Initial'
    mit_table.rows[0].cells[2].text = 'Residual'
    mit_table.rows[0].cells[3].text = 'Change'

    mit_data = [
        ('Threat Exposure',
         f'{to_metrics["threat_initial_ev"]/1e6:.1f}M',
         f'{to_metrics["threat_residual_ev"]/1e6:.1f}M',
         f'-{to_metrics["threat_reduction"]/1e6:.1f}M'),
        ('Opportunity Value',
         f'{to_metrics["opportunity_initial_ev"]/1e6:.1f}M',
         f'{to_metrics["opportunity_residual_ev"]/1e6:.1f}M',
         f'{to_metrics["opportunity_change"]/1e6:+.1f}M'),
        ('Net Exposure',
         f'{to_metrics["net_initial_exposure"]/1e6:.1f}M',
         f'{to_metrics["net_residual_exposure"]/1e6:.1f}M',
         f'-{to_metrics["net_reduction"]/1e6:.1f}M'),
        ('Risk Reduction',
         '-',
         '-',
         f'{narrative["risk_reduction_pct"]:.1f}%'),
    ]

    for i, (metric, initial, residual, change) in enumerate(mit_data, 1):
        mit_table.rows[i].cells[0].text = metric
        mit_table.rows[i].cells[1].text = initial
        mit_table.rows[i].cells[2].text = residual
        mit_table.rows[i].cells[3].text = change

    format_table_executive(mit_table, has_header=True, highlight_rows=[3, 4])

    doc.add_page_break()

def add_docx_confidence_comparison_section(doc, confidence_data, confidence_chart_img=None, selected_confidence='P80'):
    """Add confidence level comparison section to DOCX report"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Section heading
    heading = doc.add_heading('Confidence Level Comparison', 1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Introduction
    intro = doc.add_paragraph()
    intro_text = intro.add_run(
        'This section presents contingency requirements at multiple confidence levels (P50, P80, P90, P95), '
        'enabling decision-makers to explicitly evaluate the cost-confidence trade-off. This supports informed '
        'governance discussions about risk appetite and contingency adequacy.'
    )
    intro_text.font.name = 'Calibri'
    intro_text.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Confidence Level Comparison Table
    table_heading = doc.add_heading('Contingency Requirements by Confidence Level', 2)
    for run in table_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Create comparison table
    comp_table = confidence_data['comparison_table']
    levels = comp_table['levels']

    # Build column headers with star for selected level
    headers = ['Metric']
    for level in levels:
        if level == selected_confidence:
            headers.append(f'{level} ★')
        else:
            headers.append(level)

    table = doc.add_table(rows=6, cols=6)

    # Set headers
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Table data
    table_data = [
        ('Residual Exposure', [f"{v/1e6:.2f}M" for v in comp_table['residual_exposure']]),
        ('+ Mitigation Cost', [f"{v/1e6:.2f}M" for v in comp_table['mitigation_cost']]),
        ('Total Contingency', [f"{v/1e6:.2f}M" for v in comp_table['total_contingency']]),
        ('Δ from P50', ['—'] + [f"+{v/1e6:.2f}M" for v in comp_table['delta_from_p50'][1:]]),
        ('Premium vs P50', ['—'] + [f"+{v:.1f}%" for v in comp_table['premium_vs_p50'][1:]])
    ]

    for row_idx, (metric, values) in enumerate(table_data, 1):
        table.rows[row_idx].cells[0].text = metric
        for col_idx, value in enumerate(values):
            table.rows[row_idx].cells[col_idx + 1].text = value

    format_table_executive(table, has_header=True, highlight_rows=[3])

    doc.add_paragraph()  # Spacing

    # Incremental Cost Analysis Table
    inc_heading = doc.add_heading('Incremental Cost Analysis', 2)
    for run in inc_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    inc_intro = doc.add_paragraph()
    inc_intro_text = inc_intro.add_run(
        'The following table shows the additional cost required for each step up in confidence level, '
        'demonstrating the diminishing returns of higher confidence levels.'
    )
    inc_intro_text.font.name = 'Calibri'
    inc_intro_text.font.size = Pt(10)
    inc_intro_text.font.italic = True

    inc_data = confidence_data['incremental_analysis']
    inc_table = doc.add_table(rows=4, cols=4)

    # Headers
    inc_headers = ['Confidence Step', 'Additional Cost', '% Increase', 'Cost per 1% Confidence']
    for i, header in enumerate(inc_headers):
        inc_table.rows[0].cells[i].text = header

    # Data
    for row_idx, item in enumerate(inc_data, 1):
        inc_table.rows[row_idx].cells[0].text = item['step']
        inc_table.rows[row_idx].cells[1].text = f"+{item['additional_cost']/1e6:.2f}M CHF"
        inc_table.rows[row_idx].cells[2].text = f"+{item['pct_increase']:.1f}%"
        inc_table.rows[row_idx].cells[3].text = f"{item['cost_per_1pct']/1e6:.2f}M / %"

    format_table_executive(inc_table, has_header=True)

    doc.add_paragraph()  # Spacing

    # Add chart if provided
    if confidence_chart_img:
        chart_heading = doc.add_heading('Cost-Confidence Visualization', 2)
        for run in chart_heading.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(31, 78, 120)

        chart_desc = doc.add_paragraph()
        chart_desc_text = chart_desc.add_run(
            'The charts below illustrate the cost-confidence trade-off, showing how contingency requirements '
            'increase at higher confidence levels. The curve demonstrates diminishing returns: each additional '
            'percentage of confidence costs progressively more.'
        )
        chart_desc_text.font.name = 'Calibri'
        chart_desc_text.font.size = Pt(10)
        chart_desc_text.font.italic = True

        doc.add_picture(confidence_chart_img, width=Inches(6.5))

    doc.add_paragraph()  # Spacing

    # Key Insights
    insights_heading = doc.add_heading('Key Insights', 2)
    for run in insights_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 120)

    p50_cont = confidence_data['total_contingency']['P50'] / 1e6
    p95_cont = confidence_data['total_contingency']['P95'] / 1e6
    premium_p95 = confidence_data['comparison_table']['premium_vs_p50'][3]

    insights = [
        f'Moving from P50 to P95 confidence requires an additional {p95_cont - p50_cont:.1f}M CHF (+{premium_p95:.1f}% premium).',
        f'The selected confidence level ({selected_confidence}) provides a balance between cost and risk coverage.',
        'Higher confidence levels exhibit diminishing returns—each additional percentage costs progressively more.',
        'The incremental cost analysis helps quantify the "price of certainty" for governance decisions.'
    ]

    for insight in insights:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(insight)

    # Recommendation based on selected level
    doc.add_paragraph()
    rec_para = doc.add_paragraph()
    rec_para.add_run('Recommendation: ').bold = True
    if selected_confidence == 'P50':
        rec_para.add_run(
            'P50 represents the median outcome and may be suitable for organizations with high risk tolerance. '
            'Consider if additional contingency at P80 or higher is warranted for critical projects.'
        )
    elif selected_confidence == 'P80':
        rec_para.add_run(
            'P80 is a commonly used standard that provides reasonable confidence while managing costs. '
            'This level is appropriate for most project contingency allocations.'
        )
    elif selected_confidence == 'P90':
        rec_para.add_run(
            'P90 represents a conservative approach with higher certainty. '
            'This level is suitable for risk-sensitive projects or organizations with lower risk tolerance.'
        )
    else:  # P95
        rec_para.add_run(
            'P95 provides high confidence but at significant additional cost. '
            'This level is recommended for critical infrastructure or highly risk-averse organizations.'
        )

    doc.add_page_break()

def add_docx_time_phased_section(doc, phase_allocation_data, df_enhanced=None,
                                  phase_allocation_img=None, waterfall_img=None):
    """
    Add Time-Phased Contingency Profile section to DOCX report.

    Args:
        doc: Document object
        phase_allocation_data: Dictionary from calculate_phase_allocation()
        df_enhanced: Enhanced DataFrame with phase information
        phase_allocation_img: BytesIO of phase allocation chart
        waterfall_img: BytesIO of waterfall chart
    """
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    heading = doc.add_heading('Time-Phased Contingency Profile', 1)

    # Introduction
    intro = doc.add_paragraph()
    intro.add_run(
        'This section presents the distribution of contingency requirements across project phases, '
        'supporting cash flow planning and identifying when contingency reserves are most likely to be consumed.'
    )

    doc.add_paragraph()

    # Phase Statistics
    phase_stats = phase_allocation_data['phase_stats']
    confidence_level = phase_allocation_data['confidence_level']
    total_ev = phase_allocation_data['total_ev']
    total_at_confidence = phase_allocation_data['total_at_confidence']

    # Sort phases by order
    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

    # Time-Phased Allocation Table
    doc.add_heading('Phase Allocation Table', 2)

    table = doc.add_table(rows=len(sorted_phases) + 2, cols=6)  # +2 for header and total

    # Headers
    headers = ['Phase', 'Expected Value', f'{confidence_level} Allocation', '% of Total', 'Cumulative', 'Cumulative %']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_idx, (code, stats) in enumerate(sorted_phases, 1):
        table.rows[row_idx].cells[0].text = stats['name']
        table.rows[row_idx].cells[1].text = f"{stats['expected_value']/1e6:.2f}M CHF"
        table.rows[row_idx].cells[2].text = f"{stats['at_confidence']/1e6:.2f}M CHF"
        table.rows[row_idx].cells[3].text = f"{stats['confidence_percentage']:.1f}%"
        table.rows[row_idx].cells[4].text = f"{stats['cumulative_at_confidence']/1e6:.2f}M CHF"
        table.rows[row_idx].cells[5].text = f"{stats['cumulative_confidence_pct']:.1f}%"

    # Total row
    total_row = len(sorted_phases) + 1
    table.rows[total_row].cells[0].text = 'TOTAL'
    table.rows[total_row].cells[1].text = f"{total_ev/1e6:.2f}M CHF"
    table.rows[total_row].cells[2].text = f"{total_at_confidence/1e6:.2f}M CHF"
    table.rows[total_row].cells[3].text = '100.0%'
    table.rows[total_row].cells[4].text = f"{total_at_confidence/1e6:.2f}M CHF"
    table.rows[total_row].cells[5].text = '100.0%'

    # Make total row bold
    for cell in table.rows[total_row].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Apply table formatting
    format_table_executive(table, has_header=True)

    doc.add_paragraph()

    # Summary metrics
    doc.add_heading('Key Metrics', 2)

    metrics_para = doc.add_paragraph()
    metrics_para.add_run('Total Expected Value: ').bold = True
    metrics_para.add_run(f'{total_ev/1e6:.2f}M CHF\n')
    metrics_para.add_run(f'Total {confidence_level} Contingency: ').bold = True
    metrics_para.add_run(f'{total_at_confidence/1e6:.2f}M CHF\n')

    contingency_margin = (total_at_confidence - total_ev) / total_ev * 100 if total_ev > 0 else 0
    metrics_para.add_run('Contingency Margin: ').bold = True
    metrics_para.add_run(f'+{contingency_margin:.1f}%')

    doc.add_paragraph()

    # Add phase allocation chart
    if phase_allocation_img:
        doc.add_heading('Phase Allocation Visualization', 2)
        doc.add_picture(phase_allocation_img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph()
        caption.add_run('Figure: Phase-by-phase contingency allocation and S-curve showing cumulative distribution')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.size = Pt(10)
            run.font.italic = True

    doc.add_paragraph()

    # Add waterfall chart
    if waterfall_img:
        doc.add_heading('Phase Contribution Waterfall', 2)
        doc.add_picture(waterfall_img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph()
        caption.add_run('Figure: Waterfall showing each phase\'s contribution to total contingency reserve')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.size = Pt(10)
            run.font.italic = True

    doc.add_paragraph()

    # Risk distribution by phase
    if df_enhanced is not None and 'Crystallization Phase' in df_enhanced.columns:
        doc.add_heading('Risk Distribution by Phase', 2)

        phase_counts = df_enhanced['Crystallization Phase'].value_counts()
        phases_config = phase_allocation_data['phases']

        dist_table = doc.add_table(rows=len(phases_config) + 1, cols=3)
        dist_headers = ['Phase', 'Number of Risks', 'Primary Phase Risk Count']

        for i, header in enumerate(dist_headers):
            cell = dist_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_idx, (code, phase_info) in enumerate(sorted(phases_config.items(), key=lambda x: x[1]['order']), 1):
            dist_table.rows[row_idx].cells[0].text = phase_info['name']
            count = phase_counts.get(code, 0)
            dist_table.rows[row_idx].cells[1].text = str(count)
            dist_table.rows[row_idx].cells[2].text = f'{count / len(df_enhanced) * 100:.1f}%' if len(df_enhanced) > 0 else '0%'

        format_table_executive(dist_table, has_header=True)

    doc.add_paragraph()

    # Key insight
    peak_phase = max(sorted_phases, key=lambda x: x[1]['at_confidence'])

    insight_para = doc.add_paragraph()
    insight_para.add_run('Key Insight: ').bold = True
    insight_para.add_run(
        f'The {peak_phase[1]["name"]} phase has the highest contingency allocation at '
        f'{peak_phase[1]["at_confidence"]/1e6:.2f}M CHF ({peak_phase[1]["confidence_percentage"]:.1f}% of total). '
        f'Cash flow planning should ensure adequate reserves are available during this phase.'
    )

    doc.add_paragraph()

    # Cash flow planning recommendation
    doc.add_heading('Cash Flow Planning Recommendations', 2)

    rec_para = doc.add_paragraph()
    rec_para.add_run(
        'Based on the time-phased allocation analysis, the following cash flow recommendations are provided:\n'
    )

    recommendations = []
    cumulative = 0
    for code, stats in sorted_phases:
        cumulative += stats['at_confidence']
        pct = stats['confidence_percentage']
        if pct > 20:
            recommendations.append(
                f'{stats["name"]}: High allocation ({pct:.0f}%) - Ensure significant contingency reserves '
                f'({stats["at_confidence"]/1e6:.1f}M CHF) are available at phase start.'
            )
        elif pct > 10:
            recommendations.append(
                f'{stats["name"]}: Moderate allocation ({pct:.0f}%) - Plan for {stats["at_confidence"]/1e6:.1f}M CHF '
                f'contingency draw during this phase.'
            )

    if recommendations:
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')
    else:
        doc.add_paragraph(
            'Contingency is relatively evenly distributed across phases. Standard cash flow provisions should suffice.',
            style='List Bullet'
        )

    doc.add_page_break()

def add_docx_risk_register_appendix(doc, df):
    """Add full risk register as appendix"""
    from docx.shared import Pt

    doc.add_heading('Appendix A: Risk Register', 1)

    # Create table
    display_cols = ['Risk ID', 'Risk Description', 'Initial risk_Value', 'Initial_Likelihood',
                   'Residual risk_Value', 'Residual_Likelihood', 'Cost of Measures_Value']

    table = doc.add_table(rows=len(df) + 1, cols=len(display_cols))

    # Headers
    headers = ['Risk ID', 'Description', 'Initial Risk (CHF)', 'Initial Likelihood',
               'Residual Risk (CHF)', 'Residual Likelihood', 'Mitigation Cost (CHF)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data
    for i, (_, risk) in enumerate(df.iterrows(), 1):
        table.rows[i].cells[0].text = str(risk['Risk ID'])
        table.rows[i].cells[1].text = risk['Risk Description'][:60] + ('...' if len(risk['Risk Description']) > 60 else '')
        table.rows[i].cells[2].text = f'{risk["Initial risk_Value"]:,.0f}'
        table.rows[i].cells[3].text = f'{risk["Initial_Likelihood"]:.1%}'
        table.rows[i].cells[4].text = f'{risk["Residual risk_Value"]:,.0f}'
        table.rows[i].cells[5].text = f'{risk["Residual_Likelihood"]:.1%}'
        table.rows[i].cells[6].text = f'{risk["Cost of Measures_Value"]:,.0f}'

    # Apply professional formatting
    format_table_executive(table, has_header=True)

    doc.add_paragraph()  # Spacing
    note = doc.add_paragraph()
    note_run = note.add_run('Note: This table contains the complete risk register with all identified risks.')
    note_run.font.size = Pt(9)
    note_run.font.name = 'Calibri'
    note_run.font.italic = True

    doc.add_page_break()

def add_docx_methodology_appendix(doc, n_simulations, confidence_level):
    """Add methodology explanation appendix"""

    doc.add_heading('Appendix B: Methodology', 1)

    # Monte Carlo explanation
    doc.add_heading('Monte Carlo Simulation', 2)
    mc_para = doc.add_paragraph()
    mc_para.add_run(
        f'This analysis uses Monte Carlo simulation with {n_simulations:,} iterations to model '
        f'the probabilistic distribution of total risk exposure. For each iteration:\n'
    )

    steps = [
        'Each risk is evaluated independently based on its likelihood of occurrence',
        'Risks that "occur" in the simulation contribute their impact to the total',
        'The simulation aggregates all occurring risks to calculate total portfolio exposure',
        'This process repeats thousands of times to build a statistical distribution'
    ]

    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # Confidence level explanation
    doc.add_heading('Confidence Levels', 2)
    conf_para = doc.add_paragraph()
    conf_para.add_run(
        f'The selected {confidence_level} confidence level represents a percentile of the risk distribution:\n'
    )

    conf_meanings = {
        'P50': 'There is a 50% probability that actual risk exposure will be at or below this value (median/expected value)',
        'P80': 'There is an 80% probability that actual risk exposure will be at or below this value (moderately conservative)',
        'P95': 'There is a 95% probability that actual risk exposure will be at or below this value (very conservative, suitable for contingency planning)'
    }

    for level, meaning in conf_meanings.items():
        marker = ' ★ (Selected)' if level == confidence_level else ''
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(f'{level}{marker}: ').bold = True
        para.add_run(meaning)

    # Sensitivity analysis explanation
    doc.add_heading('Sensitivity Analysis', 2)
    sens_para = doc.add_paragraph()
    sens_para.add_run(
        'Sensitivity analysis uses variance decomposition to identify which risks contribute most '
        'to overall portfolio uncertainty. The methodology:\n'
    )

    sens_steps = [
        'Run baseline simulation with all risks',
        'For each risk, run simulation with that risk removed',
        'Calculate variance reduction when the risk is removed',
        'Risks with higher variance contribution are the main drivers of uncertainty'
    ]

    for step in sens_steps:
        doc.add_paragraph(step, style='List Bullet')

def generate_docx_report(initial_stats, residual_stats, df, df_with_roi, sensitivity_df,
                        initial_results, residual_results, n_simulations, confidence_level,
                        current_df=None):
    """
    Generate comprehensive DOCX report with embedded charts

    Args:
        current_df: Optional DataFrame from currently uploaded file (for phase data check)

    Returns:
        BytesIO object containing the Word document
    """
    # Use current_df for phase column checking if provided, otherwise use df
    phase_check_df = current_df if current_df is not None else df
    from docx import Document

    doc = Document()

    # Add cover page
    add_docx_cover_page(doc, confidence_level)

    # Add executive summary
    add_docx_executive_summary(doc, initial_stats, residual_stats, df, confidence_level)

    # Add contingency allocation section
    add_docx_contingency_section(doc, initial_stats, residual_stats, df, confidence_level)

    # Generate and embed charts using matplotlib (reliable, no external dependencies)
    # Side-by-side comparison charts showing initial vs residual
    with st.spinner("Generating risk heatmap comparison chart..."):
        risk_matrix_img = create_matplotlib_heatmap_combined(df)

    add_docx_risk_portfolio_section(doc, initial_stats, residual_stats, confidence_level, risk_matrix_img)

    # Add Executive Risk Narrative section
    with st.spinner("Generating executive risk narrative..."):
        add_docx_narrative_section(doc, df, initial_stats, residual_stats, confidence_level, sensitivity_df)

    # Add 3D Risk Landscape visualization section
    with st.spinner("Generating 3D risk landscape chart..."):
        chart_3d_img = create_matplotlib_3d_comparison(df)

    add_docx_3d_visualization_section(doc, df, chart_3d_img)

    # Monte Carlo section with side-by-side comparison charts
    with st.spinner("Generating Monte Carlo comparison charts..."):
        cdf_img = create_matplotlib_cdf_combined(initial_results, initial_stats, residual_results, residual_stats, confidence_level)
        hist_img = create_matplotlib_histogram_combined(initial_results, initial_stats, residual_results, residual_stats)

    add_docx_monte_carlo_section(doc, n_simulations, cdf_img, hist_img)

    # Sensitivity section with Pareto chart
    with st.spinner("Generating sensitivity analysis chart..."):
        pareto_img = create_matplotlib_pareto(sensitivity_df, top_n=20)

    add_docx_sensitivity_section(doc, sensitivity_df, pareto_img)

    # Mitigation section with ROI chart
    with st.spinner("Generating mitigation analysis chart..."):
        roi_img = create_matplotlib_roi_chart(df_with_roi, top_n=20)

    add_docx_mitigation_section(doc, df_with_roi, roi_img)

    # Confidence Level Comparison section
    with st.spinner("Generating confidence level comparison..."):
        total_mitigation_cost = df['Cost of Measures_Value'].sum()
        confidence_comparison = calculate_confidence_comparison(
            residual_results, total_mitigation_cost, confidence_level
        )
        confidence_chart_img = create_matplotlib_confidence_comparison(confidence_comparison, confidence_level)

    add_docx_confidence_comparison_section(doc, confidence_comparison, confidence_chart_img, confidence_level)

    # Time-Phased Contingency Profile section (only if uploaded data has phase columns WITH data)
    has_phase_data = (
        'Crystallization Phase' in phase_check_df.columns and
        'Phase Weight Distribution' in phase_check_df.columns and
        phase_check_df['Crystallization Phase'].notna().any() and
        (phase_check_df['Crystallization Phase'].astype(str).str.strip() != '').any() and
        phase_check_df['Phase Weight Distribution'].notna().any() and
        (phase_check_df['Phase Weight Distribution'].astype(str).str.strip() != '').any()
    )

    if has_phase_data:
        with st.spinner("Generating time-phased contingency profile..."):
            try:
                # Use phase_check_df for phase analysis
                phase_df = phase_check_df.copy()

                # Parse phase weights if not already done
                if 'Phase_Weights' not in phase_df.columns:
                    phase_df['Phase_Weights'] = phase_df['Phase Weight Distribution'].apply(parse_phase_weights)

                # We need risk occurrences - run a quick Monte Carlo for this
                # Use common random numbers for consistency
                random_nums = np.random.random((n_simulations, len(phase_df)))
                _, residual_occurrences = run_monte_carlo(
                    phase_df, n_simulations, risk_type='residual', random_numbers=random_nums
                )

                # Calculate phase allocation
                phase_allocation = calculate_phase_allocation(
                    phase_df,
                    residual_results,
                    residual_occurrences,
                    risk_type='residual',
                    confidence_level=confidence_level
                )

                # Generate charts
                phase_allocation_img = create_matplotlib_phase_allocation(phase_allocation)
                waterfall_img = create_matplotlib_phase_waterfall(phase_allocation)

                # Add section
                add_docx_time_phased_section(
                    doc, phase_allocation, phase_df,
                    phase_allocation_img, waterfall_img
                )
            except Exception as e:
                # If anything fails, just skip the section
                st.warning(f"Could not generate time-phased section: {str(e)}")

    # Appendices
    add_docx_risk_register_appendix(doc, df)
    add_docx_methodology_appendix(doc, n_simulations, confidence_level)

    # Save to BytesIO
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes

# Main application
def main():
    # Sidebar
    st.sidebar.title("⚙️ Settings")
    
    # File upload or use default
    uploaded_file = st.sidebar.file_uploader("Upload Risk Register (CSV)", type=['csv'])
    
    if uploaded_file is not None:
        df = load_risk_data(uploaded_file)
        st.sidebar.success(f"📁 Loaded: {uploaded_file.name}")
    else:
        # Use default file from same directory as app
        import os
        script_dir = os.path.dirname(os.path.abspath(__file__))
        default_file = os.path.join(script_dir, 'risk_register.csv')

        if os.path.exists(default_file):
            df = load_risk_data(default_file)
            st.sidebar.info("📁 Using risk_register.csv from application directory")
        else:
            st.sidebar.error("⚠️ Default risk register not found. Please upload your CSV file.")
            st.error("### 📁 No Risk Register Found")
            st.write("Please upload your risk register CSV file using the sidebar, or place a file named `risk_register.csv` in the same directory as the application.")
            st.stop()

    # Normalize phase column names (handle different naming conventions)
    # Support both "Crystallization Phase" and "crystallization_phase" formats
    phase_col_mapping = {
        'crystallization_phase': 'Crystallization Phase',
        'phase_weight_distribution': 'Phase Weight Distribution'
    }
    for old_name, new_name in phase_col_mapping.items():
        if old_name in df.columns and new_name not in df.columns:
            df.rename(columns={old_name: new_name}, inplace=True)

    # Check and display phase data availability
    # Columns must exist AND contain actual data (not empty)
    has_crystallization = ('Crystallization Phase' in df.columns and
                           df['Crystallization Phase'].notna().any() and
                           (df['Crystallization Phase'].astype(str).str.strip() != '').any())
    has_phase_weight = ('Phase Weight Distribution' in df.columns and
                        df['Phase Weight Distribution'].notna().any() and
                        (df['Phase Weight Distribution'].astype(str).str.strip() != '').any())

    if has_crystallization and has_phase_weight:
        st.sidebar.success("📅 Phase data detected - Time-Phased Profile enabled")
    else:
        st.sidebar.caption("📅 No phase data - Time-Phased Profile disabled")
        with st.sidebar.expander("Enable Time-Phased Profile"):
            st.write("Add these columns to your CSV with values:")
            st.code("crystallization_phase\nphase_weight_distribution")
            st.write("Example values:")
            st.code("ENG\nENG:0.5|PROC:0.3|FAB:0.2")
    
    # Simulation parameters
    st.sidebar.subheader("Monte Carlo Parameters")
    n_simulations = st.sidebar.slider("Number of Simulations", 
                                      min_value=1000, 
                                      max_value=100000, 
                                      value=10000, 
                                      step=1000)
    
    confidence_level = st.sidebar.selectbox("Confidence Level", 
                                           ["P50", "P80", "P95"], 
                                           index=2)
    
    # Run simulation button
    if st.sidebar.button("🎲 Run Simulation", type="primary"):
        with st.spinner("Running Monte Carlo simulation..."):
            # Run simulations
            initial_results, initial_occurrences = run_monte_carlo(df, n_simulations, 'initial')
            residual_results, residual_occurrences = run_monte_carlo(df, n_simulations, 'residual')
            
            # Calculate statistics
            initial_stats = calculate_statistics(initial_results)
            residual_stats = calculate_statistics(residual_results)
            
            # Calculate mitigation ROI
            df_with_roi = calculate_mitigation_roi(df)
            
            # Perform enhanced sensitivity analysis
            st.write("Performing sensitivity analysis...")
            sensitivity_df = perform_sensitivity_analysis(df, n_simulations)
            
            # Store in session state
            st.session_state['initial_results'] = initial_results
            st.session_state['residual_results'] = residual_results
            st.session_state['initial_occurrences'] = initial_occurrences
            st.session_state['residual_occurrences'] = residual_occurrences
            st.session_state['initial_stats'] = initial_stats
            st.session_state['residual_stats'] = residual_stats
            st.session_state['df'] = df
            st.session_state['df_with_roi'] = df_with_roi
            st.session_state['sensitivity_df'] = sensitivity_df
            st.session_state['confidence_level'] = confidence_level
            st.session_state['simulation_run'] = True

        st.sidebar.success("✅ Simulation completed!")

    # Store the current df (from uploaded file) before session state overwrite
    current_df = df.copy()

    # Display results if simulation has been run
    if st.session_state.get('simulation_run', False):
        initial_results = st.session_state['initial_results']
        residual_results = st.session_state['residual_results']
        initial_stats = st.session_state['initial_stats']
        residual_stats = st.session_state['residual_stats']
        df = st.session_state['df']
        df_with_roi = st.session_state['df_with_roi']
        sensitivity_df = st.session_state['sensitivity_df']
        confidence_level = st.session_state.get('confidence_level', 'P95')
        residual_occurrences = st.session_state.get('residual_occurrences')
        
        # Create tabs
        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
            "📊 Dashboard",
            "📈 Risk Matrix & Heatmap",
            "🔍 Sensitivity Analysis",
            "🎲 Monte Carlo Results",
            "💰 Cost-Benefit Analysis",
            "🔮 What-If Scenarios",
            "📋 Risk Register",
            "📥 Export"
        ])
        
        with tab1:
            st.header("Risk Portfolio Dashboard")

            # Display active confidence level
            st.info(f"📊 **Active Confidence Level: {confidence_level}** | Showing risk exposure at {confidence_level} percentile")

            # Calculate threat/opportunity metrics
            to_metrics = calculate_threat_opportunity_metrics(df_with_roi)

            # Primary metrics row - Threat EV, Opportunity EV, Net Exposure
            st.subheader("📊 Risk Exposure Summary (Expected Value)")

            primary_col1, primary_col2, primary_col3, primary_col4 = st.columns(4)

            with primary_col1:
                st.metric(
                    "Threat EV (Initial)",
                    f"{to_metrics['threat_initial_ev']/1e6:.2f}M CHF",
                    help=f"{to_metrics['threat_count']} threats identified"
                )
                st.metric(
                    "Threat EV (Residual)",
                    f"{to_metrics['threat_residual_ev']/1e6:.2f}M CHF",
                    delta=f"{-to_metrics['threat_reduction']/1e6:.2f}M" if to_metrics['threat_reduction'] > 0 else None
                )

            with primary_col2:
                # Opportunity EV is negative, so we show absolute value with a minus sign for clarity
                st.metric(
                    "Opportunity EV (Initial)",
                    f"{to_metrics['opportunity_initial_ev']/1e6:.2f}M CHF",
                    help=f"{to_metrics['opportunity_count']} opportunities identified (negative = benefit)"
                )
                st.metric(
                    "Opportunity EV (Residual)",
                    f"{to_metrics['opportunity_residual_ev']/1e6:.2f}M CHF",
                    delta=f"{-to_metrics['opportunity_change']/1e6:.2f}M" if to_metrics['opportunity_change'] != 0 else None
                )

            with primary_col3:
                # Net Exposure = Threat EV + Opportunity EV (where Opportunity is negative)
                st.metric(
                    "Net Exposure (Initial)",
                    f"{to_metrics['net_initial_exposure']/1e6:.2f}M CHF",
                    help="Threat EV + Opportunity EV"
                )
                st.metric(
                    "Net Exposure (Residual)",
                    f"{to_metrics['net_residual_exposure']/1e6:.2f}M CHF",
                    delta=f"{-to_metrics['net_reduction']/1e6:.2f}M" if to_metrics['net_reduction'] != 0 else None
                )

            with primary_col4:
                st.metric("Total Risks", len(df))
                st.metric(
                    "Threats / Opportunities",
                    f"{to_metrics['threat_count']} / {to_metrics['opportunity_count']}"
                )

            st.markdown("---")

            # Monte Carlo Results - Key metrics
            st.subheader("🎲 Monte Carlo Simulation Results")

            col1, col2, col3, col4 = st.columns(4)

            # Get confidence-specific values
            initial_confidence_value = get_confidence_value(initial_stats, confidence_level)
            residual_confidence_value = get_confidence_value(residual_stats, confidence_level)

            with col1:
                st.metric(f"Initial Risk ({confidence_level})", f"{initial_confidence_value/1e6:.2f}M CHF")
                st.metric("Initial Risk (Mean)", f"{initial_stats['mean']/1e6:.2f}M CHF")

            with col2:
                st.metric(f"Residual Risk ({confidence_level})", f"{residual_confidence_value/1e6:.2f}M CHF")
                st.metric("Residual Risk (Mean)", f"{residual_stats['mean']/1e6:.2f}M CHF")

            with col3:
                risk_reduction = ((initial_confidence_value - residual_confidence_value) / initial_confidence_value * 100) if initial_confidence_value != 0 else 0
                st.metric(f"Risk Reduction ({confidence_level})", f"{risk_reduction:.1f}%")
                total_mitigation_cost = df['Cost of Measures_Value'].sum()
                st.metric("Total Mitigation Cost", f"{total_mitigation_cost/1e6:.2f}M CHF")

            with col4:
                st.metric("Risks with Schedule Impact", df['Schedule_Impact'].sum())
                net_benefit = (initial_confidence_value - residual_confidence_value) - total_mitigation_cost
                st.metric("Net Benefit", f"{net_benefit/1e6:.2f}M CHF")

            st.markdown("---")

            # Executive Risk Narrative Section
            st.subheader("📝 Executive Risk Narrative")

            # Generate the narrative
            narrative = generate_risk_narrative(df_with_roi, initial_stats, residual_stats,
                                               confidence_level, sensitivity_df)

            # Display narrative in expandable sections
            with st.expander("📋 Executive Summary", expanded=True):
                st.markdown(narrative['executive_summary'])

            with st.expander("🔍 Critical Findings", expanded=False):
                st.markdown(narrative['critical_findings'])

            with st.expander("🏆 Top 5 Risks", expanded=False):
                st.markdown(narrative['top_risks_summary'])

            with st.expander("💡 Recommendations", expanded=False):
                st.markdown(narrative['recommendations'])

            with st.expander("📊 Mitigation Effectiveness", expanded=False):
                st.markdown(narrative['mitigation_summary'])

            if narrative['sensitivity_insights']:
                with st.expander("🎯 Sensitivity Insights", expanded=False):
                    st.markdown(narrative['sensitivity_insights'])

            st.markdown("---")

            # Tornado chart
            st.subheader("Top Risks by Expected Value")
            tornado_fig = create_tornado_chart(df_with_roi, top_n=15)
            st.plotly_chart(tornado_fig, use_container_width=True)
            
            # Statistics comparison table
            st.subheader("Statistical Summary")

            # Create metric names with indicator for selected confidence level
            metric_labels = {
                'Mean': 'Mean',
                'P50': 'Median (P50)',
                'P80': 'P80',
                'P95': 'P95',
                'Std Dev': 'Std Dev',
                'Min': 'Min',
                'Max': 'Max'
            }

            # Add visual indicator (★) to the active confidence level
            metrics_list = ['Mean', 'P50', 'P80', 'P95', 'Std Dev', 'Min', 'Max']
            display_metrics = []
            for metric in metrics_list:
                if metric == confidence_level:
                    display_metrics.append(f"★ {metric_labels.get(metric, metric)} (ACTIVE)")
                else:
                    display_metrics.append(metric_labels.get(metric, metric))

            stats_df = pd.DataFrame({
                'Metric': display_metrics,
                'Initial Risk (M CHF)': [
                    f"{initial_stats['mean']/1e6:.2f}",
                    f"{initial_stats['p50']/1e6:.2f}",
                    f"{initial_stats['p80']/1e6:.2f}",
                    f"{initial_stats['p95']/1e6:.2f}",
                    f"{initial_stats['std']/1e6:.2f}",
                    f"{initial_stats['min']/1e6:.2f}",
                    f"{initial_stats['max']/1e6:.2f}"
                ],
                'Residual Risk (M CHF)': [
                    f"{residual_stats['mean']/1e6:.2f}",
                    f"{residual_stats['p50']/1e6:.2f}",
                    f"{residual_stats['p80']/1e6:.2f}",
                    f"{residual_stats['p95']/1e6:.2f}",
                    f"{residual_stats['std']/1e6:.2f}",
                    f"{residual_stats['min']/1e6:.2f}",
                    f"{residual_stats['max']/1e6:.2f}"
                ]
            })

            st.dataframe(stats_df, use_container_width=True, hide_index=True)

            st.markdown("---")

            # Confidence Level Comparison Section
            st.subheader("📊 Confidence Level Comparison")
            st.info("""
            **Cost-Confidence Trade-off Analysis**: Compare contingency requirements at multiple confidence levels
            (P50, P80, P90, P95) to support informed governance discussions about risk appetite and contingency adequacy.
            """)

            # Calculate confidence comparison
            total_mitigation_cost = df['Cost of Measures_Value'].sum()
            confidence_comparison = calculate_confidence_comparison(
                residual_results, total_mitigation_cost, confidence_level
            )

            # Confidence Level Comparison Table
            st.markdown("#### Confidence Level Comparison Table")

            comp_table = confidence_comparison['comparison_table']
            levels_display = []
            for level in comp_table['levels']:
                if level == confidence_level:
                    levels_display.append(f"★ {level}")
                else:
                    levels_display.append(level)

            comparison_df = pd.DataFrame({
                'Metric': ['Residual Exposure', '+ Mitigation Cost', 'Total Contingency', 'Δ from P50', 'Premium vs P50'],
                levels_display[0]: [
                    f"{comp_table['residual_exposure'][0]/1e6:.2f}M",
                    f"{comp_table['mitigation_cost'][0]/1e6:.2f}M",
                    f"{comp_table['total_contingency'][0]/1e6:.2f}M",
                    "—",
                    "—"
                ],
                levels_display[1]: [
                    f"{comp_table['residual_exposure'][1]/1e6:.2f}M",
                    f"{comp_table['mitigation_cost'][1]/1e6:.2f}M",
                    f"{comp_table['total_contingency'][1]/1e6:.2f}M",
                    f"+{comp_table['delta_from_p50'][1]/1e6:.2f}M",
                    f"+{comp_table['premium_vs_p50'][1]:.1f}%"
                ],
                levels_display[2]: [
                    f"{comp_table['residual_exposure'][2]/1e6:.2f}M",
                    f"{comp_table['mitigation_cost'][2]/1e6:.2f}M",
                    f"{comp_table['total_contingency'][2]/1e6:.2f}M",
                    f"+{comp_table['delta_from_p50'][2]/1e6:.2f}M",
                    f"+{comp_table['premium_vs_p50'][2]:.1f}%"
                ],
                levels_display[3]: [
                    f"{comp_table['residual_exposure'][3]/1e6:.2f}M",
                    f"{comp_table['mitigation_cost'][3]/1e6:.2f}M",
                    f"{comp_table['total_contingency'][3]/1e6:.2f}M",
                    f"+{comp_table['delta_from_p50'][3]/1e6:.2f}M",
                    f"+{comp_table['premium_vs_p50'][3]:.1f}%"
                ],
                levels_display[4]: [
                    f"{comp_table['residual_exposure'][4]/1e6:.2f}M",
                    f"{comp_table['mitigation_cost'][4]/1e6:.2f}M",
                    f"{comp_table['total_contingency'][4]/1e6:.2f}M",
                    f"+{comp_table['delta_from_p50'][4]/1e6:.2f}M",
                    f"+{comp_table['premium_vs_p50'][4]:.1f}%"
                ]
            })
            st.dataframe(comparison_df, use_container_width=True, hide_index=True)

            # Incremental Cost Analysis Table
            st.markdown("#### Incremental Cost Analysis")

            inc_data = confidence_comparison['incremental_analysis']
            incremental_df = pd.DataFrame({
                'Confidence Step': [item['step'] for item in inc_data],
                'Additional Cost': [f"+{item['additional_cost']/1e6:.2f}M CHF" for item in inc_data],
                '% Increase': [f"+{item['pct_increase']:.1f}%" for item in inc_data],
                'Cost per 1% Confidence': [f"{item['cost_per_1pct']/1e6:.2f}M / %" for item in inc_data]
            })
            st.dataframe(incremental_df, use_container_width=True, hide_index=True)

            # Visualizations
            st.markdown("#### Visualizations")

            viz_col1, viz_col2 = st.columns(2)

            with viz_col1:
                # Confidence bar chart
                bar_chart = create_confidence_bar_chart(confidence_comparison, confidence_level)
                st.plotly_chart(bar_chart, use_container_width=True)

                # CDF with percentile markers
                cdf_chart = create_confidence_cdf_chart(residual_results, confidence_comparison, confidence_level)
                st.plotly_chart(cdf_chart, use_container_width=True)

            with viz_col2:
                # Cost-confidence curve
                curve_chart = create_cost_confidence_curve(confidence_comparison)
                st.plotly_chart(curve_chart, use_container_width=True)

                # Incremental cost chart
                inc_chart = create_incremental_cost_chart(confidence_comparison)
                st.plotly_chart(inc_chart, use_container_width=True)

            # Key insight
            p50_contingency = confidence_comparison['total_contingency']['P50'] / 1e6
            p95_contingency = confidence_comparison['total_contingency']['P95'] / 1e6
            premium = confidence_comparison['comparison_table']['premium_vs_p50'][3]

            st.success(f"""
            **💡 Key Insight**: Moving from P50 ({p50_contingency:.1f}M CHF) to P95 ({p95_contingency:.1f}M CHF)
            requires an additional {p95_contingency - p50_contingency:.1f}M CHF (+{premium:.1f}% premium).
            This represents the cost of increased confidence in covering risk exposure.
            """)

            # =================================================================
            # TIME-PHASED CONTINGENCY PROFILE SECTION
            # Only displayed if uploaded risk register contains phase columns WITH data
            # Uses current_df (from uploaded file) to check for phase columns
            # =================================================================
            has_phase_data = (
                'Crystallization Phase' in current_df.columns and
                'Phase Weight Distribution' in current_df.columns and
                current_df['Crystallization Phase'].notna().any() and
                (current_df['Crystallization Phase'].astype(str).str.strip() != '').any() and
                current_df['Phase Weight Distribution'].notna().any() and
                (current_df['Phase Weight Distribution'].astype(str).str.strip() != '').any()
            )

            if has_phase_data:
                st.markdown("---")
                st.subheader("📅 Time-Phased Contingency Profile")
                st.info("""
                **Cash Flow Planning Support**: View contingency allocation across project phases to support
                accurate cash flow forecasting and identify when contingency reserves are most likely to be consumed.
                """)

                try:
                    # Use current_df for phase analysis (has the phase columns from uploaded file)
                    phase_df = current_df.copy()

                    # Parse phase weights if not already done
                    if 'Phase_Weights' not in phase_df.columns:
                        phase_df['Phase_Weights'] = phase_df['Phase Weight Distribution'].apply(parse_phase_weights)

                    # Always generate fresh occurrences for phase analysis using phase_df
                    n_sims = len(residual_results)
                    random_nums = np.random.random((n_sims, len(phase_df)))
                    _, phase_occurrences = run_monte_carlo(
                        phase_df, n_sims, risk_type='residual', random_numbers=random_nums
                    )

                    # Calculate phase allocation using risk occurrences
                    phase_allocation = calculate_phase_allocation(
                        phase_df,
                        residual_results,
                        phase_occurrences,
                        risk_type='residual',
                        confidence_level=confidence_level
                    )

                    # Store in session state
                    st.session_state['phase_allocation'] = phase_allocation

                    # Time-Phased Allocation Table
                    st.markdown("#### Time-Phased Allocation Table")

                    phase_stats = phase_allocation['phase_stats']
                    sorted_phases = sorted(phase_stats.items(), key=lambda x: x[1]['order'])

                    phase_table_data = []
                    for code, stats in sorted_phases:
                        phase_table_data.append({
                            'Phase': stats['name'],
                            'Expected Value': f"{stats['expected_value']/1e6:.2f}M CHF",
                            f'{confidence_level} Allocation': f"{stats['at_confidence']/1e6:.2f}M CHF",
                            '% of Total': f"{stats['confidence_percentage']:.1f}%",
                            'Cumulative': f"{stats['cumulative_at_confidence']/1e6:.2f}M CHF",
                            'Cumulative %': f"{stats['cumulative_confidence_pct']:.1f}%"
                        })

                    phase_summary_df = pd.DataFrame(phase_table_data)
                    st.dataframe(phase_summary_df, use_container_width=True, hide_index=True)

                    # Summary metrics
                    total_ev = phase_allocation['total_ev']
                    total_conf = phase_allocation['total_at_confidence']

                    col_m1, col_m2, col_m3 = st.columns(3)
                    with col_m1:
                        st.metric("Total Expected Value", f"{total_ev/1e6:.2f}M CHF")
                    with col_m2:
                        st.metric(f"Total {confidence_level} Contingency", f"{total_conf/1e6:.2f}M CHF")
                    with col_m3:
                        contingency_margin = (total_conf - total_ev) / total_ev * 100 if total_ev > 0 else 0
                        st.metric("Contingency Margin", f"+{contingency_margin:.1f}%")

                    # Visualizations
                    st.markdown("#### Phase Allocation Visualizations")

                    viz_option = st.radio(
                        "Select View:",
                        ["Phase Allocation Chart", "S-Curve / Burn-Down", "Waterfall Chart", "Risk Distribution"],
                        horizontal=True,
                        key="phase_viz_option"
                    )

                    if viz_option == "Phase Allocation Chart":
                        show_cumulative = st.checkbox("Show Cumulative Values", value=False)
                        allocation_chart = create_phase_allocation_bar_chart(phase_allocation, show_cumulative)
                        st.plotly_chart(allocation_chart, use_container_width=True)

                    elif viz_option == "S-Curve / Burn-Down":
                        show_burndown = st.checkbox("Show Burn-Down (Remaining)", value=True)
                        scurve_chart = create_phase_scurve_chart(phase_allocation, show_burndown)
                        st.plotly_chart(scurve_chart, use_container_width=True)

                    elif viz_option == "Waterfall Chart":
                        waterfall_chart = create_phase_waterfall_chart(phase_allocation)
                        st.plotly_chart(waterfall_chart, use_container_width=True)

                    elif viz_option == "Risk Distribution":
                        dist_chart = create_phase_risk_distribution_chart(phase_df, phase_allocation)
                        st.plotly_chart(dist_chart, use_container_width=True)

                    # Early Warning Indicators Section
                    st.markdown("#### Early Warning Indicators")
                    st.caption("Simulated consumption tracking - set phase completion % to see warning status")

                    # Phase completion input (for demonstration)
                    ew_cols = st.columns(len(sorted_phases))

                    warning_data = []
                    for idx, (code, stats) in enumerate(sorted_phases):
                        with ew_cols[idx]:
                            phase_complete = st.slider(
                                f"{stats['name'][:4]} Complete %",
                                0, 100, 50,
                                key=f"ew_complete_{code}"
                            )
                            consumed_pct = st.slider(
                                f"{stats['name'][:4]} Consumed %",
                                0, 100, int(phase_complete * 0.9),
                                key=f"ew_consumed_{code}"
                            )

                            warning = get_early_warning_status(consumed_pct, phase_complete)
                            warning_data.append({
                                'Phase': stats['name'],
                                'Completion': f"{phase_complete}%",
                                'Consumed': f"{consumed_pct}%",
                                'Status': warning['status'],
                                'Message': warning['message']
                            })

                            # Status indicator
                            status_colors = {'Green': '🟢', 'Amber': '🟡', 'Red': '🔴'}
                            st.markdown(f"**{status_colors[warning['status']]} {warning['status']}**")

                    # Warning summary table
                    with st.expander("View Early Warning Summary Table"):
                        warning_df = pd.DataFrame(warning_data)
                        st.dataframe(warning_df, use_container_width=True, hide_index=True)

                    # Key insight for time-phased profile
                    peak_phase = max(sorted_phases, key=lambda x: x[1]['at_confidence'])
                    st.success(f"""
                    **📅 Phase Allocation Insight**: The {peak_phase[1]['name']} phase has the highest contingency
                    allocation at {peak_phase[1]['at_confidence']/1e6:.2f}M CHF ({peak_phase[1]['confidence_percentage']:.1f}% of total).
                    Plan cash flow reserves accordingly for this period.
                    """)

                except Exception as e:
                    st.error(f"Error processing time-phased data: {str(e)}")
                    st.info("Please ensure the risk register contains valid 'Crystallization Phase' and 'Phase Weight Distribution' columns.")

        with tab2:
            st.header("Risk Visualization - Matrix, Heatmap, Bubble & 3D Charts")

            # Add view selector
            view_type = st.radio(
                "Select Visualization Type:",
                ["Risk Matrix (Scatter)", "Risk Heatmap (Grid)", "Risk Bubble Chart", "3D Risk Landscape"],
                horizontal=True
            )
            
            col1, col2 = st.columns(2)
            
            if view_type == "Risk Matrix (Scatter)":
                with col1:
                    st.subheader("Initial Risk Matrix")
                    initial_matrix = create_risk_matrix(df, 'initial')
                    st.plotly_chart(initial_matrix, use_container_width=True)
                
                with col2:
                    st.subheader("Residual Risk Matrix")
                    residual_matrix = create_risk_matrix(df, 'residual')
                    st.plotly_chart(residual_matrix, use_container_width=True)
            
            elif view_type == "Risk Heatmap (Grid)":
                with col1:
                    st.subheader("Initial Risk Heatmap")
                    st.info("📊 Shows concentration and total value of risks in each likelihood-impact cell")
                    initial_heatmap = create_risk_heatmap(df, 'initial')
                    st.plotly_chart(initial_heatmap, use_container_width=True)
                
                with col2:
                    st.subheader("Residual Risk Heatmap")
                    st.info("📊 Shows risk concentration after mitigation measures")
                    residual_heatmap = create_risk_heatmap(df, 'residual')
                    st.plotly_chart(residual_heatmap, use_container_width=True)
                
                # Add summary statistics
                st.markdown("---")
                st.subheader("Heatmap Insights")
                
                # Calculate high-risk quadrant
                high_impact_high_prob_initial = df[(np.abs(df['Initial risk_Value']) > 1e6) & 
                                                    (df['Initial_Likelihood'] > 0.25)].shape[0]
                high_impact_high_prob_residual = df[(np.abs(df['Residual risk_Value']) > 1e6) & 
                                                     (df['Residual_Likelihood'] > 0.25)].shape[0]
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("High-Risk Quadrant (Initial)", 
                             high_impact_high_prob_initial,
                             help="Risks with impact >1M CHF and likelihood >25%")
                with col2:
                    st.metric("High-Risk Quadrant (Residual)", 
                             high_impact_high_prob_residual,
                             delta=high_impact_high_prob_residual - high_impact_high_prob_initial)
                with col3:
                    reduction = ((high_impact_high_prob_initial - high_impact_high_prob_residual) / 
                                high_impact_high_prob_initial * 100) if high_impact_high_prob_initial > 0 else 0
                    st.metric("High-Risk Reduction", f"{reduction:.1f}%")
            
            elif view_type == "Risk Bubble Chart":
                with col1:
                    st.subheader("Initial Risk Bubble Chart")
                    st.info("📊 Bubble size represents expected value of each risk")
                    initial_bubble = create_risk_bubble_chart(df, 'initial')
                    st.plotly_chart(initial_bubble, use_container_width=True)

                with col2:
                    st.subheader("Residual Risk Bubble Chart")
                    st.info("📊 Shows risk positioning after mitigation")
                    residual_bubble = create_risk_bubble_chart(df, 'residual')
                    st.plotly_chart(residual_bubble, use_container_width=True)

            else:  # 3D Risk Landscape
                st.subheader("3D Risk Landscape Visualization")
                st.info("""
                🌐 **3D Risk Landscape** provides an immersive view of your risk portfolio:
                - **X-axis**: Likelihood (0-100%)
                - **Y-axis**: Impact (M CHF)
                - **Z-axis**: Expected Value (M CHF)
                - **Red spheres**: Threats (positive EV)
                - **Green diamonds**: Opportunities (negative EV)
                - **Gray plane**: Zero EV threshold

                💡 **Tip**: Click and drag to rotate the 3D view. Use scroll to zoom.
                """)

                view_3d_type = st.radio(
                    "3D View Mode:",
                    ["Side-by-Side Comparison", "Initial Only", "Residual Only"],
                    horizontal=True,
                    key="3d_view_mode"
                )

                if view_3d_type == "Side-by-Side Comparison":
                    comparison_3d = create_3d_risk_comparison(df)
                    st.plotly_chart(comparison_3d, use_container_width=True)
                elif view_3d_type == "Initial Only":
                    initial_3d = create_3d_risk_surface(df, 'initial')
                    st.plotly_chart(initial_3d, use_container_width=True)
                else:
                    residual_3d = create_3d_risk_surface(df, 'residual')
                    st.plotly_chart(residual_3d, use_container_width=True)

                # Add 3D insights
                st.markdown("---")
                st.subheader("3D Landscape Insights")

                to_metrics = calculate_threat_opportunity_metrics(df)
                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    st.metric("Threats Above Zero Plane",
                             to_metrics['threat_count'],
                             help="Number of risks with positive Expected Value")
                with col2:
                    st.metric("Opportunities Below Zero Plane",
                             to_metrics['opportunity_count'],
                             help="Number of risks with negative Expected Value (benefits)")
                with col3:
                    # Count high-risk items (high EV, high likelihood)
                    high_risk_3d = len(df[(df['Initial_EV'] > 1e6) & (df['Initial_Likelihood'] > 0.5)])
                    st.metric("High-Risk Zone (>50% likelihood, >1M EV)",
                             high_risk_3d)
                with col4:
                    # Average height (EV) of threats
                    avg_threat_ev = to_metrics['threat_initial_ev'] / to_metrics['threat_count'] if to_metrics['threat_count'] > 0 else 0
                    st.metric("Avg Threat Height (EV)",
                             f"{avg_threat_ev/1e6:.2f}M CHF")

        with tab3:
            st.header("Enhanced Sensitivity Analysis")
            
            st.info("""
            🔍 **Sensitivity Analysis** identifies which risks contribute most to overall portfolio uncertainty.
            This helps prioritize risk management efforts on the risks that matter most.
            """)
            
            # Get sensitivity data
            sensitivity_df = st.session_state['sensitivity_df']
            
            # Key insights
            st.subheader("Key Findings")
            
            col1, col2, col3 = st.columns(3)
            
            # Top risk by variance
            top_risk = sensitivity_df.iloc[0]
            with col1:
                st.metric(
                    "Top Risk Driver",
                    f"Risk {top_risk['Risk ID']}",
                    f"{top_risk['Variance %']:.1f}% of variance"
                )
            
            # Number of risks for 80% variance
            risks_80 = (sensitivity_df['Cumulative %'] <= 80).sum()
            with col2:
                st.metric(
                    "Risks Driving 80% Variance",
                    f"{risks_80} risks",
                    f"{risks_80/len(df)*100:.1f}% of total"
                )
            
            # Total variance of top 5
            top5_variance = sensitivity_df.head(5)['Variance %'].sum()
            with col3:
                st.metric(
                    "Top 5 Risks Contribution",
                    f"{top5_variance:.1f}%",
                    "of total variance"
                )
            
            st.markdown("---")
            
            # Enhanced Tornado Chart
            st.subheader("Variance Contribution (Tornado Chart)")
            tornado_enhanced = create_enhanced_tornado_chart(sensitivity_df, top_n=15)
            st.plotly_chart(tornado_enhanced, use_container_width=True)
            
            st.markdown("---")
            
            # Pareto Chart
            st.subheader("Pareto Analysis (80/20 Rule)")
            st.write("""
            The **Pareto principle** suggests that roughly 80% of effects come from 20% of causes.
            This chart identifies the "vital few" risks that drive most of the uncertainty.
            """)
            pareto_chart = create_pareto_chart(sensitivity_df, top_n=20)
            st.plotly_chart(pareto_chart, use_container_width=True)
            
            # Find 80% threshold
            risks_for_80 = sensitivity_df[sensitivity_df['Cumulative %'] <= 80]
            st.success(f"""
            📊 **Pareto Insight**: {len(risks_for_80)} risks ({len(risks_for_80)/len(df)*100:.1f}% of total) 
            account for 80% of the portfolio variance.
            
            **Recommendation**: Focus intensive risk management efforts on these {len(risks_for_80)} critical risks.
            """)
            
            st.markdown("---")
            
            # Detailed sensitivity table
            st.subheader("Detailed Sensitivity Analysis")
            
            # Prepare display dataframe
            display_sensitivity = sensitivity_df[['Risk ID', 'Risk Description', 'Variance %', 
                                                   'Cumulative %', 'Mean Contribution', 'Expected Value']].copy()
            display_sensitivity['Mean Contribution (M CHF)'] = (display_sensitivity['Mean Contribution'] / 1e6).round(2)
            display_sensitivity['Expected Value (M CHF)'] = (display_sensitivity['Expected Value'] / 1e6).round(2)
            display_sensitivity['Variance %'] = display_sensitivity['Variance %'].round(2)
            display_sensitivity['Cumulative %'] = display_sensitivity['Cumulative %'].round(1)
            
            display_sensitivity = display_sensitivity[['Risk ID', 'Risk Description', 'Variance %', 
                                                       'Cumulative %', 'Mean Contribution (M CHF)', 
                                                       'Expected Value (M CHF)']]
            
            st.dataframe(
                display_sensitivity,
                use_container_width=True,
                height=400,
                hide_index=True
            )
            
            # Export sensitivity analysis
            csv_sensitivity = display_sensitivity.to_csv(index=False)
            st.download_button(
                label="📥 Download Sensitivity Analysis (CSV)",
                data=csv_sensitivity,
                file_name=f"sensitivity_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
        
        with tab4:
            st.header("Monte Carlo Simulation Results")
            
            st.info(f"📊 Performed {n_simulations:,} simulations to model probabilistic risk exposure")
            
            # Box plot
            st.subheader("Risk Exposure Variability")
            box_fig = create_box_plot(initial_results, residual_results)
            st.plotly_chart(box_fig, use_container_width=True)
            
            # Histograms
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Initial Risk Distribution")
                hist_initial = create_histogram(initial_results, initial_stats, 'initial')
                st.plotly_chart(hist_initial, use_container_width=True)
            
            with col2:
                st.subheader("Residual Risk Distribution")
                hist_residual = create_histogram(residual_results, residual_stats, 'residual')
                st.plotly_chart(hist_residual, use_container_width=True)
            
            # CDF plots
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Initial Risk CDF")
                cdf_initial = create_cdf_plot(initial_results, initial_stats, 'initial', confidence_level)
                st.plotly_chart(cdf_initial, use_container_width=True)

            with col2:
                st.subheader("Residual Risk CDF")
                cdf_residual = create_cdf_plot(residual_results, residual_stats, 'residual', confidence_level)
                st.plotly_chart(cdf_residual, use_container_width=True)
        
        with tab5:
            st.header("Mitigation Cost-Benefit Analysis")
            
            # Filter risks with mitigation measures
            df_with_measures = df_with_roi[df_with_roi['Cost of Measures_Value'] > 0].copy()
            
            if len(df_with_measures) > 0:
                # Summary metrics
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    total_reduction = df_with_measures['Risk_Reduction'].sum()
                    st.metric("Total Risk Reduction", f"{total_reduction/1e6:.2f}M CHF")
                
                with col2:
                    total_cost = df_with_measures['Cost of Measures_Value'].sum()
                    st.metric("Total Mitigation Cost", f"{total_cost/1e6:.2f}M CHF")
                
                with col3:
                    net_benefit = total_reduction - total_cost
                    st.metric("Net Benefit", f"{net_benefit/1e6:.2f}M CHF")
                
                st.markdown("---")
                
                # ROI chart
                st.subheader("Return on Investment by Risk")
                df_roi_sorted = df_with_measures.nlargest(20, 'ROI')
                
                fig_roi = px.bar(df_roi_sorted,
                                x='ROI',
                                y='Risk Description',
                                orientation='h',
                                title='Top 20 Mitigation Measures by ROI (%)',
                                color='ROI',
                                color_continuous_scale='RdYlGn',
                                labels={'ROI': 'ROI (%)'})
                
                fig_roi.update_layout(height=600, yaxis={'categoryorder': 'total ascending'})
                st.plotly_chart(fig_roi, use_container_width=True)
                
                # Benefit-Cost Ratio chart
                st.subheader("Benefit-Cost Ratio Analysis")
                df_bc_sorted = df_with_measures.nlargest(20, 'BC_Ratio')
                
                fig_bc = px.scatter(df_bc_sorted,
                                   x='Cost of Measures_Value',
                                   y='Risk_Reduction',
                                   size='BC_Ratio',
                                   color='BC_Ratio',
                                   hover_data=['Risk ID', 'Risk Description'],
                                   title='Risk Reduction vs Mitigation Cost',
                                   labels={
                                       'Cost of Measures_Value': 'Mitigation Cost (CHF)',
                                       'Risk_Reduction': 'Risk Reduction (CHF)'
                                   },
                                   color_continuous_scale='Viridis')
                
                # Add break-even line
                max_val = max(df_bc_sorted['Cost of Measures_Value'].max(), 
                             df_bc_sorted['Risk_Reduction'].max())
                fig_bc.add_trace(go.Scatter(x=[0, max_val], y=[0, max_val],
                                           mode='lines',
                                           name='Break-even line',
                                           line=dict(dash='dash', color='red')))
                
                st.plotly_chart(fig_bc, use_container_width=True)
                
                # Detailed table
                st.subheader("Detailed Cost-Benefit Analysis")
                
                display_cols = ['Risk ID', 'Risk Description', 'Initial_EV', 'Residual_EV', 
                               'Risk_Reduction', 'Cost of Measures_Value', 'ROI', 'BC_Ratio']
                
                display_df = df_with_measures[display_cols].copy()
                display_df.columns = ['Risk ID', 'Risk Description', 'Initial EV (CHF)', 
                                     'Residual EV (CHF)', 'Risk Reduction (CHF)', 
                                     'Mitigation Cost (CHF)', 'ROI (%)', 'B/C Ratio']
                
                # Format numbers
                for col in ['Initial EV (CHF)', 'Residual EV (CHF)', 'Risk Reduction (CHF)', 'Mitigation Cost (CHF)']:
                    display_df[col] = display_df[col].apply(lambda x: f"{x:,.0f}")
                
                display_df['ROI (%)'] = display_df['ROI (%)'].apply(lambda x: f"{x:.1f}")
                display_df['B/C Ratio'] = display_df['B/C Ratio'].apply(lambda x: f"{x:.2f}")

                st.dataframe(display_df, use_container_width=True, hide_index=True)

                # Sankey Diagram Section
                st.markdown("---")
                st.subheader("🔄 Risk Mitigation Flow (Sankey Diagram)")

                st.markdown("""
                The Sankey diagram below visualizes how risk flows through the mitigation process:
                - **Threats** (red tones): Potential costs/losses that flow from initial → reduced → residual
                - **Opportunities** (green/teal tones): Potential benefits that offset the threat exposure
                - **Net Exposure**: Threat EV minus Opportunity EV (shown in purple)
                """)

                # Sankey view selector
                sankey_view = st.radio(
                    "Select View:",
                    options=["Threats & Opportunities", "By Schedule Impact"],
                    horizontal=True,
                    help="Choose how to visualize the risk flow breakdown"
                )

                if sankey_view == "Threats & Opportunities":
                    sankey_fig = create_risk_sankey(df_with_roi, initial_stats, residual_stats, confidence_level)
                else:
                    sankey_fig = create_risk_sankey_detailed(df_with_roi)

                st.plotly_chart(sankey_fig, use_container_width=True)

                # Summary statistics below Sankey - now using threat/opportunity breakdown
                st.markdown("##### Flow Summary: Threats vs Opportunities")

                # Calculate threat/opportunity metrics
                to_flow_metrics = calculate_threat_opportunity_metrics(df_with_roi)

                flow_col1, flow_col2, flow_col3, flow_col4 = st.columns(4)

                with flow_col1:
                    st.metric(
                        "Threat EV (Initial)",
                        f"{to_flow_metrics['threat_initial_ev']/1e6:.2f}M CHF",
                        help=f"{to_flow_metrics['threat_count']} threats"
                    )
                    st.metric(
                        "Threat EV (Residual)",
                        f"{to_flow_metrics['threat_residual_ev']/1e6:.2f}M CHF"
                    )

                with flow_col2:
                    st.metric(
                        "Opportunity EV (Initial)",
                        f"{to_flow_metrics['opportunity_initial_ev']/1e6:.2f}M CHF",
                        help=f"{to_flow_metrics['opportunity_count']} opportunities (negative = benefit)"
                    )
                    st.metric(
                        "Opportunity EV (Residual)",
                        f"{to_flow_metrics['opportunity_residual_ev']/1e6:.2f}M CHF"
                    )

                with flow_col3:
                    st.metric(
                        "Net Exposure (Initial)",
                        f"{to_flow_metrics['net_initial_exposure']/1e6:.2f}M CHF",
                        help="Threat EV + Opportunity EV"
                    )
                    st.metric(
                        "Net Exposure (Residual)",
                        f"{to_flow_metrics['net_residual_exposure']/1e6:.2f}M CHF"
                    )

                with flow_col4:
                    threat_reduction_pct = (to_flow_metrics['threat_reduction'] / to_flow_metrics['threat_initial_ev'] * 100) if to_flow_metrics['threat_initial_ev'] > 0 else 0
                    st.metric(
                        "Threat Reduction",
                        f"{to_flow_metrics['threat_reduction']/1e6:.2f}M CHF",
                        delta=f"-{threat_reduction_pct:.1f}%"
                    )
                    mitigation_efficiency = (to_flow_metrics['threat_reduction'] / total_cost * 100) if total_cost > 0 else 0
                    st.metric("Mitigation Efficiency", f"{mitigation_efficiency:.0f}%",
                             help="Threat reduced per CHF spent")

            else:
                st.warning("No risks with mitigation measures found in the register.")

        with tab6:
            st.header("🔮 What-If Scenario Analysis")

            st.markdown("""
            Explore how changes to individual risks affect your overall portfolio exposure.
            Create scenarios by adjusting risk likelihood and/or impact values, then compare against the baseline.
            """)

            # Initialize scenario adjustments in session state
            if 'scenario_adjustments' not in st.session_state:
                st.session_state.scenario_adjustments = {}
            if 'scenario_results' not in st.session_state:
                st.session_state.scenario_results = None

            # Scenario configuration
            col_name, col_type = st.columns([2, 1])

            with col_name:
                scenario_name = st.text_input(
                    "Scenario Name",
                    value="Custom Scenario",
                    help="Give your scenario a descriptive name"
                )

            with col_type:
                scenario_type = st.selectbox(
                    "Scenario Type",
                    options=["Custom", "Optimistic (-20%)", "Pessimistic (+20%)", "Best Case (-50%)", "Worst Case (+50%)"],
                    help="Quick presets or create a custom scenario"
                )

            st.markdown("---")

            # Handle preset scenarios
            if scenario_type != "Custom":
                st.subheader("📋 Preset Scenario")

                preset_adjustments = {}
                if scenario_type == "Optimistic (-20%)":
                    factor = 0.8
                    description = "All risks reduced by 20%"
                elif scenario_type == "Pessimistic (+20%)":
                    factor = 1.2
                    description = "All risks increased by 20%"
                elif scenario_type == "Best Case (-50%)":
                    factor = 0.5
                    description = "All risks reduced by 50%"
                else:  # Worst Case
                    factor = 1.5
                    description = "All risks increased by 50%"

                st.info(f"**{scenario_type}**: {description}")

                # Apply factor to all risks
                for _, row in df.iterrows():
                    risk_id = row['Risk ID']
                    preset_adjustments[risk_id] = {
                        'likelihood': min(row['Initial_Likelihood'] * factor, 1.0),
                        'impact': row['Initial risk_Value'] * factor
                    }

                if st.button("Apply Preset & Run Simulation", type="primary", key="preset_btn"):
                    with st.spinner("Running scenario simulation..."):
                        scenario_results, _ = run_scenario_monte_carlo(
                            df, preset_adjustments, n_simulations, 'initial'
                        )
                        scenario_stats = calculate_statistics(scenario_results)

                        st.session_state.scenario_results = {
                            'name': f"{scenario_type}",
                            'results': scenario_results,
                            'stats': scenario_stats,
                            'adjustments': preset_adjustments
                        }
                        st.success("Scenario simulation complete!")
                        st.rerun()

            else:
                # Custom scenario builder
                st.subheader("📝 Custom Scenario Builder")

                # Risk selection
                st.markdown("**Select risks to modify:**")

                # Get top risks by expected value for quick selection
                top_risks = df.nlargest(10, 'Initial_EV')[['Risk ID', 'Risk Description', 'Initial risk_Value', 'Initial_Likelihood']].copy()

                selected_risks = st.multiselect(
                    "Choose risks to adjust",
                    options=df['Risk ID'].tolist(),
                    default=list(st.session_state.scenario_adjustments.keys()),
                    help="Select one or more risks to modify in this scenario"
                )

                # Display adjustment controls for selected risks
                if selected_risks:
                    st.markdown("**Adjust risk parameters:**")

                    new_adjustments = {}

                    for risk_id in selected_risks:
                        risk_row = df[df['Risk ID'] == risk_id].iloc[0]

                        with st.expander(f"📌 {risk_id}: {risk_row['Risk Description'][:60]}...", expanded=True):
                            col1, col2, col3 = st.columns([1, 1, 1])

                            # Get current adjustment or original value
                            current_adjustment = st.session_state.scenario_adjustments.get(risk_id, {})
                            orig_likelihood = risk_row['Initial_Likelihood']
                            orig_impact = risk_row['Initial risk_Value']

                            with col1:
                                st.markdown("**Original Values:**")
                                st.write(f"Likelihood: {orig_likelihood*100:.1f}%")
                                st.write(f"Impact: {orig_impact/1e6:.2f}M CHF")
                                st.write(f"EV: {(orig_likelihood * orig_impact)/1e6:.2f}M CHF")

                            with col2:
                                new_likelihood = st.slider(
                                    f"New Likelihood (%)",
                                    min_value=0,
                                    max_value=100,
                                    value=int(current_adjustment.get('likelihood', orig_likelihood) * 100),
                                    key=f"lik_{risk_id}"
                                ) / 100

                            with col3:
                                new_impact = st.number_input(
                                    f"New Impact (M CHF)",
                                    min_value=0.0,
                                    max_value=float(df['Initial risk_Value'].max() / 1e6 * 2),
                                    value=float(current_adjustment.get('impact', orig_impact) / 1e6),
                                    step=0.1,
                                    key=f"imp_{risk_id}"
                                ) * 1e6

                            # Store adjustment
                            new_adjustments[risk_id] = {
                                'likelihood': new_likelihood,
                                'impact': new_impact
                            }

                            # Show change summary
                            new_ev = new_likelihood * new_impact
                            orig_ev = orig_likelihood * orig_impact
                            ev_change = new_ev - orig_ev
                            ev_change_pct = (ev_change / orig_ev * 100) if orig_ev > 0 else 0

                            if ev_change < 0:
                                st.success(f"📉 EV Change: {ev_change/1e6:+.2f}M CHF ({ev_change_pct:+.1f}%)")
                            elif ev_change > 0:
                                st.error(f"📈 EV Change: {ev_change/1e6:+.2f}M CHF ({ev_change_pct:+.1f}%)")
                            else:
                                st.info("No change in Expected Value")

                    # Update session state
                    st.session_state.scenario_adjustments = new_adjustments

                    # Run simulation button
                    st.markdown("---")
                    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])

                    with col_btn2:
                        if st.button("🚀 Run Scenario Simulation", type="primary", use_container_width=True):
                            with st.spinner("Running Monte Carlo simulation for scenario..."):
                                scenario_results, _ = run_scenario_monte_carlo(
                                    df, new_adjustments, n_simulations, 'initial'
                                )
                                scenario_stats = calculate_statistics(scenario_results)

                                st.session_state.scenario_results = {
                                    'name': scenario_name,
                                    'results': scenario_results,
                                    'stats': scenario_stats,
                                    'adjustments': new_adjustments
                                }
                                st.success("Scenario simulation complete!")
                                st.rerun()

                else:
                    st.info("👆 Select one or more risks above to begin building your scenario.")

                    # Show top risks as suggestions
                    st.markdown("**💡 Suggested risks to explore (Top 10 by Expected Value):**")
                    suggestion_df = top_risks.copy()
                    suggestion_df['Initial risk_Value'] = suggestion_df['Initial risk_Value'].apply(lambda x: f"{x/1e6:.2f}M CHF")
                    suggestion_df['Initial_Likelihood'] = suggestion_df['Initial_Likelihood'].apply(lambda x: f"{x*100:.1f}%")
                    suggestion_df.columns = ['Risk ID', 'Description', 'Impact', 'Likelihood']
                    st.dataframe(suggestion_df, use_container_width=True, hide_index=True)

            # Display results if available
            st.markdown("---")
            st.subheader("📊 Scenario Comparison Results")

            if st.session_state.scenario_results is not None:
                scenario_data = st.session_state.scenario_results
                scenario_stats = scenario_data['stats']
                scenario_results = scenario_data['results']
                scenario_display_name = scenario_data['name']

                # Summary metrics
                col1, col2, col3, col4 = st.columns(4)

                baseline_conf_val = get_confidence_value(initial_stats, confidence_level)
                scenario_conf_val = get_confidence_value(scenario_stats, confidence_level)
                change = scenario_conf_val - baseline_conf_val
                change_pct = (change / baseline_conf_val * 100) if baseline_conf_val > 0 else 0

                with col1:
                    st.metric(
                        f"Baseline ({confidence_level})",
                        f"{baseline_conf_val/1e6:.2f}M CHF"
                    )

                with col2:
                    st.metric(
                        f"Scenario ({confidence_level})",
                        f"{scenario_conf_val/1e6:.2f}M CHF",
                        delta=f"{change/1e6:+.2f}M CHF"
                    )

                with col3:
                    st.metric(
                        "Change",
                        f"{change_pct:+.1f}%",
                        delta="Reduced" if change < 0 else "Increased",
                        delta_color="normal" if change < 0 else "inverse"
                    )

                with col4:
                    risks_modified = len(scenario_data.get('adjustments', {}))
                    st.metric("Risks Modified", risks_modified)

                # Comparison charts
                chart_col1, chart_col2 = st.columns(2)

                with chart_col1:
                    comparison_fig = create_scenario_comparison_chart(
                        initial_stats, scenario_stats, scenario_display_name, confidence_level
                    )
                    st.plotly_chart(comparison_fig, use_container_width=True)

                with chart_col2:
                    cdf_fig = create_scenario_cdf_comparison(
                        initial_results, scenario_results, initial_stats, scenario_stats,
                        scenario_display_name, confidence_level
                    )
                    st.plotly_chart(cdf_fig, use_container_width=True)

                # Modified risks table
                if scenario_data.get('adjustments'):
                    st.subheader("📋 Modified Risks Summary")
                    impact_table = create_scenario_impact_table(df, scenario_data['adjustments'], 'initial')
                    if impact_table is not None:
                        st.dataframe(impact_table, use_container_width=True, hide_index=True)

                # Clear results button
                col_clear1, col_clear2, col_clear3 = st.columns([1, 1, 1])
                with col_clear2:
                    if st.button("🗑️ Clear Scenario Results", use_container_width=True):
                        st.session_state.scenario_results = None
                        st.session_state.scenario_adjustments = {}
                        st.rerun()

            else:
                st.info("👆 Configure and run a scenario above to see comparison results here.")

                # Show quick summary of baseline
                st.markdown("**Current Baseline Summary:**")
                baseline_summary = pd.DataFrame({
                    'Metric': ['Mean', 'P50 (Median)', 'P80', 'P95'],
                    'Initial Risk (M CHF)': [
                        f"{initial_stats['mean']/1e6:.2f}",
                        f"{initial_stats['p50']/1e6:.2f}",
                        f"{initial_stats['p80']/1e6:.2f}",
                        f"{initial_stats['p95']/1e6:.2f}"
                    ]
                })
                st.dataframe(baseline_summary, use_container_width=True, hide_index=True)

        with tab7:
            st.header("Risk Register")

            # Display portfolio-level confidence information
            st.info(f"📊 **Portfolio Analysis using {confidence_level} confidence level** | "
                   f"Total Portfolio Risk ({confidence_level}): Initial = {initial_confidence_value/1e6:.2f}M CHF, "
                   f"Residual = {residual_confidence_value/1e6:.2f}M CHF")

            # Add filters
            st.subheader("Filters")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                schedule_filter = st.multiselect("Schedule Impact", 
                                                options=[True, False],
                                                default=[True, False],
                                                format_func=lambda x: "Yes" if x else "No")
            
            with col2:
                risk_threshold = st.slider("Min Initial Risk (M CHF)", 
                                          min_value=0.0,
                                          max_value=float(df['Initial risk_Value'].max()/1e6),
                                          value=0.0)
            
            with col3:
                likelihood_threshold = st.slider("Min Likelihood (%)",
                                                min_value=0,
                                                max_value=100,
                                                value=0)
            
            # Apply filters
            filtered_df = df_with_roi[
                (df_with_roi['Schedule_Impact'].isin(schedule_filter)) &
                (df_with_roi['Initial risk_Value'] >= risk_threshold * 1e6) &
                (df_with_roi['Initial_Likelihood'] >= likelihood_threshold / 100)
            ].copy()
            
            st.write(f"Showing {len(filtered_df)} of {len(df)} risks")
            
            # Display table
            display_cols = ['Risk ID', 'Risk Description', 'Initial risk_Value', 'Initial_Likelihood',
                           'Residual risk_Value', 'Residual_Likelihood', 'Cost of Measures_Value',
                           'Schedule_Impact', 'ROI', 'BC_Ratio']
            
            display_df = filtered_df[display_cols].copy()
            display_df.columns = ['Risk ID', 'Description', 'Initial Risk (CHF)', 'Initial Likelihood',
                                 'Residual Risk (CHF)', 'Residual Likelihood', 'Mitigation Cost (CHF)',
                                 'Schedule Impact', 'ROI (%)', 'B/C Ratio']
            
            # Format
            display_df['Initial Risk (CHF)'] = display_df['Initial Risk (CHF)'].apply(lambda x: f"{x:,.0f}")
            display_df['Residual Risk (CHF)'] = display_df['Residual Risk (CHF)'].apply(lambda x: f"{x:,.0f}")
            display_df['Mitigation Cost (CHF)'] = display_df['Mitigation Cost (CHF)'].apply(lambda x: f"{x:,.0f}")
            display_df['Initial Likelihood'] = display_df['Initial Likelihood'].apply(lambda x: f"{x:.1%}")
            display_df['Residual Likelihood'] = display_df['Residual Likelihood'].apply(lambda x: f"{x:.1%}")
            display_df['Schedule Impact'] = display_df['Schedule Impact'].apply(lambda x: "Yes" if x else "No")
            display_df['ROI (%)'] = display_df['ROI (%)'].apply(lambda x: f"{x:.1f}")
            display_df['B/C Ratio'] = display_df['B/C Ratio'].apply(lambda x: f"{x:.2f}")
            
            st.dataframe(display_df, use_container_width=True, hide_index=True, height=600)

        with tab8:
            st.header("Export Results")
            
            st.write("Download your risk assessment results in various formats:")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Export to Excel
                st.subheader("📊 Excel Export")
                
                # Create Excel file
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    # Summary statistics
                    summary_data = {
                        'Metric': ['Mean', 'Median (P50)', 'P80', 'P95', 'Std Dev'],
                        'Initial Risk (CHF)': [
                            initial_stats['mean'],
                            initial_stats['p50'],
                            initial_stats['p80'],
                            initial_stats['p95'],
                            initial_stats['std']
                        ],
                        'Residual Risk (CHF)': [
                            residual_stats['mean'],
                            residual_stats['p50'],
                            residual_stats['p80'],
                            residual_stats['p95'],
                            residual_stats['std']
                        ]
                    }
                    pd.DataFrame(summary_data).to_excel(writer, sheet_name='Summary', index=False)
                    
                    # Full risk register with ROI
                    df_with_roi.to_excel(writer, sheet_name='Risk Register', index=False)
                    
                    # Monte Carlo results sample
                    mc_results = pd.DataFrame({
                        'Simulation': range(1, len(initial_results) + 1),
                        'Initial Risk (CHF)': initial_results,
                        'Residual Risk (CHF)': residual_results
                    })
                    mc_results.to_excel(writer, sheet_name='Monte Carlo Results', index=False)
                
                excel_data = output.getvalue()
                
                st.download_button(
                    label="Download Excel Report",
                    data=excel_data,
                    file_name=f"risk_assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
            with col2:
                # Export to CSV
                st.subheader("📄 CSV Export")
                
                csv = df_with_roi.to_csv(index=False)
                
                st.download_button(
                    label="Download Risk Register (CSV)",
                    data=csv,
                    file_name=f"risk_register_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
                
                # Export Monte Carlo results
                mc_csv = pd.DataFrame({
                    'Simulation': range(1, len(initial_results) + 1),
                    'Initial_Risk_CHF': initial_results,
                    'Residual_Risk_CHF': residual_results
                }).to_csv(index=False)
                
                st.download_button(
                    label="Download Monte Carlo Results (CSV)",
                    data=mc_csv,
                    file_name=f"monte_carlo_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
            
            st.markdown("---")

            # Professional DOCX Report
            st.subheader("📄 Professional DOCX Report")
            st.write("Generate a comprehensive Word document with executive summary, embedded charts, and detailed analysis.")

            col_docx1, col_docx2, col_docx3 = st.columns([1, 2, 1])

            with col_docx2:
                if st.button("🎨 Generate Professional DOCX Report", type="primary", use_container_width=True):
                    with st.spinner("📄 Generating comprehensive DOCX report... This may take 30-60 seconds"):
                        try:
                            # Get sensitivity data
                            sensitivity_df = st.session_state.get('sensitivity_df')

                            if sensitivity_df is None:
                                st.error("❌ Sensitivity analysis data not found. Please run the simulation first.")
                            else:
                                # Generate the DOCX report
                                docx_report = generate_docx_report(
                                    initial_stats=initial_stats,
                                    residual_stats=residual_stats,
                                    df=df,
                                    df_with_roi=df_with_roi,
                                    sensitivity_df=sensitivity_df,
                                    initial_results=initial_results,
                                    residual_results=residual_results,
                                    n_simulations=n_simulations,
                                    confidence_level=confidence_level,
                                    current_df=current_df
                                )

                                st.success("✅ DOCX report generated successfully!")

                                st.download_button(
                                    label="📥 Download Professional Report (DOCX)",
                                    data=docx_report,
                                    file_name=f"risk_assessment_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )

                        except Exception as e:
                            st.error(f"❌ Error generating DOCX report: {str(e)}")
                            st.write("Please ensure all required libraries are installed: `pip install python-docx`")
                            import traceback
                            st.code(traceback.format_exc())

            st.info("""
                **📋 Report Contents:**
                - ✅ Professional cover page with branding
                - ✅ Executive summary with key metrics and top risks
                - ✅ Risk portfolio overview with embedded charts
                - ✅ Monte Carlo simulation results with CDF and histograms
                - ✅ Sensitivity analysis with Pareto chart
                - ✅ Mitigation cost-benefit analysis with ROI charts
                - ✅ Full risk register appendix
                - ✅ Methodology explanation appendix
            """)

            st.markdown("---")

            # Summary report
            st.subheader("📋 Summary Report")

            # Get selected confidence values
            initial_selected = get_confidence_value(initial_stats, confidence_level)
            residual_selected = get_confidence_value(residual_stats, confidence_level)
            risk_reduction_selected = initial_selected - residual_selected
            risk_reduction_pct = (risk_reduction_selected / initial_selected * 100)
            net_benefit_selected = risk_reduction_selected - df['Cost of Measures_Value'].sum()

            # Mark selected confidence level in the report
            p50_marker = " ★ (SELECTED)" if confidence_level == "P50" else ""
            p80_marker = " ★ (SELECTED)" if confidence_level == "P80" else ""
            p95_marker = " ★ (SELECTED)" if confidence_level == "P95" else ""

            report = f"""
# Risk Assessment Report
**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Monte Carlo Simulations:** {n_simulations:,}
**Selected Confidence Level:** {confidence_level}

## Executive Summary

### Primary Risk Assessment ({confidence_level} - Selected Confidence Level)
- **Initial Risk Exposure:** {initial_selected/1e6:.2f} Million CHF
- **Residual Risk Exposure (After Mitigation):** {residual_selected/1e6:.2f} Million CHF
- **Total Risk Reduction:** {risk_reduction_selected/1e6:.2f} Million CHF ({risk_reduction_pct:.1f}%)
- **Total Mitigation Investment:** {df['Cost of Measures_Value'].sum()/1e6:.2f} Million CHF
- **Net Benefit:** {net_benefit_selected/1e6:.2f} Million CHF

### Initial Risk Exposure (All Confidence Levels)
- **P50 (Median):** {initial_stats['p50']/1e6:.2f} Million CHF{p50_marker}
- **P80:** {initial_stats['p80']/1e6:.2f} Million CHF{p80_marker}
- **P95:** {initial_stats['p95']/1e6:.2f} Million CHF{p95_marker}

### Residual Risk Exposure (After Mitigation)
- **P50 (Median):** {residual_stats['p50']/1e6:.2f} Million CHF{p50_marker}
- **P80:** {residual_stats['p80']/1e6:.2f} Million CHF{p80_marker}
- **P95:** {residual_stats['p95']/1e6:.2f} Million CHF{p95_marker}

### Key Findings
- Total number of identified risks: {len(df)}
- Risks with schedule impact: {df['Schedule_Impact'].sum()}
- Risks with mitigation measures: {len(df[df['Cost of Measures_Value'] > 0])}

## Recommendations
Based on the Monte Carlo simulation with **{confidence_level}** confidence level, the project should reserve at least **{residual_selected/1e6:.2f} Million CHF**
for residual risk exposure after implementing all planned mitigation measures.

The selected {confidence_level} confidence level indicates that there is a {confidence_level.replace('P', '')}% probability that
the actual risk exposure will be at or below this amount.
"""
            
            st.markdown(report)
            
            st.download_button(
                label="Download Summary Report (MD)",
                data=report,
                file_name=f"risk_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                mime="text/markdown"
            )
    
    else:
        # Initial state - show instructions
        st.info("👈 Click **'Run Simulation'** in the sidebar to begin risk analysis")
        
        st.markdown("""
        ## How to Use This Tool
        
        1. **Upload your risk register** (CSV format) or use the default dataset
        2. **Configure simulation parameters** in the sidebar:
           - Number of Monte Carlo simulations (higher = more accurate)
           - Confidence level for reporting
        3. **Run the simulation** to generate probabilistic risk analysis
        4. **Explore the results** across multiple tabs:
           - **Dashboard**: Overview and key metrics
           - **Risk Matrix**: Visual risk mapping
           - **Monte Carlo Results**: Statistical distributions and confidence intervals
           - **Cost-Benefit**: ROI analysis for mitigation measures
           - **Risk Register**: Detailed risk data with filters
           - **Export**: Download results in Excel/CSV formats
        
        ### Features
        - ✅ Monte Carlo simulation with 1,000-100,000 iterations
        - ✅ P50/P80/P95 confidence intervals
        - ✅ Interactive risk matrices (Initial vs Residual)
        - ✅ Cost-benefit and ROI analysis
        - ✅ Distribution plots, CDF, and box plots
        - ✅ Tornado charts for sensitivity analysis
        - ✅ Export to Excel and CSV
        
        ### Risk Register Format
        Your CSV should contain these columns:
        - Risk ID
        - Risk Description
        - Initial risk (with currency)
        - Likelihood (as percentage)
        - Residual risk
        - Residual likelihood
        - Cost of Measures
        - Schedule Impact (Yes/No)
        """)

if __name__ == "__main__":
    main()
